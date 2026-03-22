import os
import pdfplumber
import logging
import hashlib
import json
import nltk
import sqlite3
from nltk.tokenize import sent_tokenize
from flask import Flask, request, jsonify, render_template, Response, stream_with_context, make_response, session, redirect, url_for, send_file
from dotenv import load_dotenv
from authlib.integrations.flask_client import OAuth
from functools import wraps

# Windows compatibility fix: Handle missing readline module
import sys
if sys.platform == 'win32':
    try:
        import readline
    except ImportError:
        try:
            # Try pyreadline3 as Windows alternative
            import pyreadline3 as readline
            sys.modules['readline'] = readline
        except ImportError:
            # Create a minimal dummy readline module if pyreadline3 is not available
            class DummyReadline:
                @staticmethod
                def add_history(*args, **kwargs):
                    pass
                @staticmethod
                def write_history_file(*args, **kwargs):
                    pass
            sys.modules['readline'] = DummyReadline()

from pinecone import Pinecone, ServerlessSpec
from langchain_community.vectorstores import Pinecone as PineconeVectorStore
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from rank_bm25 import BM25Okapi
from functools import lru_cache
import re
import pandas as pd
import warnings
import google.generativeai as genai
from openai import OpenAI
from groq import Groq
from docx import Document
import uuid
from werkzeug.utils import secure_filename
from datetime import datetime
from langchain_text_splitters import RecursiveCharacterTextSplitter
import asyncio
import aiohttp
from concurrent.futures import ThreadPoolExecutor
from asgiref.wsgi import WsgiToAsgi
import time
from io import BytesIO
from evaluation_pdf import build_evaluation_pdf_bytes
from handbook_pdf import build_handbook_pdf_bytes
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, KeepTogether
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.lib import colors
import markdown
from html import unescape
try:
    from PyPDF2 import PdfMerger, PdfReader
    PDF_MERGE_AVAILABLE = True
except ImportError:
    PDF_MERGE_AVAILABLE = False
    logging.warning("PyPDF2 not available. PDF merging will not work.")


# Suppress specific warnings
warnings.filterwarnings("ignore", category=UserWarning, message="WARNING! top_p is not default parameter.")
warnings.filterwarnings("ignore", category=UserWarning, message="WARNING! presence_penalty is not default parameter.")
warnings.filterwarnings("ignore", category=UserWarning, message="WARNING! frequency_penalty is not default parameter.")

# Load environment variables
# Get the directory where this script is located
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ENV_PATH = os.path.join(BASE_DIR, '.env')
# Load .env file, override existing env vars
load_dotenv(ENV_PATH, override=True)

# Configuration
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "").strip()
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
PINECONE_INDEX_NAME = "hr-knowledge-base"
POLICIES_FOLDER = "HR_docs/"
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx'}

# Google OAuth Configuration
GOOGLE_CLIENT_ID = os.getenv("GOOGLE_CLIENT_ID", "").strip().strip('"')
GOOGLE_CLIENT_SECRET = os.getenv("GOOGLE_CLIENT_SECRET", "").strip().strip('"')
SECRET_KEY = os.getenv("FLASK_SECRET_KEY", os.getenv("SECRET_KEY", "your-secret-key-change-in-production-12345"))

# Model Selection Configuration
MODEL_PROVIDER = os.getenv("MODEL_PROVIDER", "gemini").lower()  # Options: "gemini", "openai", or "groq"
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")  # Options: gpt-4o, gpt-4o-mini, gpt-4-turbo, gpt-3.5-turbo
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")  # Options: gemini-2.5-flash, gemini-2.5-pro, gemini-2.0-flash
GROQ_MODEL = os.getenv("GROQ_MODEL", "openai/gpt-oss-120b")  # Options: openai/gpt-oss-120b, llama-3.3-70b-versatile, llama-3.1-70b-versatile
GROQ_REASONING_EFFORT = os.getenv("GROQ_REASONING_EFFORT", "high")  # Options: low, medium, high (for reasoning models)

# Add this dictionary after imports
ACRONYM_MAP = {
    "wfh": "work from home policy",
    "pto": "paid time off policy",
    "loa": "leave of absence policy",
    "nda": "non-disclosure agreement",
    "od": "on duty policy",
    "hrbp": "human resources business partner",
    "kra": "KRA Policy - Promoting Transparency",
    "regularization": "Time change Request/ Regularization",
    "regularisation": "Time change Request/ Regularization",
    "posh": "Policy On Prevention of Sexual Harassment",
    "appraisal": "PERFORMANCE APPRAISAL & PROMOTION POLICY",
    "promotion": "PERFORMANCE APPRAISAL & PROMOTION POLICY",
    "prep": "Performance Review & Enhancement Program",
    "Grade": "GRADE STRUCTURE & FLEXIBILITY",
    "leave": "LEAVE POLICY",
    "nda": "Non Compete and Non Disclosure",
    "Office timings": "Office Timing and Attendance Policy",
    "pet": "pet policy",
    "sprint": "Weekly Sprint Policy",
    "work ethics": "WORK PLACE ETHICS"
}

# Standard behavioral questions
QUICK_CHECKS = [
    "Are you willing to relocate if applicable?",
    "What is your notice period?",
    "Can you provide details about your current organization?",
    "Please describe your current role and responsibilities.",
    "What is your current CTC (Cost to Company)?",
    "What is your expected CTC?",
    "What is your educational background?",
    "Can you describe any significant projects you've worked on?",
    "Are there any specific client requirements you want to discuss?",
    "Do you have references from colleagues who might be interested in opportunities with us?"
]
############################################################################################


# Unified prompt template - Single API call for all analyses
unified_evaluation_prompt = """
Act as a senior recruiter with 30+ years of experience.

You MUST evaluate the resume strictly against the job description (JD).
Your reasoning MUST be based ONLY on explicit evidence written in the resume.

IMPORTANT: This is a recruiter-grade evaluation, NOT an ATS keyword scan.

GENERAL PRINCIPLES (APPLY ALWAYS):
- Do NOT infer, assume, or guess skills.
- Do NOT reward buzzwords without demonstrated usage, responsibility, or outcomes.
- Do NOT hallucinate missing skills or capabilities.
- Use recruiter judgment to distinguish between:
  (a) truly missing skills
  (b) explicitly stated equivalent or functionally identical experience.

CRITICAL DISTINCTION (VERY IMPORTANT):
- Explicitly equivalent enterprise evidence COUNTS as MATCHED.
- Adjacent, implied, or loosely related experience does NOT count.

Examples:
- Kafka usage = VALID for MSK (Kafka) unless JD explicitly requires MSK administration.
- MySQL / Oracle / PostgreSQL on AWS = VALID for RDS unless JD explicitly requires RDS operations.
- CI/CD ownership, Docker, cloud deployments = VALID evidence of DevOps collaboration.
- Architecture design, solution reviews, platform ownership = VALID evidence of technical specifications.
- Leadership roles = VALID evidence of mentoring and code review when explicitly stated.

This evaluation MUST work for ANY role (technical or non-technical) and ANY model size (large or small).
Therefore, follow these STRICT rules:

STRICT EVALUATION RULES (APPLY ALL):
1. If the JD lists a MUST-HAVE skill AND there is NO direct or equivalent evidence → treat it as NOT MATCHED.
2. Equivalent enterprise-grade evidence COUNTS if it serves the same functional purpose.
3. Related, adjacent, or implied experience does NOT count.
   (Example: "MQ development" ≠ "MQ administration". "Linux exposure" ≠ "Linux admin".)
4. No assumptions. No optimism. No guessing.
5. If more than 40% of MUST-HAVE skills are NOT MATCHED → final verdict CANNOT be shortlist.
6. For non-technical roles, evaluate ONLY achievements, outcomes, metrics, stakeholder scope, and behavioral competencies.
7. For senior roles, require evidence of leadership, ownership, decision-making, mentoring, and strategic or architectural impact.
8. If the JD is specialized, the resume MUST show direct evidence of that specialization.
9. Apply overqualification analysis ONLY when experience, title, or scope significantly exceeds the JD.
10. Penalize weak resumes aggressively. Do NOT inflate scores for verbosity, repetition, or buzzwords.
11. If evidence is ambiguous → treat as NOT MATCHED.

CERTIFICATIONS:
- If certifications are NOT mentioned in the JD → set "Certification Match" = null AND exclude it from scoring.
- If JD requires certifications:
   • 100 if candidate has ALL required certifications
   • 0 if ANY required certification is missing
- Extra certifications DO NOT increase score beyond 100.

EDUCATION RULE:
- If JD says "Any Postgraduate" → postgraduate is OPTIONAL.
- Penalize education ONLY if the mandatory minimum qualification is missing.

JD MATCH CALCULATION:
- INCLUDE only applicable match factors.
- Weight all applicable factors equally.
- EXCLUDE Certification Match when set to null.
- Overall JD Match must reflect recruiter realism, not keyword completeness.

OUTPUT CONSTRAINTS:
- Your output MUST be strictly valid JSON.
- The structure, keys, and format MUST match exactly.
- Do NOT add, remove, rename, or reorder fields.
- Do NOT include explanations, markdown, or text outside the JSON.

### Output:
Return a valid JSON object ONLY. The JSON object MUST have the following structure:

{{
  "JD Match": "85%",
  "MissingKeywords": [...],
  "Profile Summary": "...",
  "Over/UnderQualification Analysis": "...",
  "Match Factors": {{
    "Skills Match": 0-100,
    "Experience Match": 0-100,
    "Education Match": 0-100,
    "Industry Knowledge": 0-100,
    "Certification Match": number or null
  }},
  "Reasoning": "...",
  "Candidate Fit Analysis": {{
    "Dimension Evaluation": [
      {{
        "Dimension": "...",
        "Evaluation": "✅ Strong / ⚠️ Moderate / ❌ Weak",
        "Recruiter Comments": "..."
      }}
    ],
    "Risk and Gaps": [
      {{
        "Area": "...",
        "Risk": "...",
        "Recruiter Strategy": "..."
      }}
    ] or null,
    "Recommendation": {{
      "Verdict": "❌ Not Recommended / ⚠️ Conditional Shortlist / ✅ Strong Shortlist",
      "Fit Level": "High / Medium / Low",
      "Rationale": "..."
    }},
    "Recruiter Narrative": "..."
  }},
  "Job Stability": {{
    "IsStable": true/false,
    "AverageJobTenure": "...",
    "JobCount": number,
    "StabilityScore": 0-100,
    "ReasoningExplanation": "...",
    "RiskLevel": "Low / Medium / High"
  }},
  "Career Progression": {{
    "progression_score": 0-100,
    "key_observations": [...],
    "career_path": [
      {{
        "title": "...",
        "company": "...",
        "duration": "...",
        "level": "Entry/Mid/Senior/Lead/Manager",
        "progression": "Promotion/Lateral/Step Back"
      }}
    ],
    "red_flags": [...],
    "reasoning": "..."
  }},
  "Interview Questions": {{
    "TechnicalQuestions": [...10 questions...],
    "NonTechnicalQuestions": [...10 questions...]
  }}
}}

Do NOT include any additional text outside the JSON object.

---
Resume:
{resume_text}

JOB DESCRIPTION:
{job_description}

ADDITIONAL CONTEXT (if any):
{additional_context_block}

"""

# Legacy prompt template (kept for backward compatibility if needed)
input_prompt_template = """

Act as a senior recruiter with 30+ years of experience.

You MUST evaluate the resume strictly against the job description (JD).
Your reasoning MUST be based ONLY on explicit evidence written in the resume.

IMPORTANT: This is a recruiter-grade evaluation, NOT an ATS keyword scan.

GENERAL PRINCIPLES (APPLY ALWAYS):
- Do NOT infer, assume, or guess skills.
- Do NOT reward buzzwords without demonstrated usage, responsibility, or outcomes.
- Do NOT hallucinate missing skills or capabilities.
- Use recruiter judgment to distinguish between:
  (a) truly missing skills
  (b) explicitly stated equivalent or functionally identical experience.

CRITICAL DISTINCTION (VERY IMPORTANT):
- Explicitly equivalent enterprise evidence COUNTS as MATCHED.
- Adjacent, implied, or loosely related experience does NOT count.

Examples:
- Kafka usage = VALID for MSK (Kafka) unless JD explicitly requires MSK administration.
- MySQL / Oracle / PostgreSQL on AWS = VALID for RDS unless JD explicitly requires RDS operations.
- CI/CD ownership, Docker, cloud deployments = VALID evidence of DevOps collaboration.
- Architecture design, solution reviews, platform ownership = VALID evidence of technical specifications.
- Leadership roles = VALID evidence of mentoring and code review when explicitly stated.

This evaluation MUST work for ANY role (technical or non-technical) and ANY model size (large or small).
Therefore, follow these STRICT rules:

STRICT EVALUATION RULES (APPLY ALL):
1. If the JD lists a MUST-HAVE skill AND there is NO direct or equivalent evidence → treat it as NOT MATCHED.
2. Equivalent enterprise-grade evidence COUNTS if it serves the same functional purpose.
3. Related, adjacent, or implied experience does NOT count.
   (Example: “MQ development” ≠ “MQ administration”. “Linux exposure” ≠ “Linux admin”.)
4. No assumptions. No optimism. No guessing.
5. If more than 40% of MUST-HAVE skills are NOT MATCHED → final verdict CANNOT be shortlist.
6. For non-technical roles, evaluate ONLY achievements, outcomes, metrics, stakeholder scope, and behavioral competencies.
7. For senior roles, require evidence of leadership, ownership, decision-making, mentoring, and strategic or architectural impact.
8. If the JD is specialized, the resume MUST show direct evidence of that specialization.
9. Apply overqualification analysis ONLY when experience, title, or scope significantly exceeds the JD.
10. Penalize weak resumes aggressively. Do NOT inflate scores for verbosity, repetition, or buzzwords.
11. If evidence is ambiguous → treat as NOT MATCHED.

CERTIFICATIONS:
- If certifications are NOT mentioned in the JD → set "Certification Match" = null AND exclude it from scoring.
- If JD requires certifications:
   • 100 if candidate has ALL required certifications
   • 0 if ANY required certification is missing
- Extra certifications DO NOT increase score beyond 100.

EDUCATION RULE:
- If JD says “Any Postgraduate” → postgraduate is OPTIONAL.
- Penalize education ONLY if the mandatory minimum qualification is missing.

JD MATCH CALCULATION:
- INCLUDE only applicable match factors.
- Weight all applicable factors equally.
- EXCLUDE Certification Match when set to null.
- Overall JD Match must reflect recruiter realism, not keyword completeness.

OUTPUT CONSTRAINTS:
- Your output MUST be strictly valid JSON.
- The structure, keys, and format MUST match exactly.
- Do NOT add, remove, rename, or reorder fields.
- Do NOT include explanations, markdown, or text outside the JSON.

### Output:
Return a valid JSON object ONLY. The JSON object MUST have the following keys:

{{
  "JD Match": "85%",
  "MissingKeywords": [...],
  "Profile Summary": "...",
  "Over/UnderQualification Analysis": "...",
  "Match Factors": {{
    "Skills Match": 0-100,
    "Experience Match": 0-100,
    "Education Match": 0-100,
    "Industry Knowledge": 0-100,
    "Certification Match": number or null
  }},
  "Reasoning": "...",
  "Candidate Fit Analysis": {{
    "Dimension Evaluation": [
      {{
        "Dimension": "...",
        "Evaluation": "✅ Strong / ⚠️ Moderate / ❌ Weak",
        "Recruiter Comments": "..."
      }}
    ],
    "Risk and Gaps": [
      {{
        "Area": "...",
        "Risk": "...",
        "Recruiter Strategy": "..."
      }}
    ] or null,
    "Recommendation": {{
      "Verdict": "❌ Not Recommended / ⚠️ Conditional Shortlist / ✅ Strong Shortlist",
      "Fit Level": "High / Medium / Low",
      "Rationale": "..."
    }},
    "Recruiter Narrative": "..."
  }}
}}

Do NOT include any additional text outside the JSON object.

---
Resume:
{resume_text}

JOB DESCRIPTION:
{job_description}

ADDITIONAL CONTEXT (if any):
{additional_context_block}

"""

interview_questions_prompt = """
You are an experienced recruiter. Generate interview questions based ONLY on information in the resume and job description. 
Do NOT assume skills.
Do NOT ask questions about skills not present in either JD or resume.

**CRITICAL**: Return ONLY valid JSON. No explanations, no reasoning, no additional text. Start with {{ and end with }}

Output ONLY this JSON:

{{
  "TechnicalQuestions": [...10 questions...],
  "NonTechnicalQuestions": [...10 questions...]
}}

Rules:
- Technical questions MUST reference only the technologies, responsibilities, and domains in the resume AND JD.
- Non-technical questions must assess behavior, ownership, teamwork, leadership, or role fit.
- No generic or irrelevant questions.
- No text outside JSON.

RESUME:
{resume_text}

JD:
{job_description}

PROFILE SUMMARY:
{profile_summary}


"""

job_stability_prompt = """
You are an HR analytics expert. Evaluate job stability strictly based on dates and job history written in the resume. 
Do NOT assume missing dates. 
Do NOT infer unstated durations.

Return ONLY this JSON:

{{
  "IsStable": true/false,
  "AverageJobTenure": "...",
  "JobCount": number,
  "StabilityScore": 0-100,
  "ReasoningExplanation": "...",
  "RiskLevel": "Low / Medium / High"
}}

RESUME:
{resume_text}

"""


# Add career progression prompt template after other prompt templates
career_prompt = """
You are an expert HR analyst. Analyze career progression strictly using evidence from the resume. 
Do NOT infer unstated roles or responsibilities.

Return ONLY:

{
  "red_flags": [...],
    "reasoning": [
    "...",
    "...",
    "..."
    ]
}

RESUME:
{resume_text}

"""

#######################################################
# Recruiter Handbook Prompt Template
recruiter_handbook_prompt = """
SYSTEM:
You are an expert technical recruiter and talent evaluator specializing in AI, analytics, product, and consulting roles. 
Your task is to generate a detailed recruiter-style evaluation and fit analysis based on a Job Description (JD) and a Candidate Resume. 
You must think and write like a senior talent partner at a top-tier consulting firm (Fractal, Deloitte, BCG, etc.) — structured, insightful, data-driven, and nuanced. 
Your output will form a Recruiter Handbook that helps internal recruiters, interviewers, and hiring managers make data-driven shortlisting decisions.

---

INSTRUCTIONS:

You will be provided with:
- JOB_DESCRIPTION_TEXT: {job_description}
- CANDIDATE_RESUME_TEXT: {resume_text}

Generate the following sections in professional markdown format:

### 1️⃣ JD Snapshot
Summarize the role in 5–6 lines:
- Role title, level, domain/vertical
- Key skills required
- Nature of role (hands-on, consulting-led, AI-driven, etc.)
- Success indicators

---

### 2️⃣ Candidate Summary
Summarize the candidate's profile in 5–6 lines:
- Education and total experience
- Functional & technical focus areas
- Domain experience (industries)
- Career trajectory highlights
- Distinguishing achievements

---

### 3️⃣ Fit Matrix (Comparison Table)
Create a markdown table with columns:  
**Dimension | JD Expectation | Candidate Evidence | Rating (1–5) | Comment**

Include these dimensions:
- Domain Expertise  
- Technical / AI / Cloud Engineering Depth  
- Consulting Gravitas / CXO Advisory Experience  
- Innovation & IP / Emerging Tech Thought Leadership  
- Hands-on IC Credibility  
- Project & Delivery Leadership  
- Business Acumen & Commercial Awareness  
- Culture / Communication Fit  

Rating legend: 5=Excellent | 3=Average | 1=Weak

---

### 4️⃣ Detailed Fit Commentary
Write 5–8 paragraphs of nuanced recruiter commentary:
- Where the candidate aligns strongly  
- Where they are weak or untested  
- How their consulting/technical balance fits  
- Domain alignment and value they could bring  
- Potential role re-alignment if not an exact fit  
Tone: confident, analytical, and evidence-driven.

---

### 5️⃣ Red Flags & Mitigation
List 3–5 potential red flags (if any) and ways to mitigate or probe during interview.

---

### 6️⃣ Interview Focus Areas
List 8–10 recommended interview questions, grouped under:
- Technical/Engineering depth  
- Consulting & stakeholder management  
- Domain expertise  
- AI/GenAI awareness  
- Leadership & delivery

---

### 7️⃣ Recruiter Summary & Pitch
Write 2 short paragraphs (recruiter voice) that summarize:
- Why this candidate could be compelling to the client
- How to position them internally
Use persuasive, professional tone for internal submission.

---

### 8️⃣ Final Verdict
Provide:
- **Fit Verdict:** Strong Shortlist / Conditional / Reconsider / Reject  
- **Fit Score (0–100%):**  
- **Confidence (High / Medium / Low):**  
- **Best-fit Domain (Retail / FSI / TMT / HLS / Other):**

---

**Output Formatting:**
- Write the recruiter handbook in professional markdown (for rendering in a web UI).  
- Keep tone polished, confident, and analytical — like a top-tier recruiter brief.  
- Be specific, never generic. Use evidence-based phrases like:
  - "Shows strong multi-client consulting maturity."
  - "Demonstrates architectural thought leadership but lacks GenAI delivery exposure."
  - "Consulting gravitas evident from CIO-level advisory roles."
- Do NOT include markdown code block markers in your response.
"""

# Initialize Gemini model
genai.configure(api_key=os.getenv('GOOGLE_API_KEY'))
gemini_model = genai.GenerativeModel(GEMINI_MODEL)

# Initialize OpenAI client
openai_client = None
if MODEL_PROVIDER == "openai" and OPENAI_API_KEY:
    try:
        openai_client = OpenAI(api_key=OPENAI_API_KEY)
    except Exception as e:
        print(f"WARNING: Failed to initialize OpenAI client: {e}")
        print("Falling back to Gemini")
elif MODEL_PROVIDER == "openai" and not OPENAI_API_KEY:
    print("WARNING: MODEL_PROVIDER is set to 'openai' but OPENAI_API_KEY is missing!")
    print("Using Gemini instead. Add OPENAI_API_KEY to your .env file to use OpenAI.")

# Initialize Groq client
groq_client = None
if MODEL_PROVIDER == "groq":
    if GROQ_API_KEY:
        try:
            groq_client = Groq(api_key=GROQ_API_KEY)
            print(f"[INIT] Groq client initialized successfully with model: {GROQ_MODEL}")
        except Exception as e:
            print(f"[ERROR] Failed to initialize Groq client: {e}")
            print("[WARNING] Groq initialization failed. Resume evaluation will fail.")
    else:
        print("[ERROR] MODEL_PROVIDER is set to 'groq' but GROQ_API_KEY is missing!")
        print("[ERROR] Please set GROQ_API_KEY in your .env file.")

# Legacy variable for backward compatibility
model = gemini_model


# ========================================
# UNIFIED MODEL ABSTRACTION LAYER
# ========================================

class UnifiedModelResponse:
    """Wrapper class to standardize responses from different models"""
    def __init__(self, text):
        self.text = text
        self.output_text = text  # For compatibility with your sample code

def generate_content_unified(prompt, stream=False, max_tokens=None):
    """
    Unified function to generate content from either Gemini or OpenAI
    
    Args:
        prompt (str): The prompt to send to the model
        stream (bool): Whether to stream the response (for real-time output)
    
    Returns:
        UnifiedModelResponse or generator: Response object with .text attribute
    """
    try:
        # Default to a high limit if not specified; individual callers (like Info Buddy)
        # can pass smaller values for short answers.
        token_limit = max_tokens or 16384

        if MODEL_PROVIDER == "openai" and openai_client:
            # OpenAI API call
            # Use max_completion_tokens for newer models (gpt-4o, gpt-4o-mini, gpt-5, o1, o3, etc.)
            # Use max_tokens for older models (gpt-3.5-turbo, gpt-4-turbo)
            # Newer models: gpt-4o*, gpt-5*, o1*, o3*
            uses_new_api = any(x in OPENAI_MODEL.lower() for x in ["gpt-4o", "gpt-5", "o1-", "o3-"])
            token_param = "max_completion_tokens" if uses_new_api else "max_tokens"
            
            # Some newer models (gpt-5*, o1*, o3*) don't support custom temperature
            # Only use temperature for models that support it (gpt-4o and older support it)
            supports_temperature = not any(x in OPENAI_MODEL.lower() for x in ["gpt-5", "o1-", "o3-"])
            
            if stream:
                # Streaming response for OpenAI
                params = {
                    "model": OPENAI_MODEL,
                    "messages": [
                        {"role": "system", "content": "You are an expert HR analyst and technical recruiter."},
                        {"role": "user", "content": prompt}
                    ],
                    "stream": True,
                    token_param: token_limit  # Adjustable token limit
                }
                if supports_temperature:
                    params["temperature"] = 0.7
                    
                    
                response = openai_client.chat.completions.create(**params)
                
                def stream_generator():
                    for chunk in response:
                        if chunk.choices[0].delta.content:
                            yield chunk.choices[0].delta.content
                
                return stream_generator()
            else:
                # Non-streaming response for OpenAI
                params = {
                    "model": OPENAI_MODEL,
                    "messages": [
                        {"role": "system", "content": "You are an expert HR analyst and technical recruiter."},
                        {"role": "user", "content": prompt}
                    ],
                    token_param: token_limit  # Adjustable token limit
                }
                if supports_temperature:
                    params["temperature"] = 0.7
                    
                response = openai_client.chat.completions.create(**params)
                logging.info(f"OpenAI response model: {response.model}, choices count: {len(response.choices)}")
                
                # Check response structure
                if not response.choices:
                    logging.error("OpenAI response has no choices!")
                    raise Exception("OpenAI returned empty choices")
                
                message = response.choices[0].message
                logging.info(f"Message role: {message.role}, has content: {message.content is not None}")
                
                # Some models might have refusal or other fields
                if hasattr(message, 'refusal') and message.refusal:
                    logging.error(f"OpenAI refused: {message.refusal}")
                    raise Exception(f"OpenAI refused to respond: {message.refusal}")
                
                content = message.content
                logging.info(f"Content length: {len(content) if content else 0}, type: {type(content)}")
                
                if not content:
                    logging.error(f"OpenAI returned empty content!")
                    logging.error(f"Full message: {message}")
                    logging.error(f"Finish reason: {response.choices[0].finish_reason}")
                    raise Exception(f"OpenAI returned empty content. Finish reason: {response.choices[0].finish_reason}")
                
                return UnifiedModelResponse(content)
        
        elif MODEL_PROVIDER == "groq" and groq_client:
            # Groq API call (same structure as OpenAI, reasoning models supported)
            # Groq uses max_completion_tokens for all models
            token_param = "max_completion_tokens"
            
            # Reasoning models (gpt-oss, o1, o3) don't support custom temperature
            is_reasoning_model = any(x in GROQ_MODEL.lower() for x in ["gpt-oss", "o1-", "o3-"])
            supports_temperature = not is_reasoning_model
            
            if stream:
                # Streaming response for Groq
                params = {
                    "model": GROQ_MODEL,
                    "messages": [
                        {"role": "system", "content": "You are an expert HR analyst and technical recruiter."},
                        {"role": "user", "content": prompt}
                    ],
                    "stream": True,
                    token_param: token_limit  # Adjustable token limit
                }
                
                # Add temperature only for non-reasoning models
                if supports_temperature:
                    params["temperature"] = 0.7
                
                # Add reasoning_effort for reasoning models
                if is_reasoning_model:
                    params["reasoning_effort"] = GROQ_REASONING_EFFORT
                
                response = groq_client.chat.completions.create(**params)
                
                def stream_generator():
                    for chunk in response:
                        if chunk.choices[0].delta.content:
                            yield chunk.choices[0].delta.content
                
                return stream_generator()
            else:
                # Non-streaming response for Groq
                params = {
                    "model": GROQ_MODEL,
                    "messages": [
                        {"role": "system", "content": "You are an expert HR analyst and technical recruiter."},
                        {"role": "user", "content": prompt}
                    ],
                    token_param: token_limit
                }
                
                # Add temperature only for non-reasoning models
                if supports_temperature:
                    params["temperature"] = 0.7
                
                # Add reasoning_effort for reasoning models
                if is_reasoning_model:
                    params["reasoning_effort"] = GROQ_REASONING_EFFORT
                
                response = groq_client.chat.completions.create(**params)
                logging.info(f"Groq response model: {response.model}, choices count: {len(response.choices)}")

                # Check response structure
                if not response.choices:
                    logging.error("Groq response has no choices!")
                    raise Exception("Groq returned empty choices")
                
                message = response.choices[0].message
                logging.info(f"Message role: {message.role}, has content: {message.content is not None}")
                
                content = message.content
                logging.info(f"Content length: {len(content) if content else 0}, type: {type(content)}")
                
                if not content:
                    logging.error(f"Groq returned empty content!")
                    logging.error(f"Full message: {message}")
                    logging.error(f"Finish reason: {response.choices[0].finish_reason}")
                    raise Exception(f"Groq returned empty content. Finish reason: {response.choices[0].finish_reason}")
                
                return UnifiedModelResponse(content)
        
        else:
            # Gemini API call (default)
            if stream:
                # Streaming response for Gemini
                response = gemini_model.generate_content(prompt, stream=True)
                return response  # Gemini already returns a generator
            else:
                # Non-streaming response for Gemini
                response = gemini_model.generate_content(prompt)
                return response  # Gemini response object already has .text
    
    except Exception as e:
        logging.error(f"Error in generate_content_unified: {str(e)}")
        # Fallback to Gemini if OpenAI or Groq fails
        if MODEL_PROVIDER in ["openai", "groq"]:
            logging.warning(f"{MODEL_PROVIDER.upper()} failed, falling back to Gemini")
            if stream:
                response = gemini_model.generate_content(prompt, stream=True)
                return response
            else:
                response = gemini_model.generate_content(prompt)
                return response
        raise

# Log which model is being used
print("=" * 60)
print(f"Model Provider Configuration: {MODEL_PROVIDER.upper()}")
if MODEL_PROVIDER == "openai" and openai_client:
    print(f"ACTUALLY USING: OpenAI Model: {OPENAI_MODEL}")
elif MODEL_PROVIDER == "openai" and not openai_client:
    print(f"ACTUALLY USING: Gemini Model: {GEMINI_MODEL} (OpenAI not available)")
elif MODEL_PROVIDER == "groq" and groq_client:
    print(f"ACTUALLY USING: Groq Model: {GROQ_MODEL}")
    if "gpt-oss" in GROQ_MODEL.lower() or "o1-" in GROQ_MODEL.lower() or "o3-" in GROQ_MODEL.lower():
        print(f"   Reasoning Effort: {GROQ_REASONING_EFFORT}")
elif MODEL_PROVIDER == "groq" and not groq_client:
    print(f"ACTUALLY USING: Gemini Model: {GEMINI_MODEL} (Groq not available)")
else:
    print(f"ACTUALLY USING: Gemini Model: {GEMINI_MODEL}")
print("=" * 60)

logging.info(f"Model Provider: {MODEL_PROVIDER.upper()}")
if MODEL_PROVIDER == "openai":
    logging.info(f"Using OpenAI Model: {OPENAI_MODEL}")
elif MODEL_PROVIDER == "groq":
    logging.info(f"Using Groq Model: {GROQ_MODEL}")
else:
    logging.info(f"Using Gemini Model: {GEMINI_MODEL}")

# ========================================


# Initialize Flask app
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['SECRET_KEY'] = SECRET_KEY
asgi_app = WsgiToAsgi(app)

# Initialize Google OAuth
oauth = OAuth(app)
google = oauth.register(
    name='google',
    client_id=GOOGLE_CLIENT_ID,
    client_secret=GOOGLE_CLIENT_SECRET,
    server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
    client_kwargs={
        'scope': 'openid email profile'
    },
    # Explicitly set the callback route to match Google Cloud Console configuration
    authorize_callback='authorize'
)

# Create uploads directory if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Initialize NLTK
nltk.download("punkt")


# Initialize Groq LLM
llm = ChatGroq(
    groq_api_key=GROQ_API_KEY,
    # model_name="mixtral-8x7b-32768",  # This generates long text,  max_tokens=4096
    # model_name=   "llama-3.1-8b-instant",#"qwen-2.5-32b", #"deepseek-r1-distill-qwen-32b",
    model_name = "qwen/qwen3-32b",
    temperature=0.377,
    max_tokens=2048,   #4096
    top_p=0.95,
    presence_penalty=0.1,
    frequency_penalty=0.1
)

# Initialize Pinecone
pc = Pinecone(api_key=PINECONE_API_KEY)
index_name = PINECONE_INDEX_NAME
if index_name not in pc.list_indexes().names():
    pc.create_index(
        name=index_name,
        dimension=384,
        metric="cosine",
        spec=ServerlessSpec(cloud="aws", region="us-east-1")
    )
index = pc.Index(index_name)

# Initialize embeddings
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Initialize vector store
vectorstore = None
try:
    from langchain_pinecone import PineconeVectorStore as NewPineconeVectorStore
    vectorstore = NewPineconeVectorStore(
        index=index,
        embedding=embeddings,
        text_key="text"
    )
    logging.info("✅ Using new langchain-pinecone vectorstore")
except ImportError:
    # Fallback to old import if new package not available
    try:
        from langchain_community.vectorstores import Pinecone as PineconeVectorStore
        vectorstore = PineconeVectorStore(
            index=index,
            embedding=embeddings,
            text_key="text"
        )
        logging.info("✅ Using old langchain-community vectorstore")
    except Exception as e:
        logging.error(f"❌ Error initializing vectorstore: {e}")
        vectorstore = None
except Exception as e:
    logging.error(f"❌ Error initializing new vectorstore: {e}")
    vectorstore = None

# Uploads & database configuration
DATABASE_NAME = 'combined_db.db'
# Ensure uploads folder is always rooted at the app directory so moves/WD changes don't break paths
app.config['UPLOAD_FOLDER'] = os.path.join(app.root_path, "uploads")

# Dedicated error log file for backend debugging (not exposed to frontend)
ERROR_LOG_DIR = os.path.join(app.root_path, "error_logs")
os.makedirs(ERROR_LOG_DIR, exist_ok=True)
error_log_path = os.path.join(ERROR_LOG_DIR, "resume_download_errors.txt")

error_file_handler = logging.FileHandler(error_log_path, encoding="utf-8")
error_file_handler.setLevel(logging.ERROR)
error_formatter = logging.Formatter(
    "%(asctime)s - %(levelname)s - %(name)s - %(message)s"
)
error_file_handler.setFormatter(error_formatter)
logging.getLogger().addHandler(error_file_handler)

def init_db():
    """Initialize database with all required tables"""
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()
    
    # Create evaluations table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS evaluations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            resume_path TEXT,
            filename TEXT,
            job_title TEXT,
            job_description TEXT,
            match_percentage REAL,
            match_factors TEXT,
            profile_summary TEXT,
            missing_keywords TEXT,
            job_stability TEXT,
            career_progression TEXT,
            technical_questions TEXT,
            nontechnical_questions TEXT,
            behavioral_questions TEXT,
            oorwin_job_id TEXT,
            candidate_fit_analysis TEXT,
            over_under_qualification TEXT,
            time_taken REAL,
            token_usage INTEGER,
            user_email TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Create qa_history table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS qa_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            question TEXT NOT NULL,
            retrieved_docs TEXT,
            final_answer TEXT,
            feedback TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Create qa_feedback table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS qa_feedback (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            question_id INTEGER,
            rating INTEGER,
            feedback TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(question_id),
            FOREIGN KEY (question_id) REFERENCES qa_history (id)
        )
    ''')
    
    # Create feedback table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS feedback (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            evaluation_id INTEGER,
            rating INTEGER,
            comments TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(evaluation_id),
            FOREIGN KEY (evaluation_id) REFERENCES evaluations (id)
        )
    ''')
    
    # Create handbook_feedback table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS handbook_feedback (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            handbook_id INTEGER NOT NULL,
            rating INTEGER NOT NULL CHECK(rating >= 1 AND rating <= 5),
            comments TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(handbook_id),
            FOREIGN KEY (handbook_id) REFERENCES recruiter_handbooks (id)
        )
    ''')
    
    # Create interview_questions table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS interview_questions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            evaluation_id INTEGER,
            technical_questions TEXT,
            nontechnical_questions TEXT,
            behavioral_questions TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (evaluation_id) REFERENCES evaluations (id)
        )
    ''')
    
    # Create recruiter_handbooks table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS recruiter_handbooks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            oorwin_job_id TEXT,
            job_title TEXT,
            job_description TEXT,
            additional_context TEXT,
            markdown_content TEXT,
            time_taken REAL,
            user_email TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Create users table for authentication and role management
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT UNIQUE NOT NULL,
            name TEXT,
            role TEXT NOT NULL DEFAULT 'Recruiter',
            team TEXT,
            manager_email TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            updated_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (manager_email) REFERENCES users(email)
        )
    ''')
    
    # Handle schema updates for existing tables
    try:
        # Check if evaluations table has new columns
        cursor.execute("PRAGMA table_info(evaluations)")
        columns = [col[1] for col in cursor.fetchall()]
        
        if 'candidate_fit_analysis' not in columns:
            cursor.execute('ALTER TABLE evaluations ADD COLUMN candidate_fit_analysis TEXT')
            print("Added candidate_fit_analysis column to evaluations")
        
        if 'over_under_qualification' not in columns:
            cursor.execute('ALTER TABLE evaluations ADD COLUMN over_under_qualification TEXT')
            print("Added over_under_qualification column to evaluations")
        
        if 'user_email' not in columns:
            cursor.execute('ALTER TABLE evaluations ADD COLUMN user_email TEXT')
            print("Added user_email column to evaluations")
    except Exception as e:
        print(f"Note: Schema update check: {e}")
    
    # Add user_email to recruiter_handbooks if not exists
    try:
        cursor.execute("PRAGMA table_info(recruiter_handbooks)")
        handbook_columns = [col[1] for col in cursor.fetchall()]
        if 'user_email' not in handbook_columns:
            cursor.execute('ALTER TABLE recruiter_handbooks ADD COLUMN user_email TEXT')
            print("Added user_email column to recruiter_handbooks")
    except Exception as e:
        print(f"Note: Schema update check for handbooks: {e}")
    
    # Add time_taken column to evaluations table
    try:
        cursor.execute("PRAGMA table_info(evaluations)")
        eval_columns = [col[1] for col in cursor.fetchall()]
        if 'time_taken' not in eval_columns:
            cursor.execute('ALTER TABLE evaluations ADD COLUMN time_taken REAL')
            print("Added time_taken column to evaluations")
        if 'token_usage' not in eval_columns:
            cursor.execute('ALTER TABLE evaluations ADD COLUMN token_usage REAL')
            print("Added token_usage column to evaluations")
        # Ensure token_usage column type is REAL (not INTEGER) for compatibility
        try:
            cursor.execute("PRAGMA table_info(evaluations)")
            eval_columns_info = cursor.fetchall()
            token_usage_col = next((col for col in eval_columns_info if col[1] == 'token_usage'), None)
            if token_usage_col and token_usage_col[2].upper() == 'INTEGER':
                # SQLite doesn't support changing column types directly, but REAL can store integers
                # So we'll just note it and continue
                pass
        except Exception as e:
            print(f"Note: Could not check token_usage column type: {e}")
        conn.commit()
    except Exception as e:
        print(f"Note: Schema update check for evaluations: {e}")
    
    # Add time_taken column to recruiter_handbooks table
    try:
        cursor.execute("PRAGMA table_info(recruiter_handbooks)")
        handbook_columns = [col[1] for col in cursor.fetchall()]
        if 'time_taken' not in handbook_columns:
            cursor.execute('ALTER TABLE recruiter_handbooks ADD COLUMN time_taken REAL')
            print("Added time_taken column to recruiter_handbooks")
    except Exception as e:
        print(f"Note: Schema update check for handbooks time_taken: {e}")
    
    # Initialize default admin user and teams
    try:
        # Insert default admin user
        cursor.execute('''
            INSERT OR IGNORE INTO users (email, name, role, team)
            VALUES (?, ?, ?, ?)
        ''', ('ritesh.m@peoplelogic.in', 'Ritesh M', 'Admin', 'Core'))
        
        # Initialize default teams (if needed, we can add team management later)
        # For now, teams are just strings: ITS, OSS, PCS, ISV, Core
        
        conn.commit()
        print("Initialized default admin user")
    except Exception as e:
        print(f"Note: User initialization: {e}")
    
    conn.commit()
    conn.close()

# Initialize database at startup
init_db()

# Authentication Helper Functions
def get_user_info(email):
    """Get user information from database"""
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()
    cursor.execute('SELECT email, name, role, team, manager_email FROM users WHERE email = ?', (email,))
    user = cursor.fetchone()
    conn.close()
    
    if user:
        return {
            'email': user[0],
            'name': user[1],
            'role': user[2],
            'team': user[3],
            'manager_email': user[4]
        }
    return None

def create_or_update_user(email, name):
    """Create or update user in database (default role: Recruiter)
    Preserves existing team and manager_email when updating to prevent data loss on login"""
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()
    
    # Check if user already exists
    cursor.execute('SELECT role, team, manager_email FROM users WHERE email = ?', (email,))
    existing_user = cursor.fetchone()
    
    if existing_user:
        # User exists - ONLY update name and updated_at, preserve everything else (role, team, manager_email)
        # This prevents team and manager_email from being cleared on login
        cursor.execute('''
            UPDATE users 
            SET name = ?, 
                updated_at = CURRENT_TIMESTAMP
            WHERE email = ?
        ''', (name, email))
        logging.info(f"Updated user {email}: preserved role, team, and manager_email")
    else:
        # New user - create with default role
        cursor.execute('''
            INSERT INTO users (email, name, role, team, manager_email, created_at, updated_at)
            VALUES (?, ?, 'Recruiter', NULL, NULL, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
        ''', (email, name))
        logging.info(f"Created new user {email} with default role 'Recruiter'")
    
    conn.commit()
    conn.close()

def get_accessible_users(current_user_email):
    """Get list of users that current user can access based on team membership only"""
    user_info = get_user_info(current_user_email)
    if not user_info:
        return []
    
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()
    
    team = user_info['team']
    
    if not team:
        # If user has no team assigned, show all users (for backward compatibility and Admin users)
        # This allows users to see all users in the dropdown for filtering
        # But data filtering will still be team-based when teams are assigned
        cursor.execute('SELECT email, name, role, team FROM users ORDER BY name')
    else:
        # Users can only see other users from the same team (team-based filtering, roles don't matter)
        cursor.execute('''
            SELECT email, name, role, team FROM users 
            WHERE team = ? 
            ORDER BY name
        ''', (team,))
    
    users = cursor.fetchall()
    conn.close()
    
    return [{'email': u[0], 'name': u[1], 'role': u[2], 'team': u[3]} for u in users]

def get_accessible_user_emails(current_user_email):
    """Get list of email addresses that current user can access"""
    users = get_accessible_users(current_user_email)
    return [u['email'] for u in users]

def filter_data_by_role(query, table_name, user_email_column, current_user_email):
    """Add WHERE clause to filter data based on user's team membership (team-based filtering)"""
    user_info = get_user_info(current_user_email)
    if not user_info:
        return query + " WHERE 1=0"  # Return no results if user not found
    
    accessible_emails = get_accessible_user_emails(current_user_email)
    
    if not accessible_emails:
        # If no accessible emails (e.g., user has no team), return no results
        return query + " WHERE 1=0"
    
    # Filter by accessible user emails (team members only)
    placeholders = ','.join(['?'] * len(accessible_emails))
    if 'WHERE' in query.upper():
        return f"{query} AND {user_email_column} IN ({placeholders})"
    else:
        return f"{query} WHERE {user_email_column} IN ({placeholders})"

# Authentication Decorators
def login_required(f):
    """Decorator to require login"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

def role_required(*allowed_roles):
    """Decorator to require specific role(s)"""
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if 'user' not in session:
                return redirect(url_for('login'))
            user_email = session['user'].get('email')
            user_info = get_user_info(user_email)
            if not user_info or user_info['role'] not in allowed_roles:
                return jsonify({'error': 'Access denied. Insufficient permissions.'}), 403
            return f(*args, **kwargs)
        return decorated_function
    return decorator

# Helper functions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_gemini_response(input_prompt):
    """Get response from selected model (Gemini or OpenAI) and clean it up."""
    try:
        response = generate_content_unified(input_prompt)
        response_text = response.text.strip()
        
        # Remove markdown code block markers if present
        if response_text.startswith('```json'):
            response_text = response_text[7:]
        if response_text.endswith('```'):
            response_text = response_text[:-3]
            
        # Clean up any extra whitespace and newlines
        response_text = response_text.strip()
        
        # Try to parse as JSON to validate
        try:
            parsed_json = json.loads(response_text)
            return json.dumps(parsed_json)  # Return properly formatted JSON string
        except json.JSONDecodeError:
            # If not valid JSON, try to extract JSON using regex
            match = re.search(r"\{.*\}", response_text, re.DOTALL)
            if match:
                try:
                    parsed_json = json.loads(match.group(0))
                    return json.dumps(parsed_json)  # Return properly formatted JSON string
                except json.JSONDecodeError:
                    raise ValueError("Invalid JSON structure in response")
            else:
                raise ValueError("No valid JSON found in response")
                
    except Exception as e:
        logging.error(f"Error in get_gemini_response: {str(e)}")
        return json.dumps({})  # Return valid empty JSON object as fallback

def sanitize_resume_text(text):
    """
    Remove email addresses and phone numbers from resume text for privacy protection.
    
    Args:
        text: Raw resume text
        
    Returns:
        Sanitized text with email and phone numbers removed
    """
    if not text:
        return text
    
    # Remove email addresses (various formats)
    # Pattern matches: user@domain.com, user.name@domain.co.uk, user+tag@domain.com, etc.
    # Note: Case-insensitive matching for better coverage
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'
    text = re.sub(email_pattern, '[EMAIL_REMOVED]', text, flags=re.IGNORECASE)
    
    # Remove phone numbers (various formats)
    # IMPORTANT: Patterns are designed to NOT match years (1900-2099) or date ranges
    # Strategy: Use negative lookahead/lookbehind to exclude common date contexts
    
    # US/Standard format: (123) 456-7890 or 123-456-7890 or 123.456.7890 or 123 456 7890
    # Exclude if it looks like a year (starts with 19xx or 20xx)
    phone_pattern2 = r'(?<!\d)(?!(?:19|20)\d{2})\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}(?!\d)'
    
    # Indian mobile format: +91-98765-43210 or 91-98765-43210 or 98765-43210 or 9876543210
    # Indian mobile numbers start with 6, 7, 8, or 9 (safe, won't match years)
    phone_pattern3 = r'(\+?91[-.\s]?)?[6-9]\d{4}[-.\s]?\d{5}'
    
    # Standalone 10-digit numbers starting with 3-9 (excluding 1xxx and 20xx which could be years)
    # Indian mobile: 6-9, US: 3-5,7-9 (excluding 1xxx and 20xx to avoid years)
    phone_pattern4 = r'\b([3-9]\d{9})\b(?!\s*(?:to|–|-)\s*(?:19|20)\d{2})'  # Exclude if followed by date range
    
    # International format with country code: +[1-9] followed by 7+ digits
    # More restrictive: require country code and sufficient digits to avoid matching dates
    phone_pattern1 = r'\+[1-9]\d{0,2}[-.\s]?\d{4,}[-.\s]?\d{3,}(?!\d)'
    
    # Apply all patterns in order
    text = re.sub(phone_pattern1, '[PHONE_REMOVED]', text)
    text = re.sub(phone_pattern2, '[PHONE_REMOVED]', text)
    text = re.sub(phone_pattern3, '[PHONE_REMOVED]', text)
    text = re.sub(phone_pattern4, '[PHONE_REMOVED]', text)
    
    # Clean up multiple consecutive replacements
    text = re.sub(r'\[EMAIL_REMOVED\](?:\s*\[EMAIL_REMOVED\])+', '[EMAIL_REMOVED]', text)
    text = re.sub(r'\[PHONE_REMOVED\](?:\s*\[PHONE_REMOVED\])+', '[PHONE_REMOVED]', text)
    
    return text

def convert_docx_to_pdf(docx_path, pdf_path=None):
    """Convert DOCX file to PDF using reportlab"""
    try:
        if not os.path.exists(docx_path):
            logging.error(f"DOCX file not found: {docx_path}")
            return None
        
        # Generate PDF path if not provided
        if not pdf_path:
            pdf_path = os.path.splitext(docx_path)[0] + '.pdf'
        
        # Read DOCX content
        doc = Document(docx_path)
        
        # Create PDF
        pdf_buffer = BytesIO()
        pdf_doc = SimpleDocTemplate(
            pdf_buffer,
            pagesize=A4,
            rightMargin=50,
            leftMargin=50,
            topMargin=50,
            bottomMargin=50
        )
        
        # Define styles
        styles = getSampleStyleSheet()
        normal_style = ParagraphStyle(
            'ResumeStyle',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
            spaceAfter=6,
            alignment=TA_LEFT
        )
        
        heading_style = ParagraphStyle(
            'ResumeHeading',
            parent=styles['Heading1'],
            fontSize=14,
            leading=16,
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        elements = []
        
        # Process paragraphs from DOCX
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                elements.append(Spacer(1, 6))
                continue
            
            # Check if it's a heading (simple heuristic: if it's short and bold-like)
            if len(text) < 100 and para.style.name.startswith('Heading'):
                elements.append(Paragraph(text, heading_style))
            else:
                # Escape HTML special characters for reportlab
                text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                elements.append(Paragraph(text, normal_style))
            elements.append(Spacer(1, 6))
        
        # Process tables if any
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    row_data.append(Paragraph(cell_text, normal_style))
                table_data.append(row_data)
            
            if table_data:
                pdf_table = Table(table_data)
                pdf_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                elements.append(Spacer(1, 12))
                elements.append(pdf_table)
                elements.append(Spacer(1, 12))
        
        # Build PDF
        pdf_doc.build(elements)
        
        # Save PDF to file
        pdf_data = pdf_buffer.getvalue()
        pdf_buffer.close()
        
        with open(pdf_path, 'wb') as f:
            f.write(pdf_data)
        
        logging.info(f"Successfully converted DOCX to PDF: {docx_path} -> {pdf_path}")
        return pdf_path
        
    except Exception as e:
        logging.error(f"Error converting DOCX to PDF: {str(e)}", exc_info=True)
        return None

def extract_text_from_file(file_path):
    try:
        # Check if file_path is None or the string "NULL"
        if not file_path or file_path == 'NULL' or file_path == 'None' or str(file_path).strip() == '':
            logging.error(f"Invalid file path provided to extract_text_from_file: {file_path}")
            return None, "Invalid file path"
        
        # Check if file exists
        if not os.path.exists(file_path):
            logging.error(f"File does not exist: {file_path}")
            return None, f"File not found: {file_path}"
        
        ext = file_path.rsplit('.', 1)[1].lower()
        if ext == 'pdf':
            try:
                with pdfplumber.open(file_path) as pdf:
                    text = ""
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                # Sanitize text: remove email and phone numbers
                text = sanitize_resume_text(text)
                return text
            except ModuleNotFoundError as e:
                if "PyCryptodome" in str(e) or "Crypto" in str(e):
                    return None, "PyCryptodome is required for encrypted PDFs. Please install it with 'pip install pycryptodome'."
                raise
        elif ext == 'docx':
            doc = Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            # Sanitize text: remove email and phone numbers
            text = sanitize_resume_text(text)
            return text
        elif ext == 'doc':
            return None, "Support for .doc files is limited. Please convert to .docx or PDF."
        else:
            return None, "Unsupported file format."
    except Exception as e:
        logging.error(f"File extraction error: {str(e)}")
        return None, str(e)

def hybrid_search(query, k=5):
    """Perform hybrid search using BM25 and vector similarity."""
    try:
        # Get vector search results
        vector_results = vectorstore.similarity_search(query, k=k)
        
        # Get BM25 results
        bm25_results = []
        if os.path.exists(POLICIES_FOLDER):
            for filename in os.listdir(POLICIES_FOLDER):
                if filename.endswith(('.txt', '.md')):
                    with open(os.path.join(POLICIES_FOLDER, filename), 'r', encoding='utf-8') as f:
                        text = f.read()
                        sentences = sent_tokenize(text)  # Use NLTK sentence tokenizer
                        bm25_results.extend(sentences)
        
        # Combine and deduplicate results
        combined_results = []
        seen_texts = set()
        
        # Add vector search results
        for doc in vector_results:
            if doc.page_content not in seen_texts:
                combined_results.append(doc.page_content)
                seen_texts.add(doc.page_content)
        
        # Add BM25 results
        for sentence in bm25_results:
            if sentence not in seen_texts:
                combined_results.append(sentence)
                seen_texts.add(sentence)
        
        # Join results with newlines
        return "\n".join(combined_results)
    
    except Exception as e:
        logging.error(f"Error in hybrid_search: {e}")
        return ""

def save_evaluation(eval_id, filename, job_title, rank_score, missing_keywords, profile_summary, match_factors, job_stability, additional_info=None, oorwin_job_id=None, candidate_fit_analysis=None, over_under_qualification=None, user_email=None, time_taken=None, token_usage=None, resume_path=None):
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # Convert data to JSON strings if they're not already
        try:
            missing_keywords_json = json.dumps(missing_keywords) if not isinstance(missing_keywords, str) else missing_keywords
        except Exception as e:
            logging.error(f"Error converting missing_keywords to JSON: {e}")
            missing_keywords_json = '[]'
            
        try:
            match_factors_json = json.dumps(match_factors) if not isinstance(match_factors, str) else match_factors
        except Exception as e:
            logging.error(f"Error converting match_factors to JSON: {e}")
            match_factors_json = '{}'
            
        try:
            job_stability_json = json.dumps(job_stability) if not isinstance(job_stability, str) else job_stability
        except Exception as e:
            logging.error(f"Error converting job_stability to JSON: {e}")
            job_stability_json = '{}'
        
        # Ensure all JSON strings are valid
        if not missing_keywords_json or missing_keywords_json == 'null':
            missing_keywords_json = '[]'
        if not match_factors_json or match_factors_json == 'null':
            match_factors_json = '{}'
        if not job_stability_json or job_stability_json == 'null':
            job_stability_json = '{}'
        
        # Convert rank_score to integer if it's a string
        rank_score_int = int(rank_score) if isinstance(rank_score, str) else rank_score
        
        # Ensure all values are strings except rank_score_int and eval_id
        filename_str = str(filename)
        job_title_str = str(job_title)
        profile_summary_str = str(profile_summary)
        
        # Handle resume_path - store ONLY filename in DB for portability
        # If resume_path is a full/relative path, strip down to basename.
        if resume_path:
            resume_path_str = os.path.basename(str(resume_path))
        else:
            # Fallback: use filename as the stored resume identifier
            resume_path_str = os.path.basename(filename_str) if filename_str else None
        
        # Handle oorwin_job_id (can be None or empty string)
        oorwin_job_id_str = str(oorwin_job_id).strip() if oorwin_job_id else None
        if oorwin_job_id_str == '' or oorwin_job_id_str == 'None':
            oorwin_job_id_str = None
        
        # Convert additional_info to JSON string if it's a dict or list
        if isinstance(additional_info, (dict, list)):
            additional_info_str = json.dumps(additional_info)
        else:
            additional_info_str = str(additional_info) if additional_info is not None else ""
        
        # Extract career progression from additional_info
        career_progression = additional_info.get('career_progression', {}) if isinstance(additional_info, dict) else {}
        career_progression_json = json.dumps(career_progression)
        
        # Convert new fields to JSON
        candidate_fit_analysis_json = json.dumps(candidate_fit_analysis) if candidate_fit_analysis else '{}'
        over_under_qualification_str = str(over_under_qualification) if over_under_qualification else ''
        
        # Convert token_usage to integer if it's a float or string
        if token_usage is not None:
            try:
                token_usage_int = int(float(token_usage)) if token_usage else None
            except (ValueError, TypeError):
                token_usage_int = None
        else:
            token_usage_int = None
        
        # Debug: Log the actual values being inserted
        logging.info(f"Values to insert - eval_id: {eval_id}, filename: {filename_str}, job_title: {job_title_str}")
        logging.info(f"JSON values - missing_keywords_json type: {type(missing_keywords_json)}, value: {missing_keywords_json[:100] if len(missing_keywords_json) > 100 else missing_keywords_json}")
        logging.info(f"JSON values - match_factors_json type: {type(match_factors_json)}")
        logging.info(f"All param types: rank_score_int={type(rank_score_int)}, oorwin_job_id_str={type(oorwin_job_id_str)}, datetime={type(datetime.now())}")
        
        # Convert datetime to string
        timestamp_str = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        # Prepare user_email
        user_email_str = user_email if user_email else None
        
        # Prepare all parameters as a tuple (excluding eval_id which is not in the INSERT)
        # Order must match the INSERT statement columns exactly:
        # resume_path, filename, job_title, job_description, match_percentage, 
        # match_factors, profile_summary, missing_keywords, 
        # job_stability, career_progression, technical_questions,
        # nontechnical_questions, behavioral_questions, oorwin_job_id, 
        # candidate_fit_analysis, over_under_qualification, time_taken, token_usage, user_email, timestamp
        params = (
            resume_path_str,                # resume_path (full path to file)
            filename_str,                    # filename
            job_title_str,                   # job_title
            "",                             # job_description
            rank_score_int,                 # match_percentage
            match_factors_json,             # match_factors
            profile_summary_str,             # profile_summary
            missing_keywords_json,           # missing_keywords
            job_stability_json,              # job_stability
            career_progression_json,         # career_progression
            None,                            # technical_questions
            None,                            # nontechnical_questions
            None,                            # behavioral_questions
            oorwin_job_id_str,               # oorwin_job_id
            candidate_fit_analysis_json,     # candidate_fit_analysis
            over_under_qualification_str,    # over_under_qualification
            time_taken,                      # time_taken (in seconds, can be None)
            token_usage_int,                 # token_usage (converted to int, can be None)
            user_email_str,                  # user_email
            timestamp_str                    # timestamp
        )
        
        # Log all parameter types
        logging.info(f"Parameter types: {[type(p).__name__ for p in params]}")
        logging.info(f"Inserting evaluation - user_email: {user_email_str}, timestamp: {timestamp_str}")
        
        cursor.execute(
            """
            INSERT INTO evaluations (
                resume_path, filename, job_title, job_description, match_percentage, 
                match_factors, profile_summary, missing_keywords, 
                job_stability, career_progression, technical_questions,
                nontechnical_questions, behavioral_questions, oorwin_job_id, 
                candidate_fit_analysis, over_under_qualification, time_taken, token_usage, user_email, timestamp
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            params
            )
        
        # Get the auto-generated ID
        db_id = cursor.lastrowid
        
        conn.commit()
        
        # Log successful save
        logging.info(f"✅ EVALUATION SAVED TO DATABASE!")
        logging.info(f"   database_id: {db_id}")
        logging.info(f"   filename: {filename_str}")
        logging.info(f"   job_title: {job_title_str}")
        logging.info(f"   match_percentage: {rank_score_int}")
        logging.info(f"   oorwin_job_id: {oorwin_job_id_str}")
        
        conn.close()
        return db_id  # Return the database ID instead of True
    except sqlite3.OperationalError as e:
        error_msg = str(e)
        logging.error(f"Database operational error in save_evaluation: {error_msg}", exc_info=True)
        # Check if it's a missing column error
        if "no such column" in error_msg.lower():
            logging.error("Database schema mismatch detected. Please run init_db() to update schema.")
            # Try to update schema automatically
            try:
                update_db_schema()
                logging.info("Attempted to update database schema")
            except Exception as schema_error:
                logging.error(f"Failed to update schema: {schema_error}")
        return False
    except Exception as e:
        logging.error(f"Database error in save_evaluation: {str(e)}", exc_info=True)
        logging.error(f"Data being saved - eval_id: {eval_id}, filename: {filename}, job_title: {job_title}")
        logging.error(f"Data types - rank_score: {type(rank_score)}, missing_keywords: {type(missing_keywords)}")
        logging.error(f"Token usage value: {token_usage}, type: {type(token_usage)}")
        return False

def save_feedback(evaluation_id, rating, comments):
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO feedback (evaluation_id, rating, comments, timestamp) VALUES (?, ?, ?, ?)",
            (evaluation_id, rating, comments, datetime.now())
        )
        logging.debug(f"Feedback inserted: evaluation_id={evaluation_id}, rating={rating}, comments={comments}")
        conn.commit()
        conn.close()
        return True
    except sqlite3.Error as e:
        logging.error(f"Database error in save_feedback: {str(e)}")
        return False
    except Exception as e:
        logging.error(f"Unexpected error in save_feedback: {str(e)}")
        return False

def save_interview_questions(evaluation_id, technical_questions, nontechnical_questions, behavioral_questions):
    """Save interview questions to database"""
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        technical_json = json.dumps(technical_questions) if not isinstance(technical_questions, str) else technical_questions
        nontechnical_json = json.dumps(nontechnical_questions) if not isinstance(nontechnical_questions, str) else nontechnical_questions
        behavioral_json = json.dumps(behavioral_questions) if not isinstance(behavioral_questions, str) else behavioral_questions
        
        cursor.execute(
            "INSERT INTO interview_questions (evaluation_id, technical_questions, nontechnical_questions, behavioral_questions, timestamp) VALUES (?, ?, ?, ?, ?)",
            (evaluation_id, technical_json, nontechnical_json, behavioral_json, datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        )
        conn.commit()
        conn.close()
        logging.debug(f"Interview questions saved successfully for: {evaluation_id}")
        return True
    except Exception as e:
        logging.error(f"Database error in save_interview_questions: {str(e)}")
        return False

def save_recruiter_handbook(evaluation_id, markdown_content, json_summary):
    """Save recruiter handbook to database"""
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # Ensure json_summary is a string
        json_summary_str = json_summary if isinstance(json_summary, str) else json.dumps(json_summary)
        
        cursor.execute(
            "INSERT INTO recruiter_handbooks (evaluation_id, markdown_content, json_summary, timestamp) VALUES (?, ?, ?, ?)",
            (evaluation_id, markdown_content, json_summary_str, datetime.now())
        )
        conn.commit()
        conn.close()
        logging.debug(f"Recruiter handbook saved successfully for: {evaluation_id}")
        return True
    except Exception as e:
        logging.error(f"Database error in save_recruiter_handbook: {str(e)}")
        return False

# Add these constants near the top with other constants
BOT_INFO = {
    "name": "PeopleBot",
    "creator": "PeopleLogic",
    "responsibility": "Help recruiters in HR policies, benefits & with any other questions!",
    "capabilities": "Help recruiters in HR policies, benefits & with any other questions"
}

GREETINGS = ["hello", "hi", "hey", "greetings", "good morning", "good afternoon", "good evening"]
IDENTITY_QUESTIONS = [
    "who are you",
    "who are u",
    "who r yu",
    "who r u",
    "who are you",
    "what is your name",
    "what are you",
    "who built you",
    "what are u",
    "who built u",
    "who created you",
    "what can you do",
    "what do you do",
    "what is your name",
    "tell me about yourself",
    "tell me about you",
     "Who Are You", "WHO ARE YOU", "Who r U", "wHo bUiLt YoU",  
    "WHAT IS YOUR NAME", "wHaT dO yOu dO",  "whoz u", "what u do", "wots ur name", "whut do u do",  
    "wats ur function", "who dun made u",  "what are you capable of", "what skills do you have",  
    "describe yourself", "what is your role", "explain yourself",  
    "what services do you provide", "how do you work",  
    "can you tell me what you do", "can you tell me your name", "can you introduce yourself",  
    "what should I call you", "how do I address you",  
    "could you tell me what you are", "what do people call you",  
    "what is your full name", "how were you created",  
    "give me some details about you", "tell me your background",  "who u", "who dis", "whats ur deal", "whats ur function", "who tf r u",  
    "who's this", "u bot?", "who you be", "ur name?", "who made u",  
    "who's your maker", "who made this bot", "whu r u", "whi r u", "who r yu", "wht is ur name", "whts ur name",  
    "whats ur name", "wat is your name", "wat r u", "whu built u",  
    "who made you", "who designed you", "who programmed you", "what are you capable of", "what skills do you have",  
    "describe yourself", "what is your role", "explain yourself",  
    "what services do you provide", "how do you work",  
    "can you tell me what you do",  

    # Mixed uppercase/lowercase variations  
    "Who Are You", "WHO ARE YOU", "Who r U", "wHo bUiLt YoU",  
    "WHAT IS YOUR NAME", "wHaT dO yOu dO",  

    # Slang & misspellin
    "whoz u", "what u do", "wots ur name", "whut do u do",  
    "wats ur function", "who dun made u",  

    # Phonetic spellings & accents  
    "hoo r u", "wat iz ur name", "hoo maid u", "wat cn u do",  
    "whut ur name", "whachu do",  

    # Extra punctuation variations  
    "who are you?", "who are you!", "who are you??",  
    "what is your name?", "who built you??", "who made you?!",
]

def handle_special_queries(question):
    """Handle greetings and identity-related questions."""
    question_lower = question.lower().strip("?!. ")
    
    # Handle greetings
    if question_lower in GREETINGS:
        return f"Hello! I'm {BOT_INFO['name']}, your HR assistant. How can I help you today?"
    
    # Handle identity questions
    if any(q in question_lower for q in IDENTITY_QUESTIONS):
        if "who" in question_lower or "what is your name" in question_lower:
            return f"I'm {BOT_INFO['name']}, an AI assistant built by {BOT_INFO['creator']}. {BOT_INFO['responsibility']}"
        elif "created" in question_lower or "built" in question_lower:
            return f"I was created by {BOT_INFO['creator']} to {BOT_INFO['responsibility']}"
        elif "can you do" in question_lower or "do you do" in question_lower:
            return f"I can {BOT_INFO['capabilities']}"
        else:
            return f"I'm {BOT_INFO['name']}, an AI assistant created by {BOT_INFO['creator']}. {BOT_INFO['capabilities']}"

    # Handle holiday list queries (static 2025 list provided by HR)
    if "holiday" in question_lower or "holidays" in question_lower:
        year = "2025"
        header = f"## Company Holidays {year}\n\nBelow are the declared holidays for {year}.\n\n"

        india_table = (
            "### India Offices (Bangalore/APAC & EU, Hyderabad, Mumbai, Delhi)\n"
            "| Date | Day | Bangalore/APAC & EU | Hyderabad | Mumbai | Delhi |\n"
            "|------|-----|----------------------|-----------|--------|-------|\n"
            "| 1-Jan-2025 | Wednesday | New Year | New Year | New Year | New Year |\n"
            "| 14-Jan-2025 | Tuesday | Pongal/ Makar Sankranti | Pongal/ Makar Sankranti | Pongal/ Makar Sankranti | Pongal/ Makar Sankranti |\n"
            "| 14-Mar-2025 | Friday | - | Holi | Holi | Holi |\n"
            "| 31-Mar-2025 | Monday | Ramzan (Id Ul Fitr) | Ramzan (Id Ul Fitr) | Ramzan (Id Ul Fitr) | - |\n"
            "| 18-Apr-2025 | Friday | Good Friday | - | - | Good Friday |\n"
            "| 1-May-2025 | Thursday | May Day | May Day | May Day | May Day |\n"
            "| 15-Aug-2025 | Friday | Independence Day | Independence Day | Independence Day | Independence Day |\n"
            "| 27-Aug-2025 | Wednesday | Ganesh Chaturthi | Ganesh Chaturthi | Ganesh Chaturthi | Ganesh Chaturthi |\n"
            "| 2-Oct-2025 | Thursday | Gandhi Jayanthi/Dasara | Gandhi Jayanthi/Dasara | Gandhi Jayanthi/Dasara | Gandhi Jayanthi/Dasara |\n"
            "| 20-Oct-2025 | Monday | Diwali-Naraka Chaturdashi | Diwali-Naraka Chaturdashi | Diwali-Naraka Chaturdashi | Diwali-Naraka Chaturdashi |\n"
            "| 25-Dec-2025 | Thursday | Christmas | Christmas | Christmas | Christmas |\n"
        )

        us_table = (
            "\n### Global Services - US\n"
            "| Date | Day | Holiday |\n"
            "|------|-----|---------|\n"
            "| 1-Jan-2025 | Wednesday | New Year |\n"
            "| 18-Apr-2025 | Friday | Good Friday |\n"
            "| 26-May-2025 | Monday | Memorial Day |\n"
            "| 4-Jul-2025 | Friday | Independence Day |\n"
            "| 1-Sep-2025 | Monday | Labour Day |\n"
            "| 20-Oct-2025 | Monday | Diwali |\n"
            "| 27-Nov-2025 | Thursday | Thanksgiving |\n"
            "| 28-Nov-2025 | Friday | Day after Thanksgiving |\n"
            "| 24-Dec-2025 | Wednesday | Christmas Eve |\n"
            "| 25-Dec-2025 | Thursday | Christmas Day |\n"
        )

        footnote = "\n> Note: If a holiday falls on a weekend, local HR guidelines on compensatory off apply."
        return header + india_table + us_table + footnote
    
    return None

# Routes
# Authentication Routes
@app.route('/login')
def login():
    """Show login page or redirect to Google OAuth"""
    if 'user' in session:
        return redirect(url_for('index'))
    # Show login page with Google button
    return render_template('login.html')

@app.route('/login/google')
def login_google():
    """Initiate Google OAuth flow"""
    # Use the standard authlib callback path: /login/google/authorized
    redirect_uri = url_for('authorize', _external=True)
    return google.authorize_redirect(redirect_uri)

@app.route('/login/google/authorized')
def authorize():
    """Handle Google OAuth callback - matches Google Cloud Console redirect URI"""
    try:
        token = google.authorize_access_token()
        user_info = token.get('userinfo')
        
        if user_info:
            email = user_info.get('email')
            name = user_info.get('name', email.split('@')[0])
            
            # Create or update user in database
            create_or_update_user(email, name)
            
            # Get user info from database
            user_info = get_user_info(email)
            
            # Store user info in session
            session['user'] = {
                'email': email,
                'name': name,
                'role': user_info['role'] if user_info else 'Recruiter',
                'team': user_info['team'] if user_info else None
            }
            
            return redirect(url_for('index'))
        else:
            return jsonify({'error': 'Failed to get user info'}), 400
    except Exception as e:
        logging.error(f"OAuth error: {str(e)}")
        return jsonify({'error': f'Authentication failed: {str(e)}'}), 500

@app.route('/logout')
def logout():
    """Logout user"""
    session.pop('user', None)
    return redirect(url_for('login'))

@app.route('/')
@login_required
def index():
    """Landing page - requires login"""
    return render_template('index.html')

# Admin Panel Routes
@app.route('/admin')
@login_required
@role_required('Admin')
def admin_panel():
    """Admin panel for user/role/team management"""
    return render_template('admin.html')

@app.route('/api/admin/users', methods=['GET'])
@login_required
@role_required('Admin')
def get_all_users():
    """Get all users for admin panel"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        cursor.execute('''
            SELECT email, name, role, team, manager_email, created_at 
            FROM users 
            ORDER BY created_at DESC
        ''')
        users = []
        for row in cursor.fetchall():
            users.append({
                'email': row[0],
                'name': row[1],
                'role': row[2],
                'team': row[3],
                'manager_email': row[4],
                'created_at': row[5]
            })
        conn.close()
        return jsonify({'success': True, 'users': users})
    except Exception as e:
        logging.error(f"Error fetching users: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/admin/users', methods=['POST'])
@login_required
@role_required('Admin')
def update_user():
    """Update user role, team, or manager"""
    try:
        data = request.json
        email = data.get('email')
        role = data.get('role')
        team = data.get('team')
        manager_email = data.get('manager_email')
        name = data.get('name')
        
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        cursor.execute('''
            UPDATE users 
            SET role = ?, team = ?, manager_email = ?, name = ?, updated_at = CURRENT_TIMESTAMP
            WHERE email = ?
        ''', (role, team, manager_email, name, email))
        conn.commit()
        conn.close()
        
        return jsonify({'success': True, 'message': 'User updated successfully'})
    except Exception as e:
        logging.error(f"Error updating user: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/admin/users/<email>', methods=['DELETE'])
@login_required
@role_required('Admin')
def delete_user(email):
    """Delete a user (admin only)"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        cursor.execute('DELETE FROM users WHERE email = ?', (email,))
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'message': 'User deleted successfully'})
    except Exception as e:
        logging.error(f"Error deleting user: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/admin/teams', methods=['GET'])
@login_required
def get_teams():
    """Get list of available teams"""
    teams = ['ITS', 'OSS', 'PCS', 'ISV', 'Core']
    return jsonify({'success': True, 'teams': teams})

@app.route('/api/admin/roles', methods=['GET'])
@login_required
def get_roles():
    """Get list of available roles"""
    roles = ['Admin', 'Business Manager', 'Anchors', 'Recruiter']
    return jsonify({'success': True, 'roles': roles})

@app.route('/api/admin/accessible-users', methods=['GET'])
@login_required
def get_accessible_users_api():
    """Get users accessible to current user"""
    user_email = session['user'].get('email')
    users = get_accessible_users(user_email)
    return jsonify({'success': True, 'users': users})

@app.route('/hr-assistant')
@login_required
def hr_assistant():
    return render_template('index1.html')

@app.route('/resume-evaluator')
@login_required
def resume_evaluator():
    return render_template('index2.html')

@app.route('/test-tabs')
def test_tabs():
    return render_template('test_tabs.html')

@app.route('/evaluation/<int:eval_id>')
def view_evaluation(eval_id):
    """View a single evaluation in detail"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Fetch evaluation data
        cursor.execute('''
            SELECT 
                e.id, e.filename, e.job_title, e.job_description,
                e.match_percentage, e.match_factors, e.profile_summary,
                e.missing_keywords, e.job_stability, e.career_progression,
                e.oorwin_job_id, e.timestamp, e.resume_path
            FROM evaluations e
            WHERE e.id = ?
        ''', (eval_id,))
        
        row = cursor.fetchone()
        
        if not row:
            conn.close()
            return "Evaluation not found", 404
        
        # Fetch interview questions
        cursor.execute('''
            SELECT technical_questions, nontechnical_questions, behavioral_questions
            FROM interview_questions
            WHERE evaluation_id = ?
        ''', (eval_id,))
        
        questions_row = cursor.fetchone()
        conn.close()
        
        # Parse JSON fields
        import json
        resume_path_from_db = row[12] if len(row) > 12 else None
        
        # Normalize the path (handle Windows backslashes, remove quotes, etc.)
        if resume_path_from_db:
            resume_path_from_db = str(resume_path_from_db).strip().strip('"').strip("'")
            # Normalize path separators
            resume_path_from_db = os.path.normpath(resume_path_from_db)
            # Check if file actually exists
            if not os.path.exists(resume_path_from_db):
                logging.warning(f"Resume file not found at path: {resume_path_from_db}")
                resume_path_from_db = None
        
        # If resume_path is missing or empty, construct it from filename (for backward compatibility)
        if not resume_path_from_db or resume_path_from_db == 'None' or resume_path_from_db == '':
            filename = row[1] if row[1] else ''
            if filename:
                # Construct path from filename
                resume_path_from_db = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                resume_path_from_db = os.path.normpath(resume_path_from_db)
                # Check if file actually exists
                if not os.path.exists(resume_path_from_db):
                    logging.warning(f"Resume file not found at constructed path: {resume_path_from_db}")
                    resume_path_from_db = None
            else:
                resume_path_from_db = None
        
        # If we have a DOCX path, try to convert it to PDF or find existing PDF
        if resume_path_from_db and resume_path_from_db.lower().endswith('.docx'):
            # Check if PDF version already exists
            pdf_path = os.path.splitext(resume_path_from_db)[0] + '.pdf'
            pdf_path = os.path.normpath(pdf_path)
            if os.path.exists(pdf_path):
                logging.info(f"[VIEW] Found existing PDF version: {pdf_path}")
                resume_path_from_db = pdf_path
            else:
                # Try to convert DOCX to PDF on-the-fly
                logging.info(f"[VIEW] Attempting to convert DOCX to PDF for old evaluation: {resume_path_from_db}")
                converted_pdf = convert_docx_to_pdf(resume_path_from_db)
                if converted_pdf and os.path.exists(converted_pdf):
                    logging.info(f"[VIEW] ✅ Successfully converted DOCX to PDF: {converted_pdf}")
                    resume_path_from_db = os.path.normpath(converted_pdf)
                else:
                    logging.warning(f"[VIEW] ⚠️ Could not convert DOCX to PDF, merging won't work: {resume_path_from_db}")
                    # Keep DOCX path but merging will fail (user will see error)
        
        # Normalize match_percentage to an integer (handle NULL/None values safely)
        try:
            match_percentage_value = int(row[4]) if row[4] is not None else 0
        except (ValueError, TypeError):
            match_percentage_value = 0

        # Normalize match_factors values to integers (avoid None/strings in template comparisons)
        raw_match_factors = {}
        if row[5]:
            try:
                raw_match_factors = json.loads(row[5])
            except Exception as jf:
                logging.warning(f"Failed to parse match_factors for evaluation {row[0]}: {jf}")
                raw_match_factors = {}

        normalized_match_factors = {}
        for factor, value in raw_match_factors.items():
            try:
                normalized_match_factors[factor] = int(value) if value is not None else 0
            except (ValueError, TypeError):
                normalized_match_factors[factor] = 0

        evaluation = {
            'id': row[0],
            'filename': row[1],
            'job_title': row[2],
            'job_description': row[3],
            'match_percentage': match_percentage_value,
            'match_factors': normalized_match_factors,
            'profile_summary': row[6],
            'missing_keywords': json.loads(row[7]) if row[7] else [],
            'job_stability': json.loads(row[8]) if row[8] else {},
            'career_progression': json.loads(row[9]) if row[9] else {},
            'oorwin_job_id': row[10],
            'timestamp': row[11],
            'resume_path': resume_path_from_db  # resume_path (constructed if missing)
        }
        
        # Helper function to normalize questions (convert objects to strings)
        def normalize_questions(questions_list):
            if not questions_list:
                return []
            normalized = []
            for q in questions_list:
                if isinstance(q, str):
                    normalized.append(q)
                elif isinstance(q, dict):
                    # Extract question text from common property names
                    normalized.append(q.get('question') or q.get('text') or q.get('content') or q.get('value') or str(q))
                else:
                    normalized.append(str(q))
            return normalized
        
        if questions_row:
            tech_raw = json.loads(questions_row[0]) if questions_row[0] else []
            nontech_raw = json.loads(questions_row[1]) if questions_row[1] else []
            behavioral_raw = json.loads(questions_row[2]) if questions_row[2] else []
            
            evaluation['technical_questions'] = normalize_questions(tech_raw)
            evaluation['nontechnical_questions'] = normalize_questions(nontech_raw)
            evaluation['behavioral_questions'] = normalize_questions(behavioral_raw)
        else:
            evaluation['technical_questions'] = []
            evaluation['nontechnical_questions'] = []
            evaluation['behavioral_questions'] = []
        
        return render_template('evaluation_view.html', evaluation=evaluation)
        
    except Exception as e:
        logging.error(f"Error viewing evaluation {eval_id}: {str(e)}")
        return f"Error loading evaluation: {str(e)}", 500

@app.route('/history')
@login_required
def history():
    conn = sqlite3.connect('combined_db.db')
    cursor = conn.cursor()
    
    try:
        # Get accessible user emails based on role
        user_email = session['user'].get('email')
        accessible_emails = get_accessible_user_emails(user_email)
        
        if not accessible_emails:
            return render_template('history.html', evaluations=[])
        
        # Build WHERE clause for filtering
        placeholders = ','.join(['?'] * len(accessible_emails))
        where_clause = f"WHERE e.user_email IN ({placeholders})"
        
        cursor.execute(f'''
            SELECT 
                e.id, 
                e.filename, 
                e.job_title, 
                e.match_percentage, 
                e.missing_keywords, 
                e.profile_summary, 
                e.job_stability,
                e.career_progression,
                e.timestamp,
                iq.technical_questions,
                iq.nontechnical_questions,
                e.oorwin_job_id
            FROM evaluations e
            LEFT JOIN interview_questions iq ON e.id = iq.evaluation_id
            {where_clause}
            ORDER BY e.timestamp DESC
        ''', accessible_emails)
        
        evaluations = []
        for row in cursor.fetchall():
            try:
                # Helper function for safe JSON parsing
                def safe_json_loads(data, default):
                    if not data:
                        logging.info(f"Empty data for field, using default: {default}")
                        return default
                    try:
                        if isinstance(data, str):
                            return json.loads(data)
                        return data
                    except json.JSONDecodeError as e:
                        logging.error(f"JSON parsing error for evaluation {row[0]}: {str(e)} - Data: {data}")
                        # Try to clean the string if it's malformed
                        if isinstance(data, str):
                            try:
                                # Remove any trailing commas, fix quotes
                                cleaned = re.sub(r',\s*}', '}', data)
                                cleaned = re.sub(r',\s*]', ']', cleaned)
                                return json.loads(cleaned)
                            except:
                                pass
                        return default
                
                # Parse JSON fields with robust error handling
                missing_keywords_raw = row[4]
                try:
                    if missing_keywords_raw:
                        missing_keywords = safe_json_loads(missing_keywords_raw, [])
                        # If it's a string that looks like a list but isn't parsed as one
                        if not isinstance(missing_keywords, list):
                            # Try to extract keywords from a string representation
                            if isinstance(missing_keywords, str):
                                # Remove brackets and split by commas
                                missing_keywords = [k.strip(' "\'') for k in missing_keywords.strip('[]').split(',')]
                            else:
                                missing_keywords = [str(missing_keywords)]
                    else:
                        missing_keywords = []
                except Exception as e:
                    logging.error(f"Error parsing missing_keywords for eval {row[0]}: {str(e)}")
                    missing_keywords = []
                
                # Log raw data for debugging
                logging.info(f"Raw job_stability data for eval {row[0]}: {row[6]}")
                logging.info(f"Raw career_progression data for eval {row[0]}: {row[7]}")
                
                # Handle job stability data
                job_stability_data = row[6]
                if job_stability_data:
                    try:
                        job_stability = safe_json_loads(job_stability_data, {})
                        # Ensure it has the expected structure
                        if not isinstance(job_stability, dict):
                            job_stability = {}
                    except Exception as e:
                        logging.error(f"Error processing job_stability for eval {row[0]}: {str(e)}")
                        job_stability = {}
                else:
                    job_stability = {}
                
                # Handle career progression data
                career_progression_data = row[7]
                if career_progression_data:
                    try:
                        career_progression = safe_json_loads(career_progression_data, {})
                        # Ensure it has the expected structure
                        if not isinstance(career_progression, dict):
                            career_progression = {}
                    except Exception as e:
                        logging.error(f"Error processing career_progression for eval {row[0]}: {str(e)}")
                        career_progression = {}
                else:
                    career_progression = {}
                
                # Handle questions
                technical_questions = safe_json_loads(row[9], [])
                nontechnical_questions = safe_json_loads(row[10], [])
                
                # Ensure profile_summary is a valid string
                profile_summary = str(row[5]) if row[5] is not None else "No summary available"
                
                # Create a default structure for job_stability if empty
                if not job_stability:
                    job_stability = {
                        "StabilityScore": 0,
                        "AverageJobTenure": "N/A",
                        "JobCount": 0,
                        "RiskLevel": "N/A",
                        "ReasoningExplanation": "No job stability data available."
                    }
                
                # Create a default structure for career_progression if empty
                if not career_progression:
                    career_progression = {
                        "progression_score": 0,
                        "key_observations": [],
                        "career_path": [],
                        "red_flags": [],
                        "reasoning": "No career progression data available."
                    }
                
                # Ensure all data is properly serialized for the template
                # This is critical to avoid issues with the template's tojson filter
                try:
                    # Test serialization to catch any issues
                    json.dumps(job_stability)
                    json.dumps(career_progression)
                    json.dumps(technical_questions)
                    json.dumps(nontechnical_questions)
                except (TypeError, ValueError) as e:
                    logging.error(f"Serialization error for evaluation {row[0]}: {str(e)}")
                    # If there's an error, convert to string representation
                    if not isinstance(job_stability, dict):
                        job_stability = {"error": "Invalid data structure", "message": str(job_stability)}
                    if not isinstance(career_progression, dict):
                        career_progression = {"error": "Invalid data structure", "message": str(career_progression)}
                    if not isinstance(technical_questions, list):
                        technical_questions = ["Error loading technical questions"]
                    if not isinstance(nontechnical_questions, list):
                        nontechnical_questions = ["Error loading non-technical questions"]
                
                evaluation = {
                    'id': row[0],
                    'filename': row[1],
                    'job_title': row[2],
                    'match_percentage': row[3],
                    'missing_keywords': missing_keywords,
                    'profile_summary': profile_summary,
                    'job_stability': job_stability,
                    'career_progression': career_progression,
                    'timestamp': row[8],
                    'technical_questions': technical_questions,
                    'nontechnical_questions': nontechnical_questions,
                    'oorwin_job_id': row[11]
                }
                evaluations.append(evaluation)
                
                # Log the processed data for debugging
                logging.info(f"Processed evaluation {row[0]}: job_stability={job_stability}, career_progression={career_progression}")
                
            except Exception as e:
                logging.error(f"Error processing row for evaluation {row[0]}: {str(e)}")
                continue
            
        return render_template('history.html', evaluations=evaluations)
    
    except Exception as e:
        logging.error(f"Error in history route: {str(e)}")
        return render_template('history.html', evaluations=[], error="Failed to load evaluations")
    
    finally:
        conn.close()

@app.route('/feedback_history')
def feedback_history():
    """Display unified feedback history page"""
    return render_template('feedback_history.html')

@app.route('/api/feedback/all')
def get_all_feedback():
    """Get all feedback from all 3 sources"""
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()

        # Get HR Assistant feedback
        cursor.execute("""
            SELECT qf.id, qf.question_id, qf.rating, qf.feedback, qf.timestamp,
                   qh.question, qh.final_answer
            FROM qa_feedback qf
            JOIN qa_history qh ON qf.question_id = qh.id
            ORDER BY qf.timestamp DESC
        """)
        hr_assistant = []
        for row in cursor.fetchall():
            hr_assistant.append({
                'id': row[0],
                'question_id': row[1],
                'rating': row[2],
                'feedback': row[3],
                'timestamp': row[4],
                'question': row[5],
                'answer': row[6]
            })

        # Get Handbook feedback
        cursor.execute("""
            SELECT hf.id, hf.handbook_id, hf.rating, hf.comments, hf.timestamp,
                   rh.job_title, rh.oorwin_job_id, rh.markdown_content
            FROM handbook_feedback hf
            JOIN recruiter_handbooks rh ON hf.handbook_id = rh.id
            ORDER BY hf.timestamp DESC
        """)
        handbooks = []
        for row in cursor.fetchall():
            handbooks.append({
                'id': row[0],
                'handbook_id': row[1],
                'rating': row[2],
                'comments': row[3],
                'timestamp': row[4],
                'job_title': row[5],
                'oorwin_job_id': row[6],
                'markdown_content': row[7]
            })

        # Get Evaluation feedback
        cursor.execute("""
            SELECT f.id, f.evaluation_id, f.rating, f.comments, f.timestamp,
                   e.filename, e.job_title, e.match_percentage, e.oorwin_job_id,
                   e.match_factors, e.profile_summary, e.missing_keywords,
                   e.job_stability, e.career_progression, e.technical_questions,
                   e.nontechnical_questions, e.behavioral_questions
            FROM feedback f
            JOIN evaluations e ON f.evaluation_id = e.id
            ORDER BY f.timestamp DESC
        """)
        evaluations = []
        for row in cursor.fetchall():
            evaluations.append({
                'id': row[0],
                'evaluation_id': row[1],
                'rating': row[2],
                'comments': row[3],
                'timestamp': row[4],
                'filename': row[5],
                'job_title': row[6],
                'match_percentage': row[7],
                'oorwin_job_id': row[8],
                'match_factors': row[9],
                'profile_summary': row[10],
                'missing_keywords': row[11],
                'job_stability': row[12],
                'career_progression': row[13],
                'technical_questions': row[14],
                'nontechnical_questions': row[15],
                'behavioral_questions': row[16]
            })

        conn.close()

        return jsonify({
            'success': True,
            'hr_assistant': hr_assistant,
            'handbooks': handbooks,
            'evaluations': evaluations
        })

    except Exception as e:
        logging.error(f"Error getting all feedback: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/feedback/check/<feedback_type>/<int:item_id>')
def check_feedback_exists(feedback_type, item_id):
    """Check if feedback already exists for a given item"""
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        if feedback_type == 'qa':
            cursor.execute("SELECT id FROM qa_feedback WHERE question_id = ?", (item_id,))
        elif feedback_type == 'handbook':
            cursor.execute("SELECT id FROM handbook_feedback WHERE handbook_id = ?", (item_id,))
        elif feedback_type == 'evaluation':
            cursor.execute("SELECT id FROM feedback WHERE evaluation_id = ?", (item_id,))
        else:
            return jsonify({'success': False, 'error': 'Invalid feedback type'}), 400
        
        exists = cursor.fetchone() is not None
        conn.close()
        
        return jsonify({'success': True, 'exists': exists})
        
    except Exception as e:
        logging.error(f"Error checking feedback: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/feedback/handbook', methods=['POST'])
def submit_handbook_feedback():
    """Submit feedback for a recruiter handbook"""
    try:
        data = request.get_json()
        
        if not data or 'handbook_id' not in data or 'rating' not in data:
            return jsonify({'success': False, 'error': 'Missing handbook_id or rating'}), 400
        
        handbook_id = data['handbook_id']
        rating = data['rating']
        comments = data.get('comments', '')
        
        # Validate rating
        if not isinstance(rating, int) or rating < 1 or rating > 5:
            return jsonify({'success': False, 'error': 'Rating must be between 1 and 5'}), 400
        
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # Check if feedback already exists
        cursor.execute("SELECT id FROM handbook_feedback WHERE handbook_id = ?", (handbook_id,))
        if cursor.fetchone():
            conn.close()
            return jsonify({'success': False, 'error': 'Feedback already submitted for this handbook'}), 400
        
        # Insert feedback
        cursor.execute("""
            INSERT INTO handbook_feedback (handbook_id, rating, comments)
            VALUES (?, ?, ?)
        """, (handbook_id, rating, comments))
        
        conn.commit()
        conn.close()
        
        return jsonify({'success': True, 'message': 'Feedback submitted successfully'})
        
    except sqlite3.IntegrityError:
        return jsonify({'success': False, 'error': 'Feedback already submitted for this handbook'}), 400
    except Exception as e:
        logging.error(f"Error submitting handbook feedback: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

# HR Assistant Routes
def get_info_buddy_generation_config(long_answer: bool) -> dict:
    """Return generation and retrieval config for Info Buddy answers.

    - Short answers: lower max_tokens and shallower retrieval
    - Long answers:  higher max_tokens and deeper retrieval
    """
    return {
        "max_tokens": 1200 if long_answer else 50,  # very strict budget for short answers (50 tokens ≈ 35-40 words)
        "retrieval_top_k": 6 if long_answer else 3,
    }


def build_info_buddy_online_prompt(expanded_question: str, long_answer: bool) -> str:
    """Build the prompt for Info Buddy in ONLINE mode."""
    if long_answer:
        # Detailed / long answer
        return f"""You are an expert AI assistant. Provide a comprehensive and detailed answer to the following question. Your response should be thorough, well-structured, and accurate.

Question: {expanded_question}

Instructions:
1. Use your knowledge to provide a complete answer.
2. If the question is about HR policies, benefits, or company-specific information, note that you may not have the latest company-specific details.
3. Format your response with clear sections and bullet points where appropriate.
4. Include relevant examples and context.
5. If you're uncertain about specific facts, mention that.
"""
    else:
        # Short / concise answer
        return f"""You are an expert AI assistant. Provide a **very concise** answer to the following question.

Question: {expanded_question}

Instructions for brevity (STRICT - ENFORCE STRICTLY):
1. Your entire answer MUST be under 40 words (approximately 30-35 words).
2. Use ONLY 2-3 very short bullet points OR a single 2-sentence paragraph.
3. NO headings, NO sections, NO deep dives, NO tables.
4. Highlight ONLY the single most critical point the user needs to know.
5. Be direct and brief. Stop immediately after the key information.
"""


def build_info_buddy_rag_prompt(expanded_question: str, context: str, long_answer: bool) -> str:
    """Build the prompt for Info Buddy in RAG (offline) mode."""
    if context:
        if long_answer:
            # Detailed / long RAG answer
            return f"""You are an expert HR Assistant. Answer the question STRICTLY based on the provided context from company policy documents.

Question: {expanded_question}

Context from Company Documents:
{context}

STRICT RULES FOR FORMATTING:

1. **ONLY use information from the context provided above.** Do not use external knowledge.

2. **TABLE FORMATTING (CRITICAL)**:
   - If the context contains tables (look for markers like [TABLE DATA] ... [END TABLE]), you MUST reproduce the markdown table **verbatim** from the context.
   - Use this EXACT format:
     ```
     | Column 1 | Column 2 | Column 3 |
     |---------|---------|----------|
     | Data 1  | Data 2  | Data 3   |
     | Data 4  | Data 5  | Data 6   |
     ```
   - Always include the header separator row (|---------|---------|).
   - Keep table columns aligned and readable.
   - If a table is complex, break it into smaller, clearer tables.
   - Add a brief title above each table (e.g., "### Performance Rating Table").

3. **RESPONSE STRUCTURE**:
   - Start with a brief 1–2 sentence summary.
   - Use clear headings (## for main sections, ### for subsections).
   - Use bullet points (• or -) for lists, NOT long paragraphs.
   - Add blank lines between sections for readability.
   - Keep paragraphs SHORT (3–4 sentences max).
   - Use bold (**text**) for key terms and important points.

4. **SOURCE CITATION**:
   - **ALWAYS cite sources using actual document names** (e.g., "According to [Leave Policy.pdf, page 5]").
   - Place citations at the END of sentences or paragraphs, not mid-sentence.
   - Format: `[Document Name.pdf, page X]`.
   - Use actual filename, not generic "Source 1" or "Source 2".

5. **READABILITY ENHANCEMENTS**:
   - Break dense information into digestible chunks.
   - Use numbered lists (1., 2., 3.) for step-by-step processes.
   - Use bullet lists (•) for features, benefits, or items.
   - Add horizontal rules (---) to separate major sections.
   - Use emojis sparingly for visual breaks (✅, 📋, 📊, etc.).

6. **EXAMPLE OF GOOD FORMATTING**:
   ```
   ## Performance Appraisal Policy
   
   The performance appraisal policy outlines how employees are evaluated annually.
   
   ### Key Features
   • Appraisals are conducted yearly (April to March).
   • Based on previously agreed KRAs.
   • Results can lead to salary hikes or promotions.
   
   ### Performance Rating Table
   | Score | Rating      | Description                           |
   |-------|-------------|---------------------------------------|
   | 5     | Exceptional | Targets met at 200% or above         |
   | 4     | Outstanding | Targets exceeded significantly       |
   | 3     | Good        | Consistently met expectations        |
   
   [APPRAISAL & PROMOTION POLICY.pdf, page 1]
   ```

7. **IMPORTANT**: If the question cannot be answered from the provided context, you MUST respond with:
   - "I'm sorry, but the information about '[topic]' is not available in our company policy documents."
   - "💡 **Suggestion**: Please enable the **'Go Online'** toggle and try asking your question again."

8. DO NOT make up information or use knowledge outside the provided context.

Answer:"""
        else:
            # Short / concise RAG answer
            return f"""You are an expert HR Assistant. Answer the question STRICTLY based on the provided context from company policy documents.

Question: {expanded_question}

Context from Company Documents:
{context}

Instructions for a SHORT answer (STRICT - ENFORCE STRICTLY):
1. Your entire answer MUST be under 40 words (approximately 30-35 words).
2. Use ONLY 2-3 very short bullet points OR a single 2-sentence paragraph.
3. Focus ONLY on the single most critical policy rule or action the user needs to know.
4. NO tables, NO headings, NO sections, NO deep dives.
5. Cite sources briefly like `[Document.pdf, page X]` at the end if space allows.
6. Do NOT invent information outside the context.
7. Be direct and brief. Stop immediately after the key information.

Short Answer:"""
    # No context found - give helpful message and gently suggest online mode
    return f"""The question "{expanded_question}" could not be answered from the available company policy documents.

Please note:
- The information may not be in the current knowledge base.
- The document may need to be updated or added.
- You can try rephrasing the question or enabling "Go Online" mode for general information.

Would you like to:
1. Try rephrasing your question
2. Enable "Go Online" mode for general information
3. Contact HR for company-specific policies not yet in the system?"""


@app.route('/api/ask', methods=['POST'])
def ask_question():
    try:
        data = request.get_json()
        question = data.get('question')
        online_mode = data.get('online_mode', False)
        # Control answer length from frontend (Info Buddy)
        # long_answer=True  -> detailed response
        # long_answer=False -> concise/short response
        long_answer = data.get('long_answer', False)

        if not question:
            return jsonify({'error': 'No question provided'}), 400

        def generate():
            complete_response = []  # Store complete response
            try:
                # Check for special queries first
                special_response = handle_special_queries(question)
                if special_response:
                    complete_response.append(special_response)
                    yield special_response
                    return

                # Expand acronyms in the question
                expanded_question = expand_acronyms(question)

                # Configure generation and retrieval depth based on answer length
                config = get_info_buddy_generation_config(long_answer)
                max_tokens = config["max_tokens"]
                retrieval_top_k = config["retrieval_top_k"]

                if online_mode:
                    # ONLINE MODE: Answer any general question using LLM knowledge
                    # No RAG constraints - can answer anything
                    detailed_prompt = build_info_buddy_online_prompt(expanded_question, long_answer)
                    response = generate_content_unified(detailed_prompt, stream=True, max_tokens=max_tokens)
                    for chunk in response:
                        if chunk.text:
                            complete_response.append(chunk.text)
                            yield chunk.text
                else:
                    # RAG MODE: Strict retrieval from local documents only
                    # Use hybrid search (BM25 + Vector) for better coverage
                    
                    # Step 1: Hybrid retrieval - combine BM25 and Vector search
                    all_retrieved_docs = []
                    
                    # Vector search (semantic similarity)
                    if vectorstore is not None:
                        vector_docs = vectorstore.similarity_search(expanded_question, k=retrieval_top_k)
                        all_retrieved_docs.extend([(doc, 'vector') for doc in vector_docs])
                        logging.info(f"🔍 Vector search retrieved {len(vector_docs)} documents")
                    
                    # BM25 search (keyword matching) - better for exact terms and tables
                    if bm25_index and bm25_corpus:
                        try:
                            query_tokens = expanded_question.lower().split()
                            bm25_scores = bm25_index.get_scores(query_tokens)
                            
                            # Get top BM25 results - depth controlled by retrieval_top_k
                            top_indices = sorted(range(len(bm25_scores)), key=lambda i: bm25_scores[i], reverse=True)[:retrieval_top_k]
                            bm25_results = []
                            for idx in top_indices:
                                if bm25_scores[idx] > 0:  # Only include relevant results
                                    text_content = " ".join(bm25_corpus[idx])
                                    # Create a Document-like object with metadata for consistency
                                    from langchain_core.documents import Document as LangchainDocument
                                    # Get metadata for this chunk if available
                                    metadata = bm25_metadata[idx] if idx < len(bm25_metadata) else {}
                                    bm25_doc = LangchainDocument(page_content=text_content, metadata=metadata)
                                    bm25_results.append(bm25_doc)
                                    all_retrieved_docs.append((bm25_doc, 'bm25'))
                            
                            logging.info(f"🔍 BM25 search retrieved {len(bm25_results)} documents")
                        except Exception as e:
                            logging.warning(f"⚠️ BM25 search failed: {e}")
                    
                    # Step 2: Deduplicate and prioritize
                    seen_content = set()
                    unique_docs = []
                    table_docs = []
                    text_docs = []
                    
                    for doc, source in all_retrieved_docs:
                        content_hash = hash(doc.page_content[:100])  # Hash first 100 chars for dedup
                        if content_hash not in seen_content:
                            seen_content.add(content_hash)
                            # Classify as table or text
                            if "[TABLE DATA]" in doc.page_content or ("|" in doc.page_content and doc.page_content.count("|") > 3):
                                table_docs.append(doc)
                            else:
                                text_docs.append(doc)
                            unique_docs.append(doc)
                    
                    # Step 3: Prioritize tables and limit context window
                    # Tables first (they're often most precise), then text chunks
                    prioritized_docs = table_docs + text_docs
                    context_docs = prioritized_docs[:retrieval_top_k]  # Depth controlled by answer length
                    
                    # Step 4: Build context with proper source citations (filename + page)
                    if context_docs:
                        context_parts = []
                        for i, doc in enumerate(context_docs):
                            # Extract actual source filename and page from metadata
                            source_name = doc.metadata.get('source', 'Unknown Document') if hasattr(doc, 'metadata') and doc.metadata else 'Unknown Document'
                            page_num = doc.metadata.get('page', 'N/A') if hasattr(doc, 'metadata') and doc.metadata else 'N/A'
                            
                            # Format citation with actual filename
                            if source_name != 'Unknown Document':
                                citation = f"{source_name}"
                                if page_num != 'N/A':
                                    citation += f", page {page_num}"
                            else:
                                citation = f"Source {i+1}"
                            
                            # Add relevance indicator for tables with proper citation
                            if doc in table_docs:
                                context_parts.append(f"[RELEVANT TABLE DATA - {citation}]\n{doc.page_content}")
                            else:
                                context_parts.append(f"[RELEVANT CONTEXT - {citation}]\n{doc.page_content}")
                        context = "\n\n---\n\n".join(context_parts)
                        
                        # Log retrieval stats
                        logging.info(f"📚 Total unique documents retrieved: {len(unique_docs)} ({len(table_docs)} tables, {len(text_docs)} text)")
                        if table_docs:
                            logging.info(f"📊 Including {len(table_docs)} table chunks in context")
                    else:
                        context = ""
                        logging.warning("⚠️ No documents retrieved from knowledge base")
                    
                    # Step 5: Build RAG prompt (short or long) and generate answer
                    prompt = build_info_buddy_rag_prompt(expanded_question, context, long_answer)
                    response = generate_content_unified(prompt, stream=True, max_tokens=max_tokens)
                    for chunk in response:
                        if chunk.text:
                            complete_response.append(chunk.text)
                            yield chunk.text

                # Store the complete Q&A in history after streaming is done
                final_answer = "".join(complete_response)
                conn = sqlite3.connect('combined_db.db')
                c = conn.cursor()
                c.execute('''INSERT INTO qa_history (question, retrieved_docs, final_answer)
                            VALUES (?, ?, ?)''', (question, None, final_answer))
                conn.commit()
                conn.close()

            except Exception as e:
                error_msg = f"Error: {str(e)}"
                yield error_msg
                # Store error in database
                conn = sqlite3.connect('combined_db.db')
                c = conn.cursor()
                c.execute('''INSERT INTO qa_history (question, final_answer)
                            VALUES (?, ?)''', (question, error_msg))
                conn.commit()
                conn.close()

        return Response(stream_with_context(generate()), mimetype='text/plain')

    except Exception as e:
        print(f"Error in ask_question: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route("/api/update_index", methods=["POST"])
def update_index_api():
    """Manually refresh the Pinecone & BM25 index."""
    try:
        # Rebuild BM25 index
        build_bm25_index(POLICIES_FOLDER)
        
        # Repopulate Pinecone
        populate_pinecone_index()
        
        return jsonify({"message": "Indexes updated successfully"}), 200
    except Exception as e:
        logging.error(f"❌ Index Update Error: {e}", exc_info=True)
        return jsonify({"error": "Failed to update indexes"}), 500

# Resume Evaluator Routes
def extract_json_from_text(text):
    """Extract JSON object from text, handling deep nesting properly"""
    # Find the first { and match with its corresponding }
    start_idx = text.find('{')
    if start_idx == -1:
        return None
    
    # Count braces to find matching closing brace
    brace_count = 0
    in_string = False
    escape_next = False
    
    for i in range(start_idx, len(text)):
        char = text[i]
        
        # Handle string detection (to ignore braces inside strings)
        if char == '\\' and not escape_next:
            escape_next = True
            continue
        
        if char == '"' and not escape_next:
            in_string = not in_string
        
        escape_next = False
        
        # Count braces only when not inside a string
        if not in_string:
            if char == '{':
                brace_count += 1
            elif char == '}':
                brace_count -= 1
                if brace_count == 0:
                    # Found matching closing brace
                    return text[start_idx:i+1]
    
    return None

async def async_groq_generate_for_scoring(prompt):
    """Async wrapper for Groq generation with temperature 0 for consistent scoring"""
    global groq_client
    try:
        # Runtime check: if client is not initialized, try to initialize it
        if not groq_client:
            # Reload environment variables
            BASE_DIR = os.path.dirname(os.path.abspath(__file__))
            ENV_PATH = os.path.join(BASE_DIR, '.env')
            load_dotenv(ENV_PATH, override=True)
            
            # Get API key again
            api_key = os.getenv("GROQ_API_KEY", "").strip()
            if not api_key:
                raise RuntimeError("Groq client not initialized. Please set GROQ_API_KEY in .env file.")
            
            # Try to initialize
            try:
                groq_client = Groq(api_key=api_key)
                logging.info("[RUNTIME] Groq client initialized successfully")
            except Exception as e:
                raise RuntimeError(f"Failed to initialize Groq client: {str(e)}. Please check your GROQ_API_KEY in .env file.")
        
        # Use Groq with temperature 0 for deterministic scoring
        is_reasoning_model = any(x in GROQ_MODEL.lower() for x in ["gpt-oss", "o1-", "o3-"])
        supports_temperature = not is_reasoning_model
        
        params = {
            "model": GROQ_MODEL,
            "messages": [
                {"role": "system", "content": "You are an expert HR analyst and technical recruiter."},
                {"role": "user", "content": prompt}
            ],
            "max_completion_tokens": 16384
        }
        
        # Set temperature to 0 for scoring (deterministic results)
        if supports_temperature:
            params["temperature"] = 0.0
        else:
            logging.warning(f"Model {GROQ_MODEL} is a reasoning model and doesn't support temperature. Using default.")
        
        # Add reasoning_effort for reasoning models
        if is_reasoning_model:
            params["reasoning_effort"] = GROQ_REASONING_EFFORT
        
        response = groq_client.chat.completions.create(**params)
        
        # Extract content and token usage
        if not response.choices:
            raise Exception("Groq returned empty choices")
        
        message = response.choices[0].message
        content = message.content
        
        if not content:
            raise Exception(f"Groq returned empty content. Finish reason: {response.choices[0].finish_reason}")
        
        # Get token usage from response
        token_usage = None
        if hasattr(response, 'usage') and response.usage:
            if hasattr(response.usage, 'total_tokens'):
                token_usage = response.usage.total_tokens
            elif isinstance(response.usage, dict) and 'total_tokens' in response.usage:
                token_usage = response.usage['total_tokens']
        
        # Return both content and token usage
        response_obj = UnifiedModelResponse(content)
        response_obj.token_usage = token_usage
        return response_obj
        
    except Exception as e:
        logging.error(f"❌ Groq generation error: {str(e)}", exc_info=True)
        raise

async def async_gemini_generate(prompt):
    """Async wrapper for model generation (Gemini or OpenAI) with improved JSON handling"""
    try:
        response = generate_content_unified(prompt)
        
        # Check if response has text attribute
        if not hasattr(response, 'text'):
            logging.error(f"❌ Response object has no 'text' attribute. Type: {type(response)}")
            logging.error(f"   Response attributes: {dir(response)}")
            if "JD Match" in str(prompt):
                return get_default_resume_evaluation()
            else:
                return get_default_career_analysis()
        
        response_text = response.text
        
        # Check if response_text is None or empty
        if response_text is None:
            logging.error("❌ response.text is None")
            if "JD Match" in str(prompt):
                return get_default_resume_evaluation()
            else:
                return get_default_career_analysis()
        
        if not isinstance(response_text, str):
            logging.error(f"❌ response.text is not a string: {type(response_text)}")
            if "JD Match" in str(prompt):
                return get_default_resume_evaluation()
            else:
                return get_default_career_analysis()
        
        response_text = response_text.strip()
        
        # Reduced logging for performance (only log on errors)
        # logging.info(f"✅ Raw response length: {len(response_text)} chars")
        # logging.info(f"📄 Raw response preview (first 300): {response_text[:300]}...")
        # logging.info(f"📄 Raw response preview (last 200): ...{response_text[-200:]}")
        
        # Remove markdown code block markers if present (do this before brace checks)
        response_text = re.sub(r'^```json\s*', '', response_text, flags=re.MULTILINE)
        response_text = re.sub(r'^```\s*', '', response_text, flags=re.MULTILINE)
        response_text = re.sub(r'\s*```$', '', response_text)
        response_text = response_text.strip()

        # If still wrapped in fences, remove them greedily
        if response_text.startswith("```") and response_text.endswith("```"):
            response_text = response_text[3:-3].strip()

        # Check if response looks like it starts with explanatory text
        first_char = response_text.strip()[0] if response_text.strip() else ''
        if first_char != '{':
            logging.warning(f"⚠️ Response doesn't start with '{{'. First 100 chars: {response_text[:100]}")
            # Try to find where JSON actually starts
            json_start = response_text.find('{')
            if json_start > 0:
                logging.info(f"   Found '{{' at position {json_start}, removing {json_start} chars before it")
                response_text = response_text[json_start:]
                logging.info(f"   After trim, first 100 chars: {response_text[:100]}")

        # Final check - must start with {
        if not response_text.startswith('{'):
            logging.error(f"❌ Response still doesn't start with '{{' after cleaning")
            logging.error(f"   First 200 chars: {response_text[:200]}")
            if "JD Match" in str(prompt):
                return get_default_resume_evaluation()
            else:
                return get_default_career_analysis()
        
        # Try to parse as JSON directly first
        try:
            parsed = json.loads(response_text)
            # logging.info("✅ Direct JSON parse successful")  # Reduced logging
            # Final validation - must be a dict
            if not isinstance(parsed, dict):
                logging.error(f"❌ Parsed JSON is not a dict: {type(parsed)}")
                if "JD Match" in str(prompt):
                    return get_default_resume_evaluation()
                else:
                    return get_default_career_analysis()
            return parsed
        except json.JSONDecodeError as e:
            # Only log if extraction also fails
            # logging.warning(f"⚠️ Direct JSON parse failed: {str(e)}")
            # logging.warning(f"   Error at position {e.pos}: {response_text[max(0,e.pos-20):e.pos+20]}")
            
            # Try to extract JSON object with proper brace matching
            json_str = extract_json_from_text(response_text)
            if json_str:
                # Validate extracted string looks like JSON
                json_str = json_str.strip()
                if not json_str.startswith('{') or not json_str.endswith('}'):
                    logging.error(f"❌ Extracted string doesn't look like JSON")
                    logging.error(f"   Starts with: {json_str[:50]}")
                    logging.error(f"   Ends with: {json_str[-50:]}")
                    if "JD Match" in str(prompt):
                        return get_default_resume_evaluation()
                    else:
                        return get_default_career_analysis()
                
                try:
                    parsed = json.loads(json_str)
                    # logging.info(f"✅ Extracted JSON successfully (length: {len(json_str)} chars)")  # Reduced logging
                    # Validate it's a dict with expected structure
                    if not isinstance(parsed, dict):
                        logging.error(f"❌ Parsed JSON is not a dict: {type(parsed)}")
                        logging.error(f"   Parsed value: {repr(parsed)[:200]}")
                        if "JD Match" in str(prompt):
                            return get_default_resume_evaluation()
                        else:
                            return get_default_career_analysis()
                    # Double-check it has expected keys for resume evaluation
                    if "JD Match" in str(prompt) and "JD Match" not in parsed:
                        logging.error(f"❌ Parsed dict missing 'JD Match' key")
                        logging.error(f"   Available keys: {list(parsed.keys())}")
                        return get_default_resume_evaluation()
                    return parsed
                except json.JSONDecodeError as e2:
                    logging.error(f"❌ Failed to parse extracted JSON: {str(e2)}")
                    logging.error(f"   Error position: {e2.pos}")
                    logging.error(f"   Extracted text (first 500 chars): {json_str[:500]}...")
                    logging.error(f"   Extracted text (last 200 chars): ...{json_str[-200:]}")
                    logging.error(f"   Extracted text around error: ...{json_str[max(0,e2.pos-50):e2.pos+50]}...")
                    # Check if this is a resume evaluation prompt (has "JD Match" in prompt)
                    if "JD Match" in str(prompt):
                        return get_default_resume_evaluation()
                    else:
                        return get_default_career_analysis()
            else:
                logging.error(f"❌ No JSON object found in response")
                logging.error(f"   Response text (first 500 chars): {response_text[:500]}...")
                # Check if this is a resume evaluation prompt
                if "JD Match" in str(prompt):
                    return get_default_resume_evaluation()
                else:
                    return get_default_career_analysis()
                
    except Exception as e:
        logging.error(f"❌ Model generation error: {str(e)}", exc_info=True)
        logging.error(f"   Error type: {type(e).__name__}")
        logging.error(f"   Error message: {str(e)}")
        # Check if this is a resume evaluation prompt
        if "JD Match" in str(prompt):
            return get_default_resume_evaluation()
        else:
            return get_default_career_analysis()

async def async_analyze_stability(resume_text):
    """Async job stability analysis"""
    try:
        stability_prompt = job_stability_prompt.format(resume_text=resume_text)
        response = await async_gemini_generate(stability_prompt)
        
        if not response:
            raise ValueError("Failed to get stability analysis")
            
        # Ensure all required fields exist
        default_data = {
            "IsStable": True,
            "AverageJobTenure": "Unknown",
            "JobCount": 0,
            "StabilityScore": 0,
            "ReasoningExplanation": "Could not analyze job stability",
            "RiskLevel": "Unknown"
        }
        
        # Merge response with defaults
        for key, default_value in default_data.items():
            if key not in response:
                response[key] = default_value
                
        return response
        
    except Exception as e:
        logging.error(f"Error in async_analyze_stability: {str(e)}")
        return {
            "IsStable": True,
            "AverageJobTenure": "Unknown",
            "JobCount": 0,
            "StabilityScore": 0,
            "ReasoningExplanation": "Could not analyze job stability",
            "RiskLevel": "Unknown"
        }

async def async_generate_questions(resume_text, job_description, profile_summary):
    """Async interview questions generation"""
    try:
        questions_prompt = interview_questions_prompt.format(
            resume_text=resume_text,
            job_description=job_description,
            profile_summary=profile_summary
        )
        response = await async_gemini_generate(questions_prompt)
        
        if not response:
            raise ValueError("Failed to generate interview questions")
            
        # Ensure we have the required fields with proper defaults
        default_data = {
            "TechnicalQuestions": [],
            "NonTechnicalQuestions": []
        }
        
        # Merge response with defaults
        for key, default_value in default_data.items():
            if key not in response:
                response[key] = default_value
            elif not isinstance(response[key], list):
                response[key] = [str(response[key])] if response[key] else []
                
        return response
        
    except Exception as e:
        logging.error(f"Error in async_generate_questions: {str(e)}")
        return {
            "TechnicalQuestions": [],
            "NonTechnicalQuestions": []
        }

async def async_generate_recruiter_handbook(resume_text, job_description):
    """Async recruiter handbook generation - returns markdown text"""
    try:
        handbook_prompt = recruiter_handbook_prompt.format(
            resume_text=resume_text,
            job_description=job_description
        )
        
        # Use selected model to generate the recruiter handbook (run in thread pool to avoid blocking)
        # This returns markdown text, not JSON
        loop = asyncio.get_event_loop()
        response = await loop.run_in_executor(None, lambda: generate_content_unified(handbook_prompt))
        response_text = response.text.strip()
        
        if not response_text:
            raise ValueError("Failed to generate recruiter handbook")
        
        # Return just the markdown content (JSON summary removed as per user request)
        return {
            "markdown_content": response_text,
            "json_summary": None
        }
        
    except Exception as e:
        logging.error(f"Error in async_generate_recruiter_handbook: {str(e)}")
        return {
            "markdown_content": f"## Error\n\nFailed to generate recruiter handbook: {str(e)}",
            "json_summary": None
        }

@app.route('/evaluate', methods=['POST'])
@login_required
async def evaluate_resume():
    try:
        if 'resume' not in request.files:
            return jsonify({'error': 'No resume file provided'}), 400
        
        file = request.files['resume']
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type'}), 400

        job_title = request.form.get('job_title')
        job_description = request.form.get('job_description')

        if not job_title or not job_description:
            return jsonify({'error': 'Missing job title or description'}), 400

        additional_context = request.form.get('additional_context', '').strip()
        additional_context_block = f"**Additional Context (client constraints/preference):** {additional_context}" if additional_context else ""

        # Save uploaded file
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        # Convert DOCX to PDF if needed, and get the PDF path for merging
        pdf_path = None
        file_ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
        if file_ext == 'docx':
            # Convert DOCX to PDF
            pdf_path = convert_docx_to_pdf(file_path)
            if pdf_path:
                logging.info(f"Converted DOCX to PDF: {pdf_path}")
            else:
                logging.warning(f"Failed to convert DOCX to PDF, will use original file: {file_path}")
                pdf_path = file_path  # Fallback to original
        elif file_ext == 'pdf':
            pdf_path = file_path  # Already a PDF
        else:
            pdf_path = file_path  # Fallback

        # Extract text from resume (use original file path for text extraction)
        resume_text = extract_text_from_file(file_path)
        if resume_text is None:
            return jsonify({'error': 'Failed to extract text from file'}), 500

        # Generate evaluation using Gemini with optimized parameters
        formatted_prompt = input_prompt_template.format(
            resume_text=resume_text,
            job_description=job_description,
            additional_context_block=additional_context_block
        )
        
        try:
            # Run all analyses concurrently using asyncio.gather
            main_response, stability_data, career_data = await asyncio.gather(
                async_gemini_generate(formatted_prompt),
                async_analyze_stability(resume_text),
                analyze_career_progression(resume_text)  # Now properly awaited
            )
            
            if not main_response:
                raise ValueError("Failed to get main evaluation response")
                
            if not career_data:
                career_data = {
                    "progression_score": 50,
                    "key_observations": ["Failed to analyze career progression"],
                    "career_path": [],
                    "red_flags": ["Analysis error"],
                    "reasoning": "Failed to process career data"
                }
                
        except Exception as e:
            logging.error(f"Error during concurrent analysis: {str(e)}")
            return jsonify({'error': 'Failed to analyze resume'}), 500
        
        # Extract values from main response
        match_percentage_str = main_response.get("JD Match", "0%")
        match_percentage = int(match_percentage_str.strip('%'))
        missing_keywords = main_response.get("MissingKeywords", [])
        profile_summary = main_response.get("Profile Summary", "No summary provided.")
        over_under_qualification = main_response.get("Over/UnderQualification Analysis", "No qualification mismatch concerns detected.")
        match_factors = main_response.get("Match Factors", {})
        candidate_fit_analysis = main_response.get("Candidate Fit Analysis", {})

        # Prepare additional information
        additional_info = {
            "job_stability": stability_data,
            "career_progression": career_data,
            "reasoning": main_response.get("Reasoning", "")
        }

        # Generate unique ID for evaluation
        eval_id = str(uuid.uuid4())

        # Get user email from session
        user_email = session.get('user', {}).get('email') if 'user' in session else None
        
        # Save evaluation to database with additional info (use PDF path for resume_path)
        db_id = save_evaluation(eval_id, filename, job_title, match_percentage, missing_keywords, profile_summary, match_factors, stability_data, additional_info, None, candidate_fit_analysis, over_under_qualification, user_email, resume_path=pdf_path)
        if db_id:
            # Generate interview questions asynchronously
            questions_data = await async_generate_questions(resume_text, job_description, profile_summary)
            
            technical_questions = questions_data.get("TechnicalQuestions", [])
            nontechnical_questions = questions_data.get("NonTechnicalQuestions", [])
            behavioral_questions = QUICK_CHECKS

            # Save interview questions with proper JSON encoding (use database ID)
            if save_interview_questions(
                db_id,
                json.dumps(technical_questions),
                json.dumps(nontechnical_questions),
                json.dumps(behavioral_questions),
            ):
                # IMPORTANT: return the actual DB primary key (db_id), not the UUID eval_id.
                # The evaluation view, /api/evaluation-full/<id>, and download endpoints
                # all expect the SQLite autoincrement id.
                return jsonify({
                    'id': db_id,
                    'match_percentage': match_percentage,
                    'match_percentage_str': match_percentage_str,
                    'missing_keywords': missing_keywords,
                    'profile_summary': profile_summary,
                    'over_under_qualification': over_under_qualification,
                    'match_factors': match_factors,
                    'candidate_fit_analysis': candidate_fit_analysis,
                    'job_stability': stability_data,
                    'career_progression': career_data,
                    'technical_questions': technical_questions,
                    'nontechnical_questions': nontechnical_questions,
                    'behavioral_questions': behavioral_questions
                })
            else:
                return jsonify({'error': 'Failed to save interview questions'}), 500
        else:
            return jsonify({'error': 'Failed to save evaluation'}), 500

    except Exception as e:
        logging.error(f"Error in evaluate_resume: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/evaluate-stream', methods=['POST'])
@login_required
def evaluate_resume_stream():
    """Streaming version of resume evaluation for better UX"""
    try:
        if 'resume' not in request.files:
            return jsonify({'error': 'No resume file provided'}), 400
        
        file = request.files['resume']
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type'}), 400

        job_title = request.form.get('job_title')
        job_description = request.form.get('job_description')
        oorwin_job_id = request.form.get('oorwin_job_id', '').strip()  # Get JobID from form
        additional_context = request.form.get('additional_context', '').strip()
        additional_context_block = f"**Additional Context (client constraints/preference):** {additional_context}" if additional_context else ""

        if not job_title or not job_description:
            return jsonify({'error': 'Missing job title or description'}), 400

        # Save uploaded file
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        file_path = os.path.normpath(file_path)  # Normalize path

        # Convert DOCX to PDF if needed, and get the PDF path for merging
        pdf_path = None
        file_ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
        logging.info(f"[UPLOAD] File extension: {file_ext}, Original file path: {file_path}")
        
        if file_ext == 'docx':
            # Convert DOCX to PDF
            logging.info(f"[CONVERSION] Attempting to convert DOCX to PDF: {file_path}")
            pdf_path = convert_docx_to_pdf(file_path)
            if pdf_path:
                pdf_path = os.path.normpath(pdf_path)  # Normalize PDF path
                if os.path.exists(pdf_path):
                    logging.info(f"[CONVERSION] ✅ Successfully converted DOCX to PDF: {pdf_path}")
                else:
                    logging.error(f"[CONVERSION] ❌ PDF file not found after conversion: {pdf_path}")
                    pdf_path = None
            else:
                logging.error(f"[CONVERSION] ❌ Conversion function returned None")
                pdf_path = None
            
            # If conversion failed, we can't merge, but we'll still save the evaluation
            if not pdf_path:
                logging.warning(f"[CONVERSION] ⚠️ Will save DOCX path (merging won't work): {file_path}")
                pdf_path = file_path  # Fallback to original
        elif file_ext == 'pdf':
            pdf_path = file_path  # Already a PDF
            logging.info(f"[UPLOAD] File is already a PDF: {pdf_path}")
        else:
            pdf_path = file_path  # Fallback
            logging.warning(f"[UPLOAD] Unknown file extension, using original path: {pdf_path}")
        
        logging.info(f"[SAVE] Final resume_path to save in database: {pdf_path}")
        logging.info(f"[DEBUG] PDF PATH: {pdf_path}")
        logging.info(f"[DEBUG] Exists at save time: {os.path.exists(pdf_path) if pdf_path else False}")

        # Extract text from resume (use original file path for text extraction)
        resume_text = extract_text_from_file(file_path)
        if resume_text is None:
            return jsonify({'error': 'Failed to extract text from file'}), 500

        # Generate unique ID for evaluation
        eval_id = str(uuid.uuid4())
        
        # Get user email from session
        user_email = session.get('user', {}).get('email') if 'user' in session else None
        
        # Track start time for performance metrics
        import time
        start_time = time.time()
        
        # Capture filename, pdf_path, and file_path for use in generator function
        captured_filename = filename
        captured_pdf_path = pdf_path
        captured_file_path = file_path

        def generate():
            try:
                # Send initial response
                yield f"data: {json.dumps({'status': 'processing', 'message': 'Analyzing resume...', 'eval_id': eval_id})}\n\n"
                
                # SINGLE API CALL: Unified evaluation (all analyses in one call)
                yield f"data: {json.dumps({'status': 'step1', 'message': 'Evaluating resume against job requirements...'})}\n\n"
                formatted_prompt = unified_evaluation_prompt.format(
                    resume_text=resume_text,
                    job_description=job_description,
                    additional_context_block=additional_context_block
                )
                
                # Use Groq with temperature 0 for consistent scoring
                total_tokens_used = 0
                try:
                    # Runtime check: if client is not initialized, try to initialize it
                    global groq_client
                    if not groq_client:
                        # Reload environment variables
                        BASE_DIR = os.path.dirname(os.path.abspath(__file__))
                        ENV_PATH = os.path.join(BASE_DIR, '.env')
                        load_dotenv(ENV_PATH, override=True)
                        
                        # Get API key again
                        api_key = os.getenv("GROQ_API_KEY", "").strip()
                        if not api_key:
                            raise RuntimeError("Groq client not initialized. Please set GROQ_API_KEY in .env file.")
                        
                        # Try to initialize
                        try:
                            groq_client = Groq(api_key=api_key)
                            logging.info("[RUNTIME] Groq client initialized successfully in stream function")
                        except Exception as e:
                            raise RuntimeError(f"Failed to initialize Groq client: {str(e)}. Please check your GROQ_API_KEY in .env file.")
                    
                    # SINGLE API CALL - Get all analyses at once
                    logging.info("[UNIFIED] Making single API call for complete evaluation...")
                    groq_response = asyncio.run(async_groq_generate_for_scoring(formatted_prompt))
                    
                    # Extract token usage
                    if hasattr(groq_response, 'token_usage') and groq_response.token_usage:
                        total_tokens_used = groq_response.token_usage
                        logging.info(f"[UNIFIED] Token usage for complete evaluation: {total_tokens_used}")
                    
                    # Parse the response using the same logic as async_gemini_generate
                    response_text = groq_response.text if hasattr(groq_response, 'text') else str(groq_response)
                    if not response_text:
                        raise ValueError("Empty response from Groq")
                    
                    response_text = response_text.strip()
                    # Remove markdown code block markers if present
                    response_text = re.sub(r'^```json\s*', '', response_text, flags=re.MULTILINE)
                    response_text = re.sub(r'^```\s*', '', response_text, flags=re.MULTILINE)
                    response_text = re.sub(r'\s*```$', '', response_text)
                    response_text = response_text.strip()
                    
                    if response_text.startswith("```") and response_text.endswith("```"):
                        response_text = response_text[3:-3].strip()
                    
                    first_char = response_text.strip()[0] if response_text.strip() else ''
                    if first_char != '{':
                        json_start = response_text.find('{')
                        if json_start > 0:
                            response_text = response_text[json_start:]
                    
                    if not response_text.startswith('{'):
                        raise ValueError("Response doesn't contain valid JSON")
                    
                    unified_response = json.loads(response_text)
                    if not isinstance(unified_response, dict):
                        raise ValueError("Parsed JSON is not a dict")
                    
                    logging.info(f"[UNIFIED] ✅ Received unified response with keys: {list(unified_response.keys())}")
                    
                except Exception as gen_error:
                    error_msg = str(gen_error)
                    logging.error(f"❌ Error calling Groq API: {type(gen_error).__name__}: {error_msg}")
                    logging.error(f"   Full traceback:", exc_info=True)
                    yield f"data: {json.dumps({'status': 'error', 'message': f'AI generation failed: {error_msg}'})}\n\n"
                    return
                
                # Verify unified response
                if not unified_response:
                    logging.error("❌ unified_response is empty/None")
                    yield f"data: {json.dumps({'status': 'error', 'message': 'Failed to analyze resume - empty response'})}\n\n"
                    return
                
                # Verify it's a dictionary
                if not isinstance(unified_response, dict):
                    logging.error(f"❌ unified_response is not a dict!")
                    logging.error(f"   Type: {type(unified_response)}")
                    logging.error(f"   Value (repr): {repr(unified_response)[:500]}")
                    yield f"data: {json.dumps({'status': 'error', 'message': f'Invalid response format from AI. Expected dict, got {type(unified_response).__name__}'})}\n\n"
                    return
                
                # Verify it has required keys
                required_keys = ["JD Match", "Match Factors", "Profile Summary"]
                missing_keys = [key for key in required_keys if key not in unified_response]
                if missing_keys:
                    logging.error(f"❌ Missing required keys in response: {missing_keys}")
                    logging.error(f"   Available keys: {list(unified_response.keys())}")
                    yield f"data: {json.dumps({'status': 'error', 'message': f'AI response missing required fields: {missing_keys}'})}\n\n"
                    return
                
                # Extract all sections from unified response
                try:
                    # Main evaluation data
                    match_percentage_str = unified_response.get("JD Match", "0%")
                    if match_percentage_str is None:
                        match_percentage_str = "0%"
                    if not isinstance(match_percentage_str, str):
                        match_percentage_str = str(match_percentage_str) if match_percentage_str else "0%"
                    match_percentage_str = match_percentage_str.strip()
                    if not match_percentage_str:
                        match_percentage_str = "0%"
                    match_percentage = int(match_percentage_str.strip('%'))
                except (ValueError, AttributeError, TypeError) as e:
                    logging.error(f"❌ Error parsing match percentage: {e}")
                    match_percentage = 0
                    match_percentage_str = "0%"
                
                # Extract all sections
                missing_keywords = unified_response.get("MissingKeywords", [])
                profile_summary = unified_response.get("Profile Summary", "No summary provided.")
                over_under_qualification = unified_response.get("Over/UnderQualification Analysis", "No qualification mismatch concerns detected.")
                match_factors = unified_response.get("Match Factors", {})
                candidate_fit_analysis = unified_response.get("Candidate Fit Analysis", {})
                
                # Extract job stability (with defaults if missing)
                stability_data = unified_response.get("Job Stability", {})
                if not stability_data or not isinstance(stability_data, dict):
                    stability_data = {
                        "IsStable": True,
                        "AverageJobTenure": "Unknown",
                        "JobCount": 0,
                        "StabilityScore": 50,
                        "ReasoningExplanation": "Stability analysis not available",
                        "RiskLevel": "Medium"
                    }
                
                # Extract career progression (with defaults if missing)
                career_data = unified_response.get("Career Progression", {})
                if not career_data or not isinstance(career_data, dict):
                    career_data = {
                        "progression_score": 50,
                        "key_observations": ["Career progression analysis not available"],
                        "career_path": [],
                        "red_flags": [],
                        "reasoning": "Analysis not available"
                    }
                
                # Extract interview questions (with defaults if missing)
                questions_data = unified_response.get("Interview Questions", {})
                if not questions_data or not isinstance(questions_data, dict):
                    questions_data = {
                        "TechnicalQuestions": [],
                        "NonTechnicalQuestions": []
                    }
                
                technical_questions = questions_data.get("TechnicalQuestions", [])
                nontechnical_questions = questions_data.get("NonTechnicalQuestions", [])
                behavioral_questions = QUICK_CHECKS
                
                # Derive resume filename to send to frontend (used for merged PDF download)
                # IMPORTANT: Always prefer PDF filename for merging - if DOCX was converted, use PDF filename
                # Use captured variables from outer scope
                resume_filename = None
                if captured_pdf_path and os.path.exists(captured_pdf_path):
                    # Prefer PDF path (even if original was DOCX, we converted it)
                    resume_filename = os.path.basename(captured_pdf_path)
                    logging.info(f"[RESUME_FILENAME] Using PDF filename: {resume_filename}")
                elif captured_file_path and os.path.exists(captured_file_path):
                    resume_filename = os.path.basename(captured_file_path)
                    logging.info(f"[RESUME_FILENAME] Using original filename: {resume_filename}")
                
                # Save evaluation EARLY (right after basic results) so we can send db_id immediately
                # This allows frontend to have the evaluation ID right away
                logging.info(f"[EARLY_SAVE] Saving evaluation early with basic results...")
                db_id = save_evaluation(
                    eval_id, 
                    captured_filename, 
                    job_title, 
                    match_percentage, 
                    missing_keywords, 
                    profile_summary, 
                    match_factors, 
                    {},  # job_stability - will update later
                    {},  # additional_info - will update later
                    oorwin_job_id, 
                    candidate_fit_analysis, 
                    over_under_qualification, 
                    user_email, 
                    None,  # time_taken - will update later
                    None,  # token_usage - will update later
                    resume_path=captured_pdf_path  # Always use PDF path for resume_path
                )
                logging.info(f"[EARLY_SAVE] Early save result: db_id={db_id}")
                
                # Determine original file extension (based on uploaded filename)
                # We will use this to control whether merging is allowed (PDF only for now)
                original_ext = os.path.splitext(captured_filename or '')[1].lower().lstrip('.')
                can_merge = (original_ext == 'pdf')
                
                # Send basic results immediately (user sees this in ~5-8 seconds instead of 20)
                # NOW includes db_id so frontend can use it immediately, along with merge capability flag
                basic_results = {
                    'status': 'basic_results',
                    'id': eval_id,  # Keep UUID for backward compatibility
                    'db_id': db_id,  # NEW: Send DB ID immediately (SQLite primary key)
                    'match_percentage': match_percentage,
                    'match_percentage_str': match_percentage_str,
                    'missing_keywords': missing_keywords if isinstance(missing_keywords, list) else [],
                    'profile_summary': profile_summary,
                    'over_under_qualification': over_under_qualification,
                    'match_factors': match_factors if isinstance(match_factors, dict) else {},
                    'candidate_fit_analysis': candidate_fit_analysis if isinstance(candidate_fit_analysis, dict) else {},
                    'filename': resume_filename,
                    'resume_path': resume_filename,  # Filename used for merging/downloads
                    'original_file_ext': original_ext,  # e.g., 'pdf' or 'docx'
                    'can_merge': can_merge            # Frontend uses this to enable/disable merge button
                }
                
                # Log the payload before sending
                logging.info(f"[STREAM PAYLOAD DEBUG] {basic_results}")
                
                # Test JSON serialization before yielding
                try:
                    json_str = json.dumps(basic_results)
                    yield f"data: {json_str}\n\n"
                except (TypeError, ValueError) as json_err:
                    logging.error(f"JSON serialization error: {json_err}")
                    logging.error(f"Problematic data types: match_factors={type(match_factors)}, candidate_fit_analysis={type(candidate_fit_analysis)}")
                    # Send error with details
                    yield f"data: {json.dumps({'status': 'error', 'message': f'Data serialization error: {str(json_err)}'})}\n\n"
                    return
                
                # Simulate progress updates (all data already extracted from single call)
                yield f"data: {json.dumps({'status': 'step2', 'message': 'Analyzing job stability and career progression...'})}\n\n"
                
                # Send stability and career data
                additional_data = {
                    'status': 'additional_data',
                    'job_stability': stability_data,
                    'career_progression': career_data
                }
                yield f"data: {json.dumps(additional_data)}\n\n"
                
                # Simulate progress for questions
                yield f"data: {json.dumps({'status': 'step3', 'message': 'Generating interview questions...'})}\n\n"
                
                # Send questions
                questions_data_response = {
                    'status': 'questions',
                    'technical_questions': technical_questions,
                    'nontechnical_questions': nontechnical_questions,
                    'behavioral_questions': behavioral_questions
                }
                yield f"data: {json.dumps(questions_data_response)}\n\n"
                
                # Step 4: Update evaluation with additional data (stability, career, questions)
                # Note: Evaluation was already saved early after basic_results, so we just update it
                yield f"data: {json.dumps({'status': 'step4', 'message': 'Finalizing results...'})}\n\n"
                
                additional_info = {
                    "job_stability": stability_data,
                    "career_progression": career_data,
                    "reasoning": unified_response.get("Reasoning", "")
                }
                
                # Calculate time taken
                time_taken = round(time.time() - start_time, 2)  # Round to 2 decimal places
                
                # Update the evaluation that was saved early with additional data
                if db_id:
                    logging.info(f"[UPDATE] Updating evaluation {db_id} with additional data...")
                    conn_update = sqlite3.connect(DATABASE_NAME)
                    cursor_update = conn_update.cursor()
                    
                    # Update job_stability, career_progression, time_taken, token_usage
                    cursor_update.execute("""
                        UPDATE evaluations 
                        SET job_stability = ?,
                            career_progression = ?,
                            time_taken = ?,
                            token_usage = ?
                        WHERE id = ?
                    """, (
                        json.dumps(stability_data) if stability_data else '{}',
                        json.dumps(career_data) if career_data else '{}',
                        time_taken,
                        total_tokens_used,
                        db_id
                    ))
                    conn_update.commit()
                    conn_update.close()
                    logging.info(f"[UPDATE] ✅ Updated evaluation {db_id} with additional data")
                else:
                    logging.warning(f"[UPDATE] ⚠️ No db_id available to update evaluation")
                
                # Verify DB save
                if db_id:
                    conn_verify = sqlite3.connect(DATABASE_NAME)
                    cursor_verify = conn_verify.cursor()
                    cursor_verify.execute("SELECT resume_path FROM evaluations WHERE id = ?", (db_id,))
                    saved_resume_path = cursor_verify.fetchone()
                    conn_verify.close()
                    logging.info(f"[VERIFY] Saved resume_path in DB for id {db_id}: {saved_resume_path}")
                
                if db_id:
                    # Use the database ID (integer) for saving interview questions
                    if save_interview_questions(db_id, 
                                             json.dumps(technical_questions), 
                                             json.dumps(nontechnical_questions), 
                                             json.dumps(behavioral_questions)):
                        # Send the database ID back to frontend for feedback submission
                        yield f"data: {json.dumps({'status': 'complete', 'message': 'Analysis complete!', 'db_id': db_id})}\n\n"
                    else:
                        yield f"data: {json.dumps({'status': 'error', 'message': 'Failed to save interview questions'})}\n\n"
                else:
                    yield f"data: {json.dumps({'status': 'error', 'message': 'Failed to save evaluation'})}\n\n"
                    
            except Exception as e:
                error_type = type(e).__name__
                error_msg = str(e)
                logging.error(f"❌ Error in streaming evaluation: {error_type}: {error_msg}")
                logging.error(f"   Error repr: {repr(error_msg)}")
                logging.error(f"   Error type: {error_type}")
                logging.error(f"   Full traceback:", exc_info=True)
                # Make sure error message is JSON-safe
                safe_error_msg = error_msg[:500] if len(error_msg) > 500 else error_msg
                yield f"data: {json.dumps({'status': 'error', 'message': safe_error_msg, 'error_type': error_type})}\n\n"

        return Response(stream_with_context(generate()), mimetype='text/event-stream')

    except Exception as e:
        logging.error(f"Error in evaluate_resume_stream: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/get_interview_questions/<evaluation_id>', methods=['GET'])
def get_interview_questions(evaluation_id):
    """Get interview questions for a specific evaluation"""
    conn = None
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # First, get the evaluation details to regenerate questions if needed
        cursor.execute(
            """
            SELECT e.resume_path, e.job_title, e.job_description, e.profile_summary 
            FROM evaluations e 
            WHERE e.id = ?
            """,
            (evaluation_id,)
        )
        eval_result = cursor.fetchone()
        
        # Then get existing questions
        cursor.execute(
            "SELECT technical_questions, nontechnical_questions, behavioral_questions FROM interview_questions WHERE evaluation_id = ?",
            (evaluation_id,)
        )
        result = cursor.fetchone()
        
        # Initialize default values
        technical_questions = []
        nontechnical_questions = []
        behavioral_questions = []
        
        if result:
            try:
                # Parse saved questions with proper error handling
                def parse_json_safely(json_str):
                    if not json_str:
                        return []
                    try:
                        data = json.loads(json_str)
                        if isinstance(data, list):
                            return data
                        elif isinstance(data, str):
                            try:
                                return json.loads(data)
                            except:
                                return [data]
                        else:
                            return [str(data)]
                    except json.JSONDecodeError:
                        try:
                            # Try to clean and parse the string
                            cleaned_str = json_str.strip('[]"\' ').replace('\\', '')
                            items = [item.strip('"\' ') for item in cleaned_str.split(',')]
                            return [item for item in items if item]
                        except:
                            return []

                technical_questions = parse_json_safely(result[0])
                nontechnical_questions = parse_json_safely(result[1])
                behavioral_questions = parse_json_safely(result[2])
            except Exception as e:
                logging.error(f"Error parsing interview questions: {str(e)}")
                technical_questions = []
                nontechnical_questions = []
                behavioral_questions = []

        # Only regenerate questions if they are completely missing (not just empty)
        # This prevents regenerating questions when they exist but are empty arrays
        if not result and eval_result:
            logging.info(f"No interview questions found in database for evaluation {evaluation_id}, generating new ones")
            resume_path = eval_result[0]
            if not resume_path or resume_path == 'NULL' or resume_path == 'None':
                logging.error(f"Invalid resume_path for evaluation {evaluation_id}: {resume_path}")
                return jsonify({'error': 'Resume file path is invalid'}), 400
            resume_text = extract_text_from_file(resume_path)
            if resume_text:
                        questions_data = asyncio.run(async_generate_questions(
                            resume_text,
                            eval_result[2],  # job_description
                            eval_result[3]   # profile_summary
                        ))
                        
            technical_questions = questions_data.get("TechnicalQuestions", [])
            nontechnical_questions = questions_data.get("NonTechnicalQuestions", [])
            behavioral_questions = QUICK_CHECKS
                        
             # Save regenerated questions
            cursor.execute(
                            """
                    INSERT INTO interview_questions 
                    (evaluation_id, technical_questions, nontechnical_questions, behavioral_questions) 
                    VALUES (?, ?, ?, ?)
                    """,
                    (evaluation_id,
                     json.dumps(technical_questions), 
                             json.dumps(nontechnical_questions), 
                     json.dumps(behavioral_questions))
                        )
            conn.commit()
            logging.info(f"Generated and saved new questions for evaluation {evaluation_id}")

            return jsonify({
                    'technical_questions': technical_questions or ["No technical questions available"],
                    'nontechnical_questions': nontechnical_questions or ["No non-technical questions available"],
                    'behavioral_questions': behavioral_questions or QUICK_CHECKS
                })

    except Exception as e:
        logging.error(f"Database error in get_interview_questions: {str(e)}")
        return jsonify({
            'technical_questions': ["Error loading technical questions"],
            'nontechnical_questions': ["Error loading non-technical questions"],
            'behavioral_questions': QUICK_CHECKS
        })
    finally:
        if conn:
            conn.close()

@app.route('/api/feedback', methods=['POST'])
def submit_feedback():
    """Handle feedback for both Q&A and resume evaluations."""
    try:
        data = request.get_json()
        logging.info(f"Received feedback data: {data}")
        
        if not data:
            logging.error("No feedback data received")
            return jsonify({'error': 'No feedback data provided'}), 400
            
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        try:
            # Check if this is Q&A feedback
            if 'question' in data:
                if 'rating' not in data:
                    return jsonify({'error': 'Missing rating'}), 400
                
                # Get question_id from qa_history
                cursor.execute("""
                    SELECT id FROM qa_history 
                    WHERE question = ? 
                    ORDER BY timestamp DESC 
                    LIMIT 1
                """, (data['question'],))
                
                result = cursor.fetchone()
                if not result:
                    # If question not found, create a new entry
                    cursor.execute("""
                        INSERT INTO qa_history (question, final_answer)
                        VALUES (?, ?)
                    """, (data['question'], ''))
                    question_id = cursor.lastrowid
                else:
                    question_id = result[0]
                
                # Check if feedback already exists for this question
                cursor.execute("SELECT id FROM qa_feedback WHERE question_id = ?", (question_id,))
                if cursor.fetchone():
                    return jsonify({'error': 'Feedback already submitted for this question'}), 400
                
                # Insert feedback
                cursor.execute("""
                    INSERT INTO qa_feedback (question_id, rating, feedback, timestamp)
                    VALUES (?, ?, ?, datetime('now'))
                """, (question_id, data['rating'], data.get('feedback', '')))
                
            else:
                # Handle resume evaluation feedback
                if 'evaluation_id' not in data or 'rating' not in data:
                    return jsonify({'error': 'Missing evaluation_id or rating'}), 400
                
                # Check if feedback already exists
                cursor.execute("SELECT id FROM feedback WHERE evaluation_id = ?", (data['evaluation_id'],))
                if cursor.fetchone():
                    return jsonify({'error': 'Feedback already submitted for this evaluation'}), 400
                
                # Insert feedback into the feedback table
                cursor.execute("""
                    INSERT INTO feedback (evaluation_id, rating, comments, timestamp)
                    VALUES (?, ?, ?, datetime('now'))
                """, (data['evaluation_id'], data['rating'], data.get('comments', '')))
            
            conn.commit()
            return jsonify({'message': 'Feedback submitted successfully'})
            
        finally:
            conn.close()
            
    except sqlite3.IntegrityError as e:
        logging.error(f"Integrity error in submit_feedback: {str(e)}")
        return jsonify({'error': 'Feedback already submitted'}), 400
    except sqlite3.Error as e:
        logging.error(f"Database error in submit_feedback: {str(e)}")
        return jsonify({'error': 'Database error occurred'}), 500
    except Exception as e:
        logging.error(f"Error in submit_feedback: {str(e)}")
        return jsonify({'error': 'An unexpected error occurred'}), 500

# --- Document Processing ---
# Optimized chunking for HR policy documents:
# - Larger chunks (1200) preserve context and complete policy explanations
# - Higher overlap (250) ensures continuity across chunks
# - Better separators prioritize paragraph/sentence boundaries
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1200,        # Increased from 400 for better context retention
    chunk_overlap=250,      # Increased from 50 (20% overlap for continuity)
    separators=["\n\n", "\n", ". ", " ", ""]  # Better paragraph/sentence awareness
)

def dataframe_to_clean_markdown(df: pd.DataFrame) -> str:
    """Convert a pandas DataFrame into a clean GitHub-flavored markdown table.

    - Fills missing/None headers with generic names (Column 1, Column 2, ...)
    - Collapses multi-line cell content into single lines
    - Ensures string-typed cells
    - Removes the index column from the markdown output
    """
    try:
        # Defensive copy
        table_df = df.copy()

        # Normalize headers
        normalized_columns = []
        for i, col in enumerate(table_df.columns):
            col_str = str(col).strip() if col is not None else ""
            if col_str == "" or col_str.lower() == "none":
                col_str = f"Column {i+1}"
            normalized_columns.append(col_str)
        table_df.columns = normalized_columns

        # Drop index-like first column if it looks like 0..N sequence
        if table_df.shape[1] > 0:
            first_col_vals = list(table_df.iloc[:, 0])
            def looks_like_sequential_index(vals):
                try:
                    ints = [int(str(v)) for v in vals]
                    return ints == list(range(len(ints)))
                except Exception:
                    return False
            if looks_like_sequential_index(first_col_vals):
                table_df = table_df.iloc[:, 1:]

        # Drop columns that are entirely empty after stripping
        if table_df.shape[1] > 0:
            non_empty_cols = []
            for c in table_df.columns:
                try:
                    # Safely convert column to string, handling DataFrames and complex objects
                    col_series = table_df[c]
                    # Check if column contains DataFrames or other complex objects
                    if col_series.dtype == 'object':
                        # Convert each cell to string individually to handle complex objects
                        col_str = col_series.apply(lambda x: str(x) if x is not None else '').str.strip()
                    else:
                        col_str = col_series.astype(str).str.strip()
                    if (col_str != "").any():
                        non_empty_cols.append(c)
                except (AttributeError, TypeError) as e:
                    # If .str accessor fails, try converting each cell individually
                    try:
                        col_str = col_series.apply(lambda x: str(x).strip() if x is not None and str(x).strip() else '')
                        if (col_str != "").any():
                            non_empty_cols.append(c)
                    except Exception:
                        # Skip this column if we can't process it
                        logging.warning(f"Skipping column '{c}' due to processing error")
                        continue
            if non_empty_cols:
                table_df = table_df[non_empty_cols]

        # Normalize cell content: string type, collapse newlines/tabs, trim
        for c in table_df.columns:
            try:
                # Safely convert column to string, handling DataFrames and complex objects
                col_series = table_df[c]
                if col_series.dtype == 'object':
                    # For object columns, convert each cell individually to handle complex objects
                    table_df[c] = col_series.apply(
                        lambda x: ' '.join(str(x).split()) if x is not None else ''
                    )
                else:
                    # For non-object columns, use standard string operations
                    table_df[c] = (
                        col_series
                        .astype(str)
                        .str.replace("\r\n|\r|\n", " ", regex=True)
                        .str.replace("\t", " ", regex=True)
                        .str.replace("\\s+", " ", regex=True)
                        .str.strip()
                    )
            except (AttributeError, TypeError) as e:
                # If .str accessor fails, convert each cell individually
                try:
                    table_df[c] = table_df[c].apply(
                        lambda x: ' '.join(str(x).split()) if x is not None else ''
                    )
                except Exception as inner_e:
                    # Last resort: just convert to string
                    table_df[c] = table_df[c].apply(lambda x: str(x) if x is not None else '')

        # Render as markdown without index
        return table_df.to_markdown(index=False, tablefmt="pipe")
    except Exception as e:
        logging.warning(f"Failed to render markdown table cleanly, falling back: {e}")
        # Fallback to basic to_markdown if anything goes wrong
        try:
            return df.to_markdown(index=False, tablefmt="pipe")
        except Exception:
            return df.to_markdown()

def process_pdf(pdf_path, documents, table_chunks):
    """Extract text and tables from a PDF file."""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                # Extract and process text
                text = page.extract_text() or ""
                if text:
                    text_chunks = text_splitter.split_text(text)
                    documents.extend(text_chunks)
                
                # Extract and process tables
                tables = page.extract_tables()
                for table in tables:
                    if table and len(table) > 1:  # Ensure table has headers and data
                        df = pd.DataFrame(table[1:], columns=table[0])
                        table_markdown = dataframe_to_clean_markdown(df)
                        
                        # Enrich table for better BM25 matching (same as Pinecone indexing)
                        column_names = " ".join([str(col).lower() for col in df.columns if col])
                        sample_values = []
                        # Use positional access to avoid issues with non-standard/duplicate column labels
                        num_sample_cols = min(3, len(df.columns))
                        for col_idx in range(num_sample_cols):
                            series = df.iloc[:, col_idx]
                            sample_values.extend([str(val).lower() for val in series.dropna().head(3).tolist()])
                        sample_context = " ".join(sample_values[:10])
                        enriched_table = f"[TABLE DATA] Topic: {column_names} {sample_context}\n\n{table_markdown}\n\n[END TABLE]"
                        table_chunks.append(enriched_table)
    except Exception as e:
        logging.error(f"❌ Error processing PDF {pdf_path}: {e}")

def populate_pinecone_index():
    """Extract content from PDF documents and populate Pinecone index."""
    try:
        documents = []
        table_chunks = []
        texts = []
        metadatas = []
        
        # Get all PDF files from the policies folder
        if not os.path.exists(POLICIES_FOLDER):
            logging.warning(f"Policies folder {POLICIES_FOLDER} does not exist")
            return
            
        pdf_files = [f for f in os.listdir(POLICIES_FOLDER) if f.endswith('.pdf')]
        if not pdf_files:
            logging.warning(f"No PDF files found in {POLICIES_FOLDER}")
            return
            
        total_files = len(pdf_files)
        logging.info(f"📚 Processing {total_files} PDF files")
        
        for idx, filename in enumerate(pdf_files, 1):
            pdf_path = os.path.join(POLICIES_FOLDER, filename)
            logging.info(f"📄 Processing file {idx}/{total_files}: {filename}")
            
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text() or ""
                    if text:
                        chunks = text_splitter.split_text(text)
                        documents.extend(chunks)
                        for chunk in chunks:
                            texts.append(chunk)
                            metadatas.append({
                                "source": filename,
                                "page": page_num,
                                "type": "text"
                            })
                        logging.info(f"   Page {page_num}: Added {len(chunks)} text chunks")
                    
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables, 1):
                        if table and len(table) > 1:
                            df = pd.DataFrame(table[1:], columns=table[0])
                            table_markdown = dataframe_to_clean_markdown(df)
                            
                            # Improve table representation for better retrieval:
                            # 1. Extract column names as descriptive keywords
                            column_names = " ".join([str(col).lower() for col in df.columns if col])
                            
                            # 2. Extract sample data values as context
                            sample_values = []
                            # Use positional access to avoid issues with non-standard/duplicate column labels
                            num_sample_cols = min(3, len(df.columns))
                            for col_idx in range(num_sample_cols):  # First up to 3 columns
                                series = df.iloc[:, col_idx]
                                sample_values.extend([str(val).lower() for val in series.dropna().head(3).tolist()])
                            sample_context = " ".join(sample_values[:10])
                            
                            # 3. Create enriched table chunk with context
                            enriched_table = f"[TABLE DATA] Topic: {column_names} {sample_context}\n\n{table_markdown}\n\n[END TABLE]"
                            table_chunks.append(enriched_table)
                            texts.append(enriched_table)
                            metadatas.append({
                                "source": filename,
                                "page": page_num,
                                "type": "table"
                            })
                            logging.info(f"   Page {page_num}: Added table {table_num} (enriched with context)")
    except Exception as e:
        logging.error(f"❌ Error in document processing: {str(e)}")
        raise

    try:
        all_chunks = documents + table_chunks
        total_chunks = len(texts)
        
        if total_chunks == 0:
            raise ValueError("No content extracted from documents")
        
        logging.info(f"📊 Preparing to insert {total_chunks} chunks into Pinecone")
        
        # Initialize Pinecone components
        pc = Pinecone(api_key=PINECONE_API_KEY)
        index = pc.Index(PINECONE_INDEX_NAME)
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        
        # Insert in batches
        batch_size = 50  # Reduced batch size for better reliability
        for i in range(0, total_chunks, batch_size):
            batch_texts = texts[i:i + batch_size]
            batch_metas = metadatas[i:i + batch_size]
            PineconeVectorStore.from_texts(
                texts=batch_texts,
                embedding=embeddings,
                index_name=PINECONE_INDEX_NAME,
                metadatas=batch_metas
            )
            logging.info(f"✅ Inserted batch {i//batch_size + 1}/{(total_chunks-1)//batch_size + 1}")
        
        # Verify insertion
        stats = index.describe_index_stats()
        vector_count = stats['total_vector_count']
        logging.info(f"🎉 Successfully populated index with {vector_count} vectors")
        
    except Exception as e:
        logging.error(f"❌ Error in Pinecone operations: {str(e)}")
        raise
#         total_files = len(pdf_files)
#         logging.info(f"📚 Processing {total_files} PDF files")
        
#         for idx, filename in enumerate(pdf_files, 1):
#             pdf_path = os.path.join(POLICIES_FOLDER, filename)
#             logging.info(f"📄 Processing file {idx}/{total_files}: {filename}")
            
#             with pdfplumber.open(pdf_path) as pdf:
#                 for page_num, page in enumerate(pdf.pages, 1):
#                     text = page.extract_text() or ""
#                     if text:
#                         chunks = text_splitter.split_text(text)
#                         documents.extend(chunks)
#                         logging.info(f"   Page {page_num}: Added {len(chunks)} text chunks")
#                     
#                     tables = page.extract_tables()
#                     for table_num, table in enumerate(tables, 1):
#                         if table and len(table) > 1:
#                             df = pd.DataFrame(table[1:], columns=table[0])
#                             table_chunks.append(df.to_markdown())
#                             logging.info(f"   Page {page_num}: Added table {table_num}")
#     except Exception as e:
#         logging.error(f"❌ Error in document processing: {str(e)}")
#         raise

#     try:
#     all_chunks = documents + table_chunks
#     total_chunks = len(all_chunks)
        
#     if total_chunks == 0:
#             raise ValueError("No content extracted from documents")
        
#     logging.info(f"📊 Preparing to insert {total_chunks} chunks into Pinecone")
        
#         # Initialize Pinecone components
#     pc = Pinecone(api_key=PINECONE_API_KEY)
#     index = pc.Index(PINECONE_INDEX_NAME)
#     embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        
#     # Insert in batches
#     batch_size = 50  # Reduced batch size for better reliability
#     for i in range(0, total_chunks, batch_size):
#             batch = all_chunks[i:i + batch_size]
#             PineconeVectorStore.from_texts(
#                 texts=batch,
#                 embedding=embeddings,
#                 index_name=PINECONE_INDEX_NAME
#             )
#             logging.info(f"✅ Inserted batch {i//batch_size + 1}/{(total_chunks-1)//batch_size + 1}")
        
#         # Verify insertion
#             stats = index.describe_index_stats()
#             vector_count = stats['total_vector_count']
#             logging.info(f"🎉 Successfully populated index with {vector_count} vectors")
        
#     except Exception as e:
#         logging.error(f"❌ Error in Pinecone operations: {str(e)}")
#         raise

def initialize_pinecone():
    """Initialize Pinecone. Create and populate index if it doesn't exist."""
    try:
        logging.info("🔧 Initializing Pinecone...")
        
        # Check if index exists
        pc = Pinecone(api_key=PINECONE_API_KEY)
        index_name = PINECONE_INDEX_NAME
        
        if index_name in pc.list_indexes().names():
            logging.info(f"📋 Index '{index_name}' already exists - using existing index")
            # Check if index has data
            index = pc.Index(index_name)
            stats = index.describe_index_stats()
            vector_count = stats['total_vector_count']
            
            if vector_count > 0:
                logging.info(f"✅ Index '{index_name}' has {vector_count} vectors - no need to populate")
                return True
            else:
                logging.info(f"⚠️ Index '{index_name}' exists but is empty - populating...")
                populate_pinecone_index()
                return True
        else:
            # Create new index only if it doesn't exist
            logging.info(f"🆕 Index '{index_name}' doesn't exist - creating new index...")
            pc.create_index(
                name=index_name,
                dimension=384,
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1")
            )
            
            # Wait for index to be ready
            time.sleep(10)
            logging.info(f"✅ Index '{index_name}' created successfully")
            
            # Populate the new index
            logging.info("📚 Populating new Pinecone index...")
            populate_pinecone_index()
            
            return True
        
    except Exception as e:
        logging.error(f"❌ Error initializing Pinecone: {str(e)}")
        return False

# --- BM25 Setup ---
bm25_index = None
bm25_corpus = None
bm25_metadata = []  # Store metadata for each BM25 chunk (filename, page, type)

def build_bm25_index(folder_path):
    """Builds BM25 index from policy documents with metadata tracking."""
    global bm25_index, bm25_corpus, bm25_metadata
    
    all_texts = []
    table_chunks = []
    text_metadata = []  # Track metadata for text chunks
    table_metadata = []  # Track metadata for table chunks
    
    # Process all PDF files with metadata
    for filename in os.listdir(folder_path):
        if filename.endswith(".pdf"):
            pdf_path = os.path.join(folder_path, filename)
            
            # Process PDF with metadata tracking
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    text = page.extract_text() or ""
                    if text:
                        text_chunks = text_splitter.split_text(text)
                        all_texts.extend(text_chunks)
                        # Add metadata for each text chunk
                        for chunk in text_chunks:
                            text_metadata.append({
                                "source": filename,
                                "page": page_num,
                                "type": "text"
                            })
                    
                    # Extract tables
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables, 1):
                        if table and len(table) > 1:
                            df = pd.DataFrame(table[1:], columns=table[0])
                            table_markdown = dataframe_to_clean_markdown(df)
                            
                            # Enrich table for better BM25 matching
                            column_names = " ".join([str(col).lower() for col in df.columns if col])
                            sample_values = []
                            num_sample_cols = min(3, len(df.columns))
                            for col_idx in range(num_sample_cols):
                                series = df.iloc[:, col_idx]
                                sample_values.extend([str(val).lower() for val in series.dropna().head(3).tolist()])
                            sample_context = " ".join(sample_values[:10])
                            enriched_table = f"[TABLE DATA] Topic: {column_names} {sample_context}\n\n{table_markdown}\n\n[END TABLE]"
                            table_chunks.append(enriched_table)
                            # Add metadata for table chunk
                            table_metadata.append({
                                "source": filename,
                                "page": page_num,
                                "type": "table"
                            })
    
    # Combine text and tables for indexing
    all_chunks = all_texts + table_chunks
    bm25_metadata = text_metadata + table_metadata
    
    if all_chunks:
        # Tokenize for BM25
        bm25_corpus = [text.split() for text in all_chunks]
        bm25_index = BM25Okapi(bm25_corpus)
        logging.info(f"✅ BM25 index built with {len(bm25_corpus)} document chunks (with metadata tracking)")
    else:
        logging.warning("⚠️ No content found for BM25 indexing")

def expand_query_with_llm(question, llm):
    """Expands user query using LLM to include synonyms but retains original meaning."""
    expansion_prompt = f"""
    Provide alternative phrasings and related terms for: '{question}', 
    ensuring the original word is always included. Include HR-specific terms if applicable.
    """
    try:
        expanded_query = llm.invoke(expansion_prompt).content
        logging.info(f"🔍 Query Expansion: {expanded_query}")
        return expanded_query
    except Exception as e:
        logging.error(f"❌ Query Expansion Failed: {e}")
        return question  # Fall back to the original question

def hybrid_search(question, llm, retriever):
    """Performs hybrid retrieval using BM25 and Pinecone vectors."""
    global bm25_index, bm25_corpus
    
    # Expand query
    expanded_query = expand_query_with_llm(question, llm)
    
    results = []
    
    # Step 1: BM25 Keyword Search
    if bm25_index and bm25_corpus:
        bm25_results = bm25_index.get_top_n(expanded_query.split(), bm25_corpus, n=5)
        bm25_texts = [" ".join(text) for text in bm25_results]
        results.extend(bm25_texts)
        logging.info(f"🔍 BM25 Retrieved {len(bm25_texts)} results")
    
    # Step 2: Vector Search
    pinecone_results = retriever.invoke(expanded_query)
    pinecone_texts = [doc.page_content for doc in pinecone_results]
    results.extend(pinecone_texts)
    
    # Prioritize table content (tables contain | character in markdown)
    table_texts = [text for text in results if "|" in text]
    non_table_texts = [text for text in results if "|" not in text]
    
    # Combine results: tables first, then other content
    combined_results = table_texts + non_table_texts
    
    # Remove duplicates while preserving order
    unique_results = []
    seen = set()
    for text in combined_results:
        # Use a hash of the text as a unique identifier
        text_hash = hash(text)
        if text_hash not in seen:
            seen.add(text_hash)
            unique_results.append(text)
    
    # Join and truncate to avoid token limits
    final_text = "\n\n".join(unique_results)[:5000]
    
    return final_text

def save_qa_to_db(question, retrieved_docs, final_answer, feedback=None):
    """Stores a Q&A pair in SQLite with optional feedback."""
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        logging.info(f"Saving Q&A to DB - Question: {question[:50]}...")  # Debug log
        
        query = """
        INSERT INTO qa_history (question, retrieved_docs, final_answer, feedback) 
        VALUES (?, ?, ?, ?)
        """
        cursor.execute(query, (question, retrieved_docs, final_answer, feedback))
        conn.commit()
        
        question_id = cursor.lastrowid
        logging.info(f"✅ Q&A stored successfully with ID: {question_id}")
        return question_id
    except Exception as e:
        logging.error(f"❌ Error saving Q&A to DB: {e}", exc_info=True)
        return None
    finally:
        conn.close()

def setup_llm_chain():
    """Initialize the LLM and retrieval chain."""
    # Initialize LLM with optimized parameters
    llm = ChatGroq(
        # model_name= "mixtral-8x7b-32768", 
        #  model_name= "llama-3.1-8b-instant",
        model_name =  "qwen/qwen3-32b",     #"qwen-2.5-32b",
        groq_api_key=GROQ_API_KEY,
        temperature=0.377,
        max_tokens=32768,
        top_p=0.95,
        presence_penalty=0.1,
        frequency_penalty=0.1,
        streaming=True
    )
    
    # Initialize retriever only if vectorstore is available
    retriever = None
    if vectorstore is not None:
        try:
            retriever = vectorstore.as_retriever(search_kwargs={"k": 5})
            logging.info("✅ Retriever initialized successfully")
        except Exception as e:
            logging.error(f"❌ Error initializing retriever: {e}")
            retriever = None
    else:
        logging.warning("⚠️ Vectorstore not available, retriever will be None")
    
    return llm, None, retriever  # Return llm, qa_chain (None for now), retriever

def expand_acronyms(question):
    """Expand HR-related acronyms in the question."""
    expanded_question = question.lower()
    for acronym, full_form in ACRONYM_MAP.items():
        expanded_question = expanded_question.replace(acronym.lower(), full_form.lower())
    return expanded_question

async def analyze_career_progression(resume_text):
    """Analyze career progression from resume text using Gemini."""
    try:
        formatted_prompt = f"""You are an expert HR analyst. Analyze this candidate's career progression.
Return ONLY a JSON object with the following structure, no other text:
{{
    "progression_score": <number 0-100>,
    "key_observations": [<list of string observations>],
    "career_path": [
        {{
            "title": "<job title>",
            "company": "<company name>",
            "duration": "<time period>",
            "level": "<Entry/Mid/Senior/Lead/Manager>",
            "progression": "<Promotion/Lateral/Step Back>"
        }}
    ],
    "red_flags": [<list of string concerns>],
    "reasoning": "<analysis explanation>"
}}

Resume text:
{resume_text}"""

        # Get response from Gemini
        response = await async_gemini_generate(formatted_prompt)
        
        # If response is already a dict (from async_gemini_generate)
        if isinstance(response, dict):
            parsed_response = response
        else:
            try:
                parsed_response = json.loads(response) if isinstance(response, str) else {}
            except json.JSONDecodeError:
                logging.error(f"Failed to parse response as JSON: {response}")
                return get_default_career_analysis()

        # Validate and clean the response data
        cleaned_data = {
            "progression_score": validate_progression_score(parsed_response.get("progression_score", 50)),
            "key_observations": validate_list(parsed_response.get("key_observations", [])) or ["No key observations found"],
            "career_path": validate_career_path(parsed_response.get("career_path", [])),
            "red_flags": validate_list(parsed_response.get("red_flags", [])) or ["No red flags identified"],
            "reasoning": str(parsed_response.get("reasoning", "No analysis provided")).strip()
        }

        # Ensure we have valid data
        if cleaned_data["progression_score"] == 50 and not cleaned_data["career_path"]:
            return get_default_career_analysis()
            
        return cleaned_data

    except Exception as e:
        logging.error(f"Career progression analysis error: {str(e)}")
        logging.error(f"Full traceback:", exc_info=True)
        return get_default_career_analysis()

def get_default_career_analysis():
    """Return default career analysis structure"""
    return {
        "progression_score": 50,
        "key_observations": ["Unable to analyze career progression"],
        "career_path": [],
        "red_flags": ["Analysis encountered technical issues"],
        "reasoning": "Analysis failed to complete"
    }

def get_default_resume_evaluation():
    """Return default resume evaluation structure when JSON parsing fails"""
    return {
        "JD Match": "0%",
        "MissingKeywords": [],
        "Profile Summary": "Unable to analyze resume due to technical error. Please try again.",
        "Over/UnderQualification Analysis": "Analysis unavailable",
        "Match Factors": {
            "Skills Match": 0,
            "Experience Match": 0,
            "Education Match": 0,
            "Industry Knowledge": 0,
            "Certification Match": None
        },
        "Reasoning": "JSON parsing failed. The AI response could not be properly parsed.",
        "Candidate Fit Analysis": {
            "Dimension Evaluation": [],
            "Risk and Gaps": None,
            "Recommendation": {
                "Verdict": "❌ Analysis Failed",
                "Fit Level": "Unknown",
                "Rationale": "Technical error prevented proper evaluation"
            },
            "Recruiter Narrative": "Unable to generate evaluation due to technical issues."
        }
    }

def validate_progression_score(score):
    """Validate and normalize progression score"""
    try:
        if isinstance(score, str):
            score = score.strip('%')
        score = float(score)
        return int(max(0, min(100, score)))
    except (ValueError, TypeError):
        return 50

def validate_list(items):
    """Validate and clean list items"""
    if not isinstance(items, list):
        return []
    return [str(item).strip() for item in items if item and str(item).strip()]

def validate_career_path(path):
    """Validate and clean career path entries"""
    if not isinstance(path, list):
        return []
    
    cleaned_path = []
    required_fields = ["title", "company", "duration", "level", "progression"]
    
    for entry in path:
        if not isinstance(entry, dict):
            continue
        
        cleaned_entry = {}
        for field in required_fields:
            cleaned_entry[field] = str(entry.get(field, "Not specified")).strip()
        cleaned_path.append(cleaned_entry)
    
    return cleaned_path

def update_db_schema():
    """Update database schema if needed"""
    conn = sqlite3.connect('combined_db.db')
    cursor = conn.cursor()
    
    # Add new columns if they don't exist
    try:
        cursor.execute('''
            ALTER TABLE evaluations 
            ADD COLUMN job_stability TEXT;
        ''')
    except sqlite3.OperationalError:
        pass  # Column already exists
        
    try:
        cursor.execute('''
            ALTER TABLE evaluations 
            ADD COLUMN career_progression TEXT;
        ''')
    except sqlite3.OperationalError:
        pass  # Column already exists
    
    # Add oorwin_job_id column to evaluations table
    try:
        cursor.execute('''
            ALTER TABLE evaluations 
            ADD COLUMN oorwin_job_id TEXT;
        ''')
        logging.info("Added oorwin_job_id column to evaluations table")
    except sqlite3.OperationalError:
        pass  # Column already exists
    
    # Add token_usage column to evaluations table
    try:
        cursor.execute('''
            ALTER TABLE evaluations 
            ADD COLUMN token_usage INTEGER;
        ''')
        logging.info("Added token_usage column to evaluations table")
    except sqlite3.OperationalError:
        pass  # Column already exists
    
    # Create recruiter_handbooks table
    try:
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS recruiter_handbooks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                oorwin_job_id TEXT,
                job_title TEXT,
                job_description TEXT,
                additional_context TEXT,
                markdown_content TEXT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        logging.info("Created recruiter_handbooks table")
    except sqlite3.OperationalError:
        pass  # Table already exists
    
    # Create index on oorwin_job_id for faster queries
    try:
        cursor.execute('''
            CREATE INDEX IF NOT EXISTS idx_evaluations_job_id 
            ON evaluations(oorwin_job_id)
        ''')
        cursor.execute('''
            CREATE INDEX IF NOT EXISTS idx_handbooks_job_id 
            ON recruiter_handbooks(oorwin_job_id)
        ''')
        logging.info("Created indexes on oorwin_job_id")
    except sqlite3.OperationalError:
        pass  # Index already exists
    
    conn.commit()
    conn.close()

@app.route('/api/evaluation/<evaluation_id>', methods=['GET'])
def get_evaluation_details(evaluation_id):
    """API endpoint to get evaluation details by ID"""
    conn = None
    try:
        logging.info(f"Fetching evaluation details for ID: {evaluation_id}")
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # Helper function for parsing JSON safely
        def parse_json_safely(json_str):
            if not json_str:
                logging.info("Empty JSON string, returning empty list")
                return []
            try:
                data = json.loads(json_str)
                if isinstance(data, list):
                    logging.info(f"Successfully parsed list with {len(data)} items")
                    return data
                elif isinstance(data, str):
                    try:
                        parsed_data = json.loads(data)
                        logging.info(f"Successfully parsed nested JSON string")
                        return parsed_data
                    except:
                        logging.info(f"Failed to parse nested JSON, treating as single item")
                        return [data]
                else:
                    logging.info(f"Non-list data type: {type(data)}, converting to string")
                    return [str(data)]
            except json.JSONDecodeError as e:
                logging.warning(f"JSON decode error: {str(e)}, attempting cleanup")
                try:
                    # Try to clean and parse the string
                    cleaned_str = json_str.strip('[]"\' ').replace('\\', '')
                    items = [item.strip('"\' ') for item in cleaned_str.split(',')]
                    result = [item for item in items if item]
                    logging.info(f"Cleanup successful, extracted {len(result)} items")
                    return result
                except Exception as e2:
                    logging.error(f"Cleanup failed: {str(e2)}")
                    return []
        
        # Get evaluation details first to get job title for default questions
        cursor.execute('''
            SELECT 
                e.id, 
                e.filename, 
                e.job_title, 
                e.match_percentage, 
                e.profile_summary, 
                e.job_stability,
                e.career_progression,
                e.timestamp,
                e.missing_keywords,
                e.behavioral_questions,
                e.technical_questions,
                e.nontechnical_questions
            FROM evaluations e
            WHERE e.id = ?
        ''', (evaluation_id,))
        
        row = cursor.fetchone()
        if not row:
            logging.warning(f"No evaluation found with ID: {evaluation_id}")
            return jsonify({'error': 'Evaluation not found'}), 404
        
        logging.info(f"Found evaluation with ID: {row[0]}, filename: {row[1]}")
        job_title = row[2]
        
        # Parse JSON fields
        try:
            job_stability = json.loads(row[5]) if row[5] else {}
            logging.info(f"Parsed job_stability: {type(job_stability)}")
        except Exception as e:
            logging.error(f"Error parsing job_stability: {str(e)}")
            job_stability = {}
            
        try:
            career_progression = json.loads(row[6]) if row[6] else {}
            logging.info(f"Parsed career_progression: {type(career_progression)}")
        except Exception as e:
            logging.error(f"Error parsing career_progression: {str(e)}")
            career_progression = {}
        
        # Parse missing keywords with special handling
        try:
            missing_keywords_raw = row[8]
            if missing_keywords_raw:
                try:
                    missing_keywords = json.loads(missing_keywords_raw)
                    logging.info(f"Parsed missing_keywords: {type(missing_keywords)}")
                    # If it's not a list, try to convert it
                    if not isinstance(missing_keywords, list):
                        if isinstance(missing_keywords, str):
                            # Remove brackets and split by commas
                            missing_keywords = [k.strip(' "\'') for k in missing_keywords.strip('[]').split(',')]
                        else:
                            missing_keywords = [str(missing_keywords)]
                except Exception as e:
                    logging.error(f"Error parsing missing_keywords JSON: {str(e)}")
                    # If JSON parsing fails, try to extract from string
                    if isinstance(missing_keywords_raw, str):
                        # Check if it looks like a list
                        if missing_keywords_raw.startswith('[') and missing_keywords_raw.endswith(']'):
                            # Remove brackets and split by commas
                            missing_keywords = [k.strip(' "\'') for k in missing_keywords_raw.strip('[]').split(',')]
                        else:
                            missing_keywords = [missing_keywords_raw]
                    else:
                        missing_keywords = []
            else:
                missing_keywords = []
        except Exception as e:
            logging.error(f"Error processing missing_keywords: {str(e)}")
            missing_keywords = []
        
        # Initialize question variables
        technical_questions = []
        nontechnical_questions = []
        behavioral_questions = []
        
        # PRIORITY 1: Try to get interview questions from interview_questions table FIRST
        # This is the dedicated table for storing interview questions
        cursor.execute(
            "SELECT technical_questions, nontechnical_questions, behavioral_questions FROM interview_questions WHERE evaluation_id = ?",
            (evaluation_id,)
        )
        iq_result = cursor.fetchone()
        
        if not iq_result:
            logging.info(f"No interview questions found with numeric ID, trying string ID")
            # If no results, try with the string representation of the ID
            cursor.execute(
                "SELECT technical_questions, nontechnical_questions, behavioral_questions FROM interview_questions WHERE evaluation_id = ?",
                (str(evaluation_id),)
            )
            iq_result = cursor.fetchone()
        
        if iq_result:
            logging.info(f"Found interview questions in interview_questions table for evaluation {evaluation_id}")
            technical_questions = parse_json_safely(iq_result[0])
            nontechnical_questions = parse_json_safely(iq_result[1])
            behavioral_questions = parse_json_safely(iq_result[2])
            logging.info(f"Retrieved from interview_questions table: {len(technical_questions)} technical, {len(nontechnical_questions)} non-technical, {len(behavioral_questions)} behavioral questions")
        else:
            logging.info(f"No interview questions found in interview_questions table for evaluation {evaluation_id}")
            
            # PRIORITY 2: Fallback to evaluations table if interview_questions table is empty
            logging.info("Falling back to evaluations table for questions")
        
        # Try to get behavioral questions from evaluations
        try:
            behavioral_questions_raw = row[9]
            if behavioral_questions_raw:
                behavioral_questions = parse_json_safely(behavioral_questions_raw)
                logging.info(f"Parsed behavioral_questions from evaluations: {len(behavioral_questions)} questions")
        except Exception as e:
            logging.error(f"Error parsing behavioral_questions from evaluations: {str(e)}")
        
        # Try to get technical questions from evaluations
        try:
            if row[10]:
                technical_questions = parse_json_safely(row[10])
                logging.info(f"Parsed technical_questions from evaluations: {len(technical_questions)} questions")
        except Exception as e:
            logging.error(f"Error parsing technical_questions from evaluations: {str(e)}")
        
        # Try to get non-technical questions from evaluations
        try:
            if row[11]:
                nontechnical_questions = parse_json_safely(row[11])
                logging.info(f"Parsed nontechnical_questions from evaluations: {len(nontechnical_questions)} questions")
        except Exception as e:
            logging.error(f"Error parsing nontechnical_questions from evaluations: {str(e)}")
        
        # If still no behavioral questions, use default QUICK_CHECKS
        if not behavioral_questions:
            logging.info("No behavioral questions found, using QUICK_CHECKS")
            behavioral_questions = QUICK_CHECKS
        
        
        # Only generate default questions if we still don't have any questions at all
        if not technical_questions and not nontechnical_questions:
            logging.info(f"Generating default questions for job title: {job_title}")
            default_technical, default_nontechnical = get_default_interview_questions(job_title)
            
            if not technical_questions:
                technical_questions = default_technical
                logging.info(f"Using default technical questions: {len(technical_questions)} questions")
            
            if not nontechnical_questions:
                nontechnical_questions = default_nontechnical
                logging.info(f"Using default non-technical questions: {len(nontechnical_questions)} questions")
        
        # Create response
        response = {
            'id': row[0],
            'filename': row[1],
            'job_title': row[2],
            'match_percentage': row[3],
            'profile_summary': row[4] or "No summary available",
            'job_stability': job_stability,
            'career_progression': career_progression,
            'timestamp': row[7],
            'missing_keywords': missing_keywords,
            'technical_questions': technical_questions,
            'nontechnical_questions': nontechnical_questions,
            'behavioral_questions': behavioral_questions
        }
        
        logging.info(f"Returning response with {len(technical_questions)} technical questions, {len(nontechnical_questions)} non-technical questions, {len(behavioral_questions)} behavioral questions")
        return jsonify(response)
    
    except Exception as e:
        logging.error(f"Error fetching evaluation details: {str(e)}")
        return jsonify({'error': str(e)}), 500
    
    finally:
        if conn:
            conn.close()

@app.route('/api/generate_questions/<evaluation_id>', methods=['POST'])
async def generate_questions_api(evaluation_id):
    """API endpoint to generate interview questions for an evaluation"""
    conn = None
    try:
        logging.info(f"Generating questions for evaluation ID: {evaluation_id}")
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # Get evaluation details
        cursor.execute(
            """
            SELECT resume_path, job_title, job_description, profile_summary 
            FROM evaluations 
            WHERE id = ?
            """,
            (evaluation_id,)
        )
        eval_result = cursor.fetchone()
        
        if not eval_result:
            logging.warning(f"No evaluation found with ID: {evaluation_id}")
            return jsonify({'error': 'Evaluation not found'}), 404
        
        # Extract resume text
        resume_path = eval_result[0]
        job_description = eval_result[2]
        profile_summary = eval_result[3]
        
        if not resume_path or resume_path == 'NULL' or resume_path == 'None' or str(resume_path).strip() == '':
            logging.error(f"Invalid resume_path for evaluation {evaluation_id}: {resume_path}")
            return jsonify({'error': 'No valid resume path found for this evaluation'}), 400
        
        resume_text = extract_text_from_file(resume_path)
        if not resume_text:
            return jsonify({'error': 'Failed to extract text from resume'}), 400
        
        # Generate questions
        logging.info(f"Generating questions for resume: {resume_path}")
        questions_data = await async_generate_questions(
            resume_text,
            job_description,
            profile_summary
        )
        
        technical_questions = questions_data.get("TechnicalQuestions", [])
        nontechnical_questions = questions_data.get("NonTechnicalQuestions", [])
        behavioral_questions = QUICK_CHECKS
        
        # Save questions to database
        try:
            # First check if there's an existing entry
            cursor.execute(
                "SELECT id FROM interview_questions WHERE evaluation_id = ?",
                (evaluation_id,)
            )
            existing = cursor.fetchone()
            
            if existing:
                # Update existing entry
                cursor.execute(
                    """
                    UPDATE interview_questions 
                    SET technical_questions = ?,
                        nontechnical_questions = ?,
                        behavioral_questions = ?
                    WHERE evaluation_id = ?
                    """,
                    (
                        json.dumps(technical_questions), 
                        json.dumps(nontechnical_questions), 
                        json.dumps(behavioral_questions), 
                        evaluation_id
                    )
                )
            else:
                # Insert new entry
                cursor.execute(
                    """
                    INSERT INTO interview_questions 
                    (evaluation_id, technical_questions, nontechnical_questions, behavioral_questions) 
                    VALUES (?, ?, ?, ?)
                    """,
                    (
                        evaluation_id,
                        json.dumps(technical_questions),
                        json.dumps(nontechnical_questions),
                        json.dumps(behavioral_questions)
                    )
                )
            
            conn.commit()
            logging.info(f"Saved questions for evaluation ID: {evaluation_id}")
        except Exception as e:
            logging.error(f"Error saving questions to database: {str(e)}")
            conn.rollback()
        
        # Return the generated questions
        return jsonify({
            'technical_questions': technical_questions,
            'nontechnical_questions': nontechnical_questions,
            'behavioral_questions': behavioral_questions
        })
        
    except Exception as e:
        logging.error(f"Error generating questions: {str(e)}")
        return jsonify({'error': str(e)}), 500
        
    finally:
        if conn:
            conn.close()

@app.route('/api/generate-recruiter-handbook', methods=['POST'])
@login_required
def generate_recruiter_handbook():
    """API endpoint to generate a comprehensive recruiter handbook"""
    try:
        data = request.get_json()
        job_title = data.get('job_title', '').strip()
        job_description = data.get('job_description', '').strip()
        additional_context = data.get('additional_context', '').strip()
        oorwin_job_id = data.get('oorwin_job_id', '').strip()
        
        if not job_title:
            return jsonify({
                'success': False,
                'message': 'Job title is required'
            }), 400
            
        if not job_description:
            return jsonify({
                'success': False,
                'message': 'Job description is required'
            }), 400
        
        # Check if handbook already exists for this job ID
        if oorwin_job_id:
            try:
                conn = sqlite3.connect('combined_db.db')
                cursor = conn.cursor()
                
                cursor.execute('''
                    SELECT id, markdown_content, job_title, timestamp, user_email
                    FROM recruiter_handbooks
                    WHERE oorwin_job_id = ?
                    ORDER BY timestamp DESC
                    LIMIT 1
                ''', (oorwin_job_id,))
                
                existing_handbook = cursor.fetchone()
                conn.close()
                
                if existing_handbook:
                    handbook_id, markdown_content, existing_job_title, timestamp, creator_email = existing_handbook
                    logging.info(f"Handbook already exists for JobID {oorwin_job_id} (ID: {handbook_id}, created: {timestamp})")
                    
                    # Get creator name if available
                    creator_name = None
                    if creator_email:
                        try:
                            conn = sqlite3.connect('combined_db.db')
                            cursor = conn.cursor()
                            cursor.execute('SELECT name FROM users WHERE email = ?', (creator_email,))
                            user_row = cursor.fetchone()
                            if user_row:
                                creator_name = user_row[0] or creator_email
                            else:
                                creator_name = creator_email
                            conn.close()
                        except:
                            creator_name = creator_email
                    
                    return jsonify({
                        'success': True,
                        'existing': True,
                        'message': f'A handbook already exists for Job ID "{oorwin_job_id}". Showing existing handbook.',
                        'handbook_id': handbook_id,
                        'markdown_content': markdown_content,
                        'job_title': existing_job_title or job_title,
                        'oorwin_job_id': oorwin_job_id,
                        'created_at': timestamp,
                        'created_by': creator_name or creator_email or 'Unknown',
                        'created_by_email': creator_email
                    })
            except Exception as check_error:
                logging.error(f"Error checking for existing handbook: {str(check_error)}")
                # Continue with generation if check fails
            
        logging.info(f"Generating new recruiter handbook for JobID: {oorwin_job_id or 'None'}...")
        
        # Track start time for performance metrics
        import time
        handbook_start_time = time.time()
        
        # Create the prompt for Gemini
        handbook_prompt = f"""You are an expert recruitment specialist with deep experience in creating recruiter playbooks and handbooks for various roles across industries. Your task is to analyze a provided Job Description (JD) and generate a comprehensive "Recruiter Playbook & Handbook" in a structured, professional format that mirrors the style, structure, and content of the example playbooks you have been trained on (e.g., for roles like Head of Engineering at Fractal, Product Manager – Monogastric at Jubilant Ingrevia, and Senior Research Scientist – Analytical Chemistry at Jubilant Ingrevia).

Key guidelines for the output:

**CRITICAL ORDER - Follow this exact sequence (NO duplicates, NO extra sections):**

1. **Title**: Start with an emoji like 📖, followed by "Recruiter Playbook & Handbook: [Role Title] ([Company Name])".

2. **Mini Table of Contents**: Immediately after the title, add a compact "Mini Table of Contents" with markdown links to each main section (H2 headings). Keep it to one line per section. Format as: "- [Section Name](#section-name)"

3. **Introduction**: A brief paragraph (ONLY ONE paragraph) explaining that the handbook equips recruiters with JD analysis, screening framework, sourcing tactics, red flags, a recruiter sales pitch, and more to engage candidates effectively. Include the job link if provided in the JD; otherwise, omit it. Do NOT repeat the title or add any other content before this introduction.

4. Structure the content exactly as follows (use numbered sections, bullet points, and sub-bullets for clarity; incorporate emojis like ✨ for sales pitch, ✅ for closing, • for lists):

   **1. Primary Sourcing Parameters (Must-Have)**: Produce a compact, scannable table using GitHub-Flavored Markdown with EXACTLY these columns and order:
      | # | Skill / Experience | Recruiter Cue | Why It Matters |
      Then add 6–8 rows, where:
      - "#" is a running number starting at 1 (plain numbers only)
      - "Skill / Experience" is the must-have capability (e.g., Microfrontend, React + Next.js/Remix, Frontend Platform Development, System Design, Tech Leadership & Mentorship, Strong Communication)
      - "Recruiter Cue" lists concrete signals/cues (comma-separated; keep concise, e.g., "Module Federation, SingleSPA, NX, Turborepo")
      - "Why It Matters" explains impact/value in a short phrase (one line)
      CRITICAL: Output MUST be a valid markdown table with header separator row (|---|---|---|---|). No extra prose between the H2 and the table.

   **2. Screening Framework**: Categorize into sections like A. Background & Motivation, B. Domain Experience, etc. (up to G. Practicalities). Each section should have 1-3 bullet-pointed screening questions derived from the JD. Keep questions open-ended and probing.

   **3. Target Talent Pools**:
      - **Likely Companies**: List 4-8 relevant companies, ONE per line as separate bullets. Do NOT put multiple companies on one line.
      - **Likely Titles**: List 3-5 alternative job titles, ONE per line as separate bullets. Do NOT put multiple titles on one line.
      - **Boolean Search Samples**: Provide **EXACTLY 3 DIFFERENT Boolean search strings**, each one **STRICTLY under 200 characters**. Format each as follows:
        * **Sample 1 (Skills-focused):** `[your boolean string here]`
        * **Sample 2 (Company-focused):** `[your boolean string here]`
        * **Sample 3 (Title-focused):** `[your boolean string here]`
        CRITICAL: Each string MUST be under 200 characters. Count characters carefully. Focus on key terms only.

   **4. Red Flags to Watch**: List 4-6 bullet-pointed red flags (e.g., lack of specific experience) based on potential mismatches from the JD.

   **5. Recruiter Sales Pitch (to candidates)**: Start with ✨ **Why [Company]?** List 5-7 bullet points highlighting company strengths, role impact, growth opportunities, etc. Infer from the JD or common knowledge; make it engaging and positive. End with a closing tagline.

   **6. Recruiter Checklist (Pre-call)**: List 4-6 bullet points of key pre-call actions (e.g., confirm experience, probe specifics).

   **7. Overqualification/Overkill Risk Assessment**: ⚠️ **CRITICAL** - Analyze if candidates for this role might be overqualified or "overkill". Consider:
      - **Experience Level Mismatch**: If JD asks for 3-5 years but attracts 10+ year candidates
      - **Title Level Gap**: If role is Manager level but attracts Director/VP candidates  
      - **Compensation Concerns**: If role budget might not match senior candidates' expectations
      - **Flight Risk Indicators**: Red flags that candidate may leave quickly (overqualified, career plateau, lateral move)
      - **When to Proceed Anyway**: Scenarios where overqualified candidates are worth considering (career change, lifestyle choice, genuine interest)
      - **Screening Questions**: 2-3 specific questions to probe motivation and flight risk for overqualified candidates
      
      End with ✅ followed by: "This handbook provides recruiters with JD analysis, structured screening questions, sourcing pools, red flags, and overqualification assessment. Use as a starting document and conduct your own research before commencing the search."

**Style**: Professional, concise, actionable. Use bold for section headers. Incorporate industry-specific nuances from the JD. Assume good intent and focus on fit. Do not add unrelated content. Output in markdown format for readability.

---

**Job Description:**
{job_description}

{"**Additional Context:**" if additional_context else ""}
{additional_context if additional_context else ""}

---

Generate the complete Recruiter Playbook & Handbook now:"""

        # Generate handbook using selected model
        try:
            response = generate_content_unified(handbook_prompt)
            logging.info(f"Response received: {type(response)}, has text attr: {hasattr(response, 'text')}")
            
            if not response:
                raise Exception("No response received from AI model")
            
            # Debug: Check what's in response.text
            if hasattr(response, 'text'):
                logging.info(f"response.text type: {type(response.text)}, value: '{response.text[:100] if response.text else 'None or empty'}'")
        
            if not hasattr(response, 'text'):
                logging.error(f"Response object has no 'text' attribute")
                logging.error(f"Response attributes: {dir(response)}")
                raise Exception(f"Response object missing 'text' attribute. Type: {type(response)}")
            
            if not response.text:
                logging.error(f"response.text is empty or None")
                logging.error(f"response.text value: {repr(response.text)}")
                # Try output_text as fallback
                if hasattr(response, 'output_text') and response.output_text:
                    logging.info("Using output_text as fallback")
                    handbook_content = response.output_text
                else:
                    raise Exception(f"AI response text is empty. response.text={repr(response.text)}")
            else:
                handbook_content = response.text
            
            logging.info(f"Handbook content length: {len(handbook_content)} characters")
        except Exception as api_error:
            logging.error(f"API call failed: {str(api_error)}", exc_info=True)
            raise
        
        logging.info("Recruiter handbook generated successfully")
        
        # Calculate time taken
        handbook_time_taken = round(time.time() - handbook_start_time, 2)  # Round to 2 decimal places
        logging.info(f"Handbook generation took {handbook_time_taken} seconds")
        
        # Save handbook to database
        try:
            conn = sqlite3.connect('combined_db.db')
            cursor = conn.cursor()
            
            # Get user email from session
            user_email = session.get('user', {}).get('email') if 'user' in session else None
            
            cursor.execute('''
                INSERT INTO recruiter_handbooks (
                    oorwin_job_id, job_title, job_description, 
                    additional_context, markdown_content, time_taken, user_email, timestamp
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                oorwin_job_id if oorwin_job_id else None,
                job_title,
                job_description,
                additional_context,
                handbook_content,
                handbook_time_taken,
                user_email,
                datetime.now()
            ))
            
            conn.commit()
            handbook_id = cursor.lastrowid
            conn.close()
            
            logging.info(f"Handbook saved to database with ID: {handbook_id}")
            
            # Return success with handbook_id
            return jsonify({
                'success': True,
                'markdown_content': handbook_content,
                'handbook_id': handbook_id
            })
        except Exception as e:
            logging.error(f"Error saving handbook to database: {str(e)}")
            # Return without handbook_id if save fails
            return jsonify({
                'success': True,
                'markdown_content': handbook_content
            })
        
    except Exception as e:
        logging.error(f"Error generating recruiter handbook: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/get-job-ids', methods=['GET'])
def get_job_ids():
    """API endpoint to get all unique JobIDs for auto-suggest"""
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # Get unique JobIDs from evaluations
        cursor.execute('''
            SELECT DISTINCT oorwin_job_id 
            FROM evaluations 
            WHERE oorwin_job_id IS NOT NULL AND oorwin_job_id != ''
            ORDER BY oorwin_job_id
        ''')
        eval_job_ids = [row[0] for row in cursor.fetchall()]
        
        # Get unique JobIDs from handbooks
        cursor.execute('''
            SELECT DISTINCT oorwin_job_id 
            FROM recruiter_handbooks 
            WHERE oorwin_job_id IS NOT NULL AND oorwin_job_id != ''
            ORDER BY oorwin_job_id
        ''')
        handbook_job_ids = [row[0] for row in cursor.fetchall()]
        
        # Combine and remove duplicates
        all_job_ids = list(set(eval_job_ids + handbook_job_ids))
        all_job_ids.sort()
        
        conn.close()
        
        return jsonify({
            'success': True,
            'job_ids': all_job_ids
        })
        
    except Exception as e:
        logging.error(f"Error fetching JobIDs: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/get-job-data/<job_id>', methods=['GET'])
def get_job_data(job_id):
    """API endpoint to get job description for auto-fill based on JobID"""
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # First try to get from handbooks (they always have full JD)
        cursor.execute('''
            SELECT job_title, job_description, additional_context 
            FROM recruiter_handbooks 
            WHERE oorwin_job_id = ? 
            ORDER BY timestamp DESC 
            LIMIT 1
        ''', (job_id,))
        
        result = cursor.fetchone()
        
        if result and result[1]:  # Check if job_description is not empty
            conn.close()
            return jsonify({
                'success': True,
                'job_title': result[0],
                'job_description': result[1],
                'additional_context': result[2] if result[2] else '',
                'source': 'handbook'
            })
        
        # If not found in handbooks or JD is empty, try evaluations
        cursor.execute('''
            SELECT job_title, job_description 
            FROM evaluations 
            WHERE oorwin_job_id = ? AND job_description IS NOT NULL AND job_description != ''
            ORDER BY timestamp DESC 
            LIMIT 1
        ''', (job_id,))
        
        result = cursor.fetchone()
        conn.close()
        
        if result:
            return jsonify({
                'success': True,
                'job_title': result[0],
                'job_description': result[1],
                'source': 'evaluation'
            })
        else:
            return jsonify({
                'success': False,
                'message': 'No data found for this JobID'
            }), 404
        
    except Exception as e:
        logging.error(f"Error fetching job data for {job_id}: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/get-handbooks', methods=['GET'])
@login_required
def get_handbooks():
    """API endpoint to get all recruiter handbooks (filtered by role)"""
    try:
        # Get accessible user emails based on role
        user_email = session['user'].get('email')
        accessible_emails = get_accessible_user_emails(user_email)
        
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        if not accessible_emails:
            return jsonify({
                'success': True,
                'handbooks': []
            })
        
        # Build WHERE clause for filtering
        placeholders = ','.join(['?'] * len(accessible_emails))
        where_clause = f"WHERE user_email IN ({placeholders})"
        
        cursor.execute(f'''
            SELECT id, oorwin_job_id, job_title, job_description, 
                   additional_context, markdown_content, timestamp
            FROM recruiter_handbooks
            {where_clause}
            ORDER BY timestamp DESC
        ''', accessible_emails)
        
        rows = cursor.fetchall()
        conn.close()
        
        handbooks = []
        for row in rows:
            handbooks.append({
                'id': row[0],
                'oorwin_job_id': row[1],
                'job_title': row[2],
                'job_description': row[3],
                'additional_context': row[4],
                'markdown_content': row[5],
                'timestamp': row[6]
            })
        
        return jsonify({
            'success': True,
            'handbooks': handbooks
        })
        
    except Exception as e:
        logging.error(f"Error fetching handbooks: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

# ===== DASHBOARD & ANALYTICS ROUTES =====

@app.route('/dashboard')
def dashboard():
    """Analytics dashboard showing usage metrics"""
    return render_template('dashboard.html')

@app.route('/api/analytics/overview', methods=['GET'])
@login_required
def get_analytics_overview():
    """API endpoint for dashboard overview metrics with user, team, and date filtering"""
    try:
        user_email = session['user'].get('email')
        filter_user = request.args.get('user_email', '')
        filter_team = request.args.get('team', '')
        date_from = request.args.get('date_from', '')
        date_to = request.args.get('date_to', '')
        
        # Get accessible user emails based on role
        accessible_emails = get_accessible_user_emails(user_email)
        
        # Debug logging
        logging.info(f"Analytics overview - Current user: {user_email}, Accessible emails: {accessible_emails}")
        
        # Filter by team if specified
        if filter_team:
            conn = sqlite3.connect(DATABASE_NAME)
            cursor = conn.cursor()
            cursor.execute('SELECT email FROM users WHERE team = ?', (filter_team,))
            team_emails = [row[0] for row in cursor.fetchall()]
            conn.close()
            # Intersect with accessible emails
            accessible_emails = [e for e in accessible_emails if e in team_emails]
        
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # Ensure we always include the current user's email if accessible_emails is empty
        if not accessible_emails:
            accessible_emails = [user_email]
            logging.info(f"Analytics overview - No accessible emails, using current user: {user_email}")
        
        logging.info(f"Analytics overview - Using accessible emails: {accessible_emails}")
        
        # Build WHERE clauses first, then build params to match
        eval_where_parts = []
        handbook_where_parts = []
        
        # Build user_email filter
        if filter_user and (filter_user in accessible_emails or filter_user == user_email):
            # When filtering by specific user, use exact match (don't include NULL records)
            eval_where_parts.append("user_email = ?")
            handbook_where_parts.append("user_email = ?")
        else:
            # Include accessible emails and NULL records
            placeholders = ','.join(['?'] * len(accessible_emails))
            eval_where_parts.append(f"(user_email IN ({placeholders}) OR user_email IS NULL)")
            handbook_where_parts.append(f"(user_email IN ({placeholders}) OR user_email IS NULL)")
        
        # Add date filters
        if date_from:
            eval_where_parts.append("DATE(timestamp) >= ?")
            handbook_where_parts.append("DATE(timestamp) >= ?")
        
        if date_to:
            eval_where_parts.append("DATE(timestamp) <= ?")
            handbook_where_parts.append("DATE(timestamp) <= ?")
        
        eval_where_clause = " WHERE " + " AND ".join(eval_where_parts) if eval_where_parts else ""
        handbook_where_clause = " WHERE " + " AND ".join(handbook_where_parts) if handbook_where_parts else ""
        
        logging.info(f"Analytics overview - Eval WHERE clause: {eval_where_clause}")
        logging.info(f"Analytics overview - Handbook WHERE clause: {handbook_where_clause}")
        
        # Build params to match WHERE clause placeholders (in order)
        eval_params = []
        handbook_params = []
        
        # Add user_email params (either filter_user or accessible_emails)
        if filter_user and (filter_user in accessible_emails or filter_user == user_email):
            eval_params.append(filter_user)
            handbook_params.append(filter_user)
        else:
            eval_params.extend(accessible_emails)
            handbook_params.extend(accessible_emails)
        
        # Add date params
        if date_from:
            eval_params.append(date_from)
            handbook_params.append(date_from)
        
        if date_to:
            eval_params.append(date_to)
            handbook_params.append(date_to)
        
        logging.info(f"Analytics overview - Eval params: {eval_params}")
        logging.info(f"Analytics overview - Handbook params: {handbook_params}")
        
        # Get total counts with filters
        cursor.execute(f'SELECT COUNT(*) FROM evaluations{eval_where_clause}', eval_params)
        total_evaluations = cursor.fetchone()[0]
        
        cursor.execute(f'SELECT COUNT(*) FROM recruiter_handbooks{handbook_where_clause}', handbook_params)
        total_handbooks = cursor.fetchone()[0]
        
        # Get unique jobs with filters
        # Need to add AND clause for oorwin_job_id check
        eval_job_where = eval_where_clause + (" AND " if eval_where_clause else " WHERE ") + "oorwin_job_id IS NOT NULL AND oorwin_job_id != ''"
        handbook_job_where = handbook_where_clause + (" AND " if handbook_where_clause else " WHERE ") + "oorwin_job_id IS NOT NULL AND oorwin_job_id != ''"
        
        cursor.execute(f'SELECT COUNT(DISTINCT oorwin_job_id) FROM evaluations{eval_job_where}', eval_params)
        unique_jobs_evals = cursor.fetchone()[0]
        
        cursor.execute(f'SELECT COUNT(DISTINCT oorwin_job_id) FROM recruiter_handbooks{handbook_job_where}', handbook_params)
        unique_jobs_handbooks = cursor.fetchone()[0]
        
        total_jobs = max(unique_jobs_evals, unique_jobs_handbooks)
        
        # Get average match score with filters
        # Need to add AND clause for match_percentage check
        eval_score_where = eval_where_clause + (" AND " if eval_where_clause else " WHERE ") + "match_percentage IS NOT NULL"
        cursor.execute(f'SELECT AVG(match_percentage) FROM evaluations{eval_score_where}', eval_params)
        avg_match_score = cursor.fetchone()[0] or 0
        avg_match_score = round(avg_match_score, 1)
        
        # Calculate previous period for comparison (if date range is specified)
        prev_total_evaluations = None
        prev_total_handbooks = None
        prev_total_jobs = None
        prev_avg_match_score = None
        
        if date_from and date_to:
            try:
                from datetime import datetime, timedelta
                # Calculate previous period (same duration before the selected period)
                from_date_obj = datetime.strptime(date_from, '%Y-%m-%d')
                to_date_obj = datetime.strptime(date_to, '%Y-%m-%d')
                period_days = (to_date_obj - from_date_obj).days + 1
                
                prev_to_date = from_date_obj - timedelta(days=1)
                prev_from_date = prev_to_date - timedelta(days=period_days - 1)
                
                prev_date_from = prev_from_date.strftime('%Y-%m-%d')
                prev_date_to = prev_to_date.strftime('%Y-%m-%d')
                
                # Build previous period WHERE clauses
                prev_eval_where_parts = []
                prev_handbook_where_parts = []
                prev_eval_params = []
                prev_handbook_params = []
                
                if accessible_emails:
                    placeholders = ','.join(['?'] * len(accessible_emails))
                    prev_eval_where_parts.append(f"user_email IN ({placeholders})")
                    prev_handbook_where_parts.append(f"user_email IN ({placeholders})")
                    prev_eval_params.extend(accessible_emails)
                    prev_handbook_params.extend(accessible_emails)
                
                if filter_user and filter_user in accessible_emails:
                    prev_eval_where_parts.append("user_email = ?")
                    prev_handbook_where_parts.append("user_email = ?")
                    prev_eval_params.append(filter_user)
                    prev_handbook_params.append(filter_user)
                
                prev_eval_where_parts.append("DATE(timestamp) >= ?")
                prev_eval_where_parts.append("DATE(timestamp) <= ?")
                prev_handbook_where_parts.append("DATE(timestamp) >= ?")
                prev_handbook_where_parts.append("DATE(timestamp) <= ?")
                prev_eval_params.extend([prev_date_from, prev_date_to])
                prev_handbook_params.extend([prev_date_from, prev_date_to])
                
                prev_eval_where_clause = " WHERE " + " AND ".join(prev_eval_where_parts) if prev_eval_where_parts else ""
                prev_handbook_where_clause = " WHERE " + " AND ".join(prev_handbook_where_parts) if prev_handbook_where_parts else ""
                
                cursor.execute(f'SELECT COUNT(*) FROM evaluations{prev_eval_where_clause}', prev_eval_params)
                prev_total_evaluations = cursor.fetchone()[0]
                
                cursor.execute(f'SELECT COUNT(*) FROM recruiter_handbooks{prev_handbook_where_clause}', prev_handbook_params)
                prev_total_handbooks = cursor.fetchone()[0]
                
                cursor.execute(f'SELECT COUNT(DISTINCT oorwin_job_id) FROM evaluations{prev_eval_where_clause} AND oorwin_job_id IS NOT NULL AND oorwin_job_id != ""', prev_eval_params)
                prev_unique_jobs_evals = cursor.fetchone()[0]
                
                cursor.execute(f'SELECT COUNT(DISTINCT oorwin_job_id) FROM recruiter_handbooks{prev_handbook_where_clause} AND oorwin_job_id IS NOT NULL AND oorwin_job_id != ""', prev_handbook_params)
                prev_unique_jobs_handbooks = cursor.fetchone()[0]
                
                prev_total_jobs = max(prev_unique_jobs_evals, prev_unique_jobs_handbooks)
                
                cursor.execute(f'SELECT AVG(match_percentage) FROM evaluations{prev_eval_where_clause} AND match_percentage IS NOT NULL', prev_eval_params)
                prev_avg_match_score = cursor.fetchone()[0] or 0
                prev_avg_match_score = round(prev_avg_match_score, 1)
            except Exception as e:
                logging.error(f"Error calculating previous period: {str(e)}")
        
        # Get conversion rate (jobs with handbook that also have evaluations) with filters
        # Build join query with filters
        join_where_parts = []
        join_params = []
        
        if accessible_emails:
            placeholders = ','.join(['?'] * len(accessible_emails))
            join_where_parts.append(f"h.user_email IN ({placeholders})")
            join_where_parts.append(f"e.user_email IN ({placeholders})")
            join_params.extend(accessible_emails)
            join_params.extend(accessible_emails)
        
        if filter_user and filter_user in accessible_emails:
            join_where_parts.append("h.user_email = ?")
            join_where_parts.append("e.user_email = ?")
            join_params.append(filter_user)
            join_params.append(filter_user)
        
        if date_from:
            join_where_parts.append("DATE(h.timestamp) >= ?")
            join_where_parts.append("DATE(e.timestamp) >= ?")
            join_params.append(date_from)
            join_params.append(date_from)
        
        if date_to:
            join_where_parts.append("DATE(h.timestamp) <= ?")
            join_where_parts.append("DATE(e.timestamp) <= ?")
            join_params.append(date_to)
            join_params.append(date_to)
        
        join_where_clause = " WHERE " + " AND ".join(join_where_parts) if join_where_parts else ""
        
        cursor.execute(f'''
            SELECT COUNT(DISTINCT h.oorwin_job_id) 
            FROM recruiter_handbooks h
            INNER JOIN evaluations e ON h.oorwin_job_id = e.oorwin_job_id
            {join_where_clause}
            AND h.oorwin_job_id IS NOT NULL AND h.oorwin_job_id != ""
        ''', join_params)
        jobs_with_both = cursor.fetchone()[0]
        conversion_rate = (jobs_with_both / total_handbooks * 100) if total_handbooks > 0 else 0
        
        # Get average evaluations per job with filters
        cursor.execute(f'''
            SELECT AVG(eval_count) FROM (
                SELECT oorwin_job_id, COUNT(*) as eval_count 
                FROM evaluations 
                {eval_where_clause}
                AND oorwin_job_id IS NOT NULL AND oorwin_job_id != ""
                GROUP BY oorwin_job_id
            )
        ''', eval_params)
        avg_evals_per_job = cursor.fetchone()[0] or 0
        
        # Get active jobs (activity in last 7 days) with filters
        active_where_parts = list(eval_where_parts)
        active_params = list(eval_params)
        active_where_parts.append("datetime(timestamp) >= datetime('now', '-7 days')")
        
        active_where_clause = " WHERE " + " AND ".join(active_where_parts) if active_where_parts else ""
        cursor.execute(f'''
            SELECT COUNT(DISTINCT oorwin_job_id) 
            FROM evaluations 
            {active_where_clause}
            AND oorwin_job_id IS NOT NULL 
            AND oorwin_job_id != ""
        ''', active_params)
        active_jobs_7d = cursor.fetchone()[0]
        
        # Get average eval time (calculate from time_taken column)
        eval_time_where = eval_where_clause + (" AND " if eval_where_clause else " WHERE ") + "time_taken IS NOT NULL AND time_taken > 0"
        cursor.execute(f'SELECT AVG(time_taken) FROM evaluations{eval_time_where}', eval_params)
        avg_eval_time_result = cursor.fetchone()[0]
        avg_eval_time = round(avg_eval_time_result, 1) if avg_eval_time_result else 0
        
        # Get average handbook generation time
        handbook_time_where = handbook_where_clause + (" AND " if handbook_where_clause else " WHERE ") + "time_taken IS NOT NULL AND time_taken > 0"
        cursor.execute(f'SELECT AVG(time_taken) FROM recruiter_handbooks{handbook_time_where}', handbook_params)
        avg_handbook_time_result = cursor.fetchone()[0]
        avg_handbook_time = round(avg_handbook_time_result, 1) if avg_handbook_time_result else 0
        
        # Get token usage metrics
        token_where = eval_where_clause + (" AND " if eval_where_clause else " WHERE ") + "token_usage IS NOT NULL AND token_usage > 0"
        cursor.execute(f'SELECT SUM(token_usage), AVG(token_usage), COUNT(*) FROM evaluations{token_where}', eval_params)
        token_result = cursor.fetchone()
        total_tokens_used = int(token_result[0]) if token_result[0] else 0
        avg_tokens_per_eval = round(token_result[1], 0) if token_result[1] else 0
        evals_with_tokens = token_result[2] if token_result[2] else 0
        
        # Get latest evaluation token usage
        latest_token_where = eval_where_clause + (" AND " if eval_where_clause else " WHERE ") + "token_usage IS NOT NULL AND token_usage > 0"
        cursor.execute(f'SELECT token_usage FROM evaluations{latest_token_where} ORDER BY timestamp DESC LIMIT 1', eval_params)
        latest_token_result = cursor.fetchone()
        latest_eval_tokens = int(latest_token_result[0]) if latest_token_result else None
        
        conn.close()
        
        # Calculate trends
        def calculate_trend(current, previous):
            if previous is None or previous == 0:
                return None, None
            change = current - previous
            percent_change = (change / previous) * 100 if previous > 0 else 0
            return change, round(percent_change, 1)
        
        eval_change, eval_percent = calculate_trend(total_evaluations, prev_total_evaluations)
        handbook_change, handbook_percent = calculate_trend(total_handbooks, prev_total_handbooks)
        jobs_change, jobs_percent = calculate_trend(total_jobs, prev_total_jobs)
        score_change, score_percent = calculate_trend(avg_match_score, prev_avg_match_score)
        
        return jsonify({
            'success': True,
            'metrics': {
                'total_evaluations': total_evaluations,
                'total_handbooks': total_handbooks,
                'total_jobs': total_jobs,
                'active_jobs': active_jobs_7d,
                'avg_match_score': round(avg_match_score, 1),
                'conversion_rate': round(conversion_rate, 1),
                'avg_evals_per_job': round(avg_evals_per_job, 1),
                'avg_eval_time': avg_eval_time,
                'avg_handbook_time': avg_handbook_time,
                'total_tokens_used': total_tokens_used,
                'tokens_per_eval': int(avg_tokens_per_eval),
                'latest_eval_tokens': latest_eval_tokens,
                'trends': {
                    'evaluations': {
                        'change': eval_change,
                        'percent': eval_percent
                    },
                    'handbooks': {
                        'change': handbook_change,
                        'percent': handbook_percent
                    },
                    'jobs': {
                        'change': jobs_change,
                        'percent': jobs_percent
                    },
                    'match_score': {
                        'change': score_change,
                        'percent': score_percent
                    }
                }
            }
        })
        
    except Exception as e:
        logging.error(f"Error fetching analytics overview: {str(e)}", exc_info=True)
        import traceback
        logging.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({
            'success': False,
            'message': str(e),
            'error': traceback.format_exc()
        }), 500

@app.route('/api/analytics/export-csv', methods=['GET'])
@login_required
def export_analytics_csv():
    """Export team/user data to CSV with evaluations and handbooks"""
    try:
        import csv
        from io import StringIO
        
        user_email = session['user'].get('email')
        filter_user = request.args.get('user_email', '')
        filter_team = request.args.get('team', '')
        date_from = request.args.get('date_from', '')
        date_to = request.args.get('date_to', '')
        
        # Get accessible user emails based on role
        accessible_emails = get_accessible_user_emails(user_email)
        
        # Debug logging
        logging.info(f"CSV Export - User: {user_email}, Accessible emails: {len(accessible_emails) if accessible_emails else 0}")
        logging.info(f"CSV Export - Filters: team={filter_team}, user={filter_user}, date_from={date_from}, date_to={date_to}")
        
        # Filter by team if specified
        if filter_team:
            conn_temp = sqlite3.connect(DATABASE_NAME)
            cursor_temp = conn_temp.cursor()
            cursor_temp.execute('SELECT email FROM users WHERE team = ?', (filter_team,))
            team_emails = [row[0] for row in cursor_temp.fetchall()]
            conn_temp.close()
            logging.info(f"CSV Export - Team emails: {len(team_emails)}")
            if accessible_emails:
                accessible_emails = [e for e in accessible_emails if e in team_emails]
                logging.info(f"CSV Export - After team filter: {len(accessible_emails)} accessible emails")
            else:
                # If no accessible users but team filter specified, use team emails
                accessible_emails = team_emails
                logging.info(f"CSV Export - Using team emails directly: {len(accessible_emails)}")
        
        # Filter by specific user if selected
        if filter_user:
            if accessible_emails and filter_user in accessible_emails:
                accessible_emails = [filter_user]
                logging.info(f"CSV Export - Filtered to specific user: {filter_user}")
            elif not accessible_emails:
                # If no accessible users restriction, allow if user exists
                conn_temp = sqlite3.connect(DATABASE_NAME)
                cursor_temp = conn_temp.cursor()
                cursor_temp.execute('SELECT email FROM users WHERE email = ?', (filter_user,))
                if cursor_temp.fetchone():
                    accessible_emails = [filter_user]
                    logging.info(f"CSV Export - Using filter_user directly: {filter_user}")
                conn_temp.close()
        
        logging.info(f"CSV Export - Final accessible emails count: {len(accessible_emails) if accessible_emails else 0}")
        
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Get user info for all accessible users (or all users if no restriction)
        if accessible_emails:
            user_placeholders = ','.join(['?'] * len(accessible_emails))
            cursor.execute(f'''
                SELECT email, name, role, team FROM users 
                WHERE email IN ({user_placeholders})
            ''', accessible_emails)
        else:
            # If no accessible users restriction, get all users (for Admin)
            cursor.execute('SELECT email, name, role, team FROM users')
        users_dict = {row[0]: {'name': row[1], 'role': row[2], 'team': row[3]} for row in cursor.fetchall()}
        
        # Build WHERE clause for evaluations and handbooks
        where_parts = []
        params = []
        
        if accessible_emails:
            placeholders = ','.join(['?'] * len(accessible_emails))
            where_parts.append(f"user_email IN ({placeholders})")
            params.extend(accessible_emails)
        else:
            # If no accessible users, still allow CSV export but with no data
            # This handles cases where user has no accessible data (e.g., new Recruiter)
            where_parts.append("1=0")  # Always false condition - returns no rows
        
        if date_from:
            where_parts.append("DATE(timestamp) >= ?")
            params.append(date_from)
        
        if date_to:
            where_parts.append("DATE(timestamp) <= ?")
            params.append(date_to)
        
        where_clause = " WHERE " + " AND ".join(where_parts) if where_parts else ""
        
        # Get evaluations
        cursor.execute(f'''
            SELECT user_email, filename, job_title, oorwin_job_id, 
                   match_percentage, timestamp
            FROM evaluations
            {where_clause}
            ORDER BY timestamp DESC, user_email
        ''', params)
        evaluations = cursor.fetchall()
        
        # Get handbooks
        cursor.execute(f'''
            SELECT user_email, job_title, oorwin_job_id, timestamp
            FROM recruiter_handbooks
            {where_clause}
            ORDER BY timestamp DESC, user_email
        ''', params)
        handbooks = cursor.fetchall()
        
        conn.close()
        
        # Create CSV content
        output = StringIO()
        writer = csv.writer(output)
        
        # Write header
        writer.writerow([
            'User Email', 'User Name', 'Role', 'Team', 'Date', 
            'Type', 'Job Title', 'Job ID', 'Match Score', 'Filename'
        ])
        
        # Write evaluation rows
        for eval_row in evaluations:
            user_email_eval, filename, job_title, job_id, match_score, timestamp = eval_row
            user_info = users_dict.get(user_email_eval, {'name': '', 'role': '', 'team': ''})
            writer.writerow([
                user_email_eval,
                user_info['name'],
                user_info['role'],
                user_info['team'],
                timestamp,
                'Evaluation',
                job_title or '',
                job_id or '',
                match_score or '',
                filename or ''
            ])
        
        # Write handbook rows
        for hb_row in handbooks:
            user_email_hb, job_title, job_id, timestamp = hb_row
            user_info = users_dict.get(user_email_hb, {'name': '', 'role': '', 'team': ''})
            writer.writerow([
                user_email_hb,
                user_info['name'],
                user_info['role'],
                user_info['team'],
                timestamp,
                'Handbook',
                job_title or '',
                job_id or '',
                '',
                ''
            ])
        
        # Create response
        output.seek(0)
        response = make_response(output.getvalue())
        response.headers['Content-Type'] = 'text/csv'
        response.headers['Content-Disposition'] = f'attachment; filename=team_data_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv'
        
        return response
        
    except Exception as e:
        logging.error(f"Error exporting CSV: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/analytics/timeline', methods=['GET'])
@login_required
def get_analytics_timeline():
    """API endpoint for activity timeline chart data"""
    try:
        days = int(request.args.get('days', 30))
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # Get daily evaluation counts
        cursor.execute(f'''
            SELECT DATE(timestamp) as date, COUNT(*) as count
            FROM evaluations
            WHERE datetime(timestamp) >= datetime('now', '-{days} days')
            GROUP BY DATE(timestamp)
            ORDER BY date
        ''')
        eval_data = cursor.fetchall()
        
        # Get daily handbook counts
        cursor.execute(f'''
            SELECT DATE(timestamp) as date, COUNT(*) as count
            FROM recruiter_handbooks
            WHERE datetime(timestamp) >= datetime('now', '-{days} days')
            GROUP BY DATE(timestamp)
            ORDER BY date
        ''')
        handbook_data = cursor.fetchall()
        
        conn.close()
        
        # Convert to dict for easy merging
        eval_dict = {row[0]: row[1] for row in eval_data}
        handbook_dict = {row[0]: row[1] for row in handbook_data}
        
        # Merge and create timeline
        from datetime import datetime, timedelta
        timeline = []
        start_date = datetime.now() - timedelta(days=days)
        
        for i in range(days + 1):
            current_date = (start_date + timedelta(days=i)).strftime('%Y-%m-%d')
            timeline.append({
                'date': current_date,
                'evaluations': eval_dict.get(current_date, 0),
                'handbooks': handbook_dict.get(current_date, 0)
            })
        
        return jsonify({
            'success': True,
            'timeline': timeline
        })
        
    except Exception as e:
        logging.error(f"Error fetching timeline data: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/analytics/team-performance', methods=['GET'])
@login_required
def get_team_performance():
    """API endpoint for team performance comparison"""
    try:
        user_email = session['user'].get('email')
        date_from = request.args.get('date_from', '')
        date_to = request.args.get('date_to', '')
        
        accessible_emails = get_accessible_user_emails(user_email)
        
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Get all teams from accessible users
        if accessible_emails:
            placeholders = ','.join(['?'] * len(accessible_emails))
            cursor.execute(f'''
                SELECT DISTINCT team FROM users 
                WHERE email IN ({placeholders}) AND team IS NOT NULL AND team != ""
            ''', accessible_emails)
            teams = [row[0] for row in cursor.fetchall()]
        else:
            teams = []
        
        team_performance = []
        for team in teams:
            # Get team members
            cursor.execute('SELECT email FROM users WHERE team = ?', (team,))
            team_emails = [row[0] for row in cursor.fetchall()]
            
            if not team_emails:
                continue
            
            # Build WHERE clause
            team_placeholders = ','.join(['?'] * len(team_emails))
            eval_where = [f"user_email IN ({team_placeholders})"]
            handbook_where = [f"user_email IN ({team_placeholders})"]
            eval_params = list(team_emails)
            handbook_params = list(team_emails)
            
            if date_from:
                eval_where.append("DATE(timestamp) >= ?")
                handbook_where.append("DATE(timestamp) >= ?")
                eval_params.append(date_from)
                handbook_params.append(date_from)
            
            if date_to:
                eval_where.append("DATE(timestamp) <= ?")
                handbook_where.append("DATE(timestamp) <= ?")
                eval_params.append(date_to)
                handbook_params.append(date_to)
            
            eval_where_clause = " WHERE " + " AND ".join(eval_where)
            handbook_where_clause = " WHERE " + " AND ".join(handbook_where)
            
            # Get counts
            cursor.execute(f'SELECT COUNT(*) FROM evaluations{eval_where_clause}', eval_params)
            eval_count = cursor.fetchone()[0]
            
            cursor.execute(f'SELECT COUNT(*) FROM recruiter_handbooks{handbook_where_clause}', handbook_params)
            handbook_count = cursor.fetchone()[0]
            
            cursor.execute(f'SELECT AVG(match_percentage) FROM evaluations{eval_where_clause} AND match_percentage IS NOT NULL', eval_params)
            avg_score = cursor.fetchone()[0] or 0
            
            team_performance.append({
                'team': team,
                'evaluations': eval_count,
                'handbooks': handbook_count,
                'avg_score': round(avg_score, 1)
            })
        
        conn.close()
        
        return jsonify({
            'success': True,
            'teams': team_performance
        })
        
    except Exception as e:
        logging.error(f"Error fetching team performance: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/analytics/match-score-distribution', methods=['GET'])
@login_required
def get_match_score_distribution():
    """API endpoint for match score distribution"""
    try:
        user_email = session['user'].get('email')
        filter_user = request.args.get('user_email', '')
        filter_team = request.args.get('team', '')
        date_from = request.args.get('date_from', '')
        date_to = request.args.get('date_to', '')
        
        accessible_emails = get_accessible_user_emails(user_email)
        
        # Filter by team if specified
        if filter_team:
            conn = sqlite3.connect(DATABASE_NAME)
            cursor = conn.cursor()
            cursor.execute('SELECT email FROM users WHERE team = ?', (filter_team,))
            team_emails = [row[0] for row in cursor.fetchall()]
            conn.close()
            accessible_emails = [e for e in accessible_emails if e in team_emails]
        
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # Build WHERE clause
        where_parts = []
        params = []
        
        if accessible_emails:
            placeholders = ','.join(['?'] * len(accessible_emails))
            where_parts.append(f"user_email IN ({placeholders})")
            params.extend(accessible_emails)
        
        if filter_user and filter_user in accessible_emails:
            where_parts.append("user_email = ?")
            params.append(filter_user)
        
        if date_from:
            where_parts.append("DATE(timestamp) >= ?")
            params.append(date_from)
        
        if date_to:
            where_parts.append("DATE(timestamp) <= ?")
            params.append(date_to)
        
        where_clause = " WHERE " + " AND ".join(where_parts) if where_parts else ""
        
        # Get score distribution (buckets: 0-20, 21-40, 41-60, 61-80, 81-100)
        cursor.execute(f'''
            SELECT 
                CASE 
                    WHEN match_percentage <= 20 THEN '0-20'
                    WHEN match_percentage <= 40 THEN '21-40'
                    WHEN match_percentage <= 60 THEN '41-60'
                    WHEN match_percentage <= 80 THEN '61-80'
                    ELSE '81-100'
                END as score_range,
                COUNT(*) as count
            FROM evaluations
            {where_clause}
            AND match_percentage IS NOT NULL
            GROUP BY score_range
            ORDER BY score_range
        ''', params)
        
        distribution = cursor.fetchall()
        conn.close()
        
        # Format for chart
        score_ranges = ['0-20', '21-40', '41-60', '61-80', '81-100']
        distribution_dict = {row[0]: row[1] for row in distribution}
        
        result = [{'range': r, 'count': distribution_dict.get(r, 0)} for r in score_ranges]
        
        return jsonify({
            'success': True,
            'distribution': result
        })
        
    except Exception as e:
        logging.error(f"Error fetching match score distribution: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/analytics/user-activity', methods=['GET'])
@login_required
def get_user_activity():
    """API endpoint for user activity breakdown"""
    try:
        user_email = session['user'].get('email')
        filter_team = request.args.get('team', '')
        date_from = request.args.get('date_from', '')
        date_to = request.args.get('date_to', '')
        
        accessible_emails = get_accessible_user_emails(user_email)
        
        # Filter by team if specified
        if filter_team:
            conn = sqlite3.connect(DATABASE_NAME)
            cursor = conn.cursor()
            cursor.execute('SELECT email FROM users WHERE team = ?', (filter_team,))
            team_emails = [row[0] for row in cursor.fetchall()]
            conn.close()
            accessible_emails = [e for e in accessible_emails if e in team_emails]
        
        if not accessible_emails:
            return jsonify({
                'success': True,
                'users': []
            })
        
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Get user names
        placeholders = ','.join(['?'] * len(accessible_emails))
        cursor.execute(f'''
            SELECT email, name FROM users
            WHERE email IN ({placeholders})
        ''', accessible_emails)
        users_dict = {row[0]: row[1] or row[0] for row in cursor.fetchall()}
        conn.close()
        
        # Get activity counts
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        eval_where_parts = [f"user_email IN ({placeholders})"]
        handbook_where_parts = [f"user_email IN ({placeholders})"]
        eval_params = list(accessible_emails)
        handbook_params = list(accessible_emails)
        
        if date_from:
            eval_where_parts.append("DATE(timestamp) >= ?")
            handbook_where_parts.append("DATE(timestamp) >= ?")
            eval_params.append(date_from)
            handbook_params.append(date_from)
        
        if date_to:
            eval_where_parts.append("DATE(timestamp) <= ?")
            handbook_where_parts.append("DATE(timestamp) <= ?")
            eval_params.append(date_to)
            handbook_params.append(date_to)
        
        eval_where_clause = " WHERE " + " AND ".join(eval_where_parts)
        handbook_where_clause = " WHERE " + " AND ".join(handbook_where_parts)
        
        # Get evaluation counts per user
        cursor.execute(f'''
            SELECT user_email, COUNT(*) as count
            FROM evaluations
            {eval_where_clause}
            GROUP BY user_email
        ''', eval_params)
        eval_counts = {row[0]: row[1] for row in cursor.fetchall()}
        
        # Get handbook counts per user
        cursor.execute(f'''
            SELECT user_email, COUNT(*) as count
            FROM recruiter_handbooks
            {handbook_where_clause}
            GROUP BY user_email
        ''', handbook_params)
        handbook_counts = {row[0]: row[1] for row in cursor.fetchall()}
        
        conn.close()
        
        # Combine data
        user_activity = []
        for email in accessible_emails:
            user_activity.append({
                'email': email,
                'name': users_dict.get(email, email),
                'evaluations': eval_counts.get(email, 0),
                'handbooks': handbook_counts.get(email, 0),
                'total': eval_counts.get(email, 0) + handbook_counts.get(email, 0)
            })
        
        # Sort by total activity
        user_activity.sort(key=lambda x: x['total'], reverse=True)
        
        return jsonify({
            'success': True,
            'users': user_activity
        })
        
    except Exception as e:
        logging.error(f"Error fetching user activity: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/analytics/recent-activity', methods=['GET'])
@login_required
def get_recent_activity():
    """API endpoint for recent activity feed"""
    try:
        user_email = session['user'].get('email')
        limit = int(request.args.get('limit', 20))
        filter_team = request.args.get('team', '')
        filter_user = request.args.get('user_email', '')
        date_from = request.args.get('date_from', '')
        date_to = request.args.get('date_to', '')
        
        accessible_emails = get_accessible_user_emails(user_email)
        
        # Filter by team if specified
        if filter_team:
            conn = sqlite3.connect(DATABASE_NAME)
            cursor = conn.cursor()
            cursor.execute('SELECT email FROM users WHERE team = ?', (filter_team,))
            team_emails = [row[0] for row in cursor.fetchall()]
            conn.close()
            accessible_emails = [e for e in accessible_emails if e in team_emails]
        
        if not accessible_emails:
            return jsonify({
                'success': True,
                'activities': []
            })
        
        # Get user names
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        placeholders = ','.join(['?'] * len(accessible_emails))
        cursor.execute(f'''
            SELECT email, name FROM users
            WHERE email IN ({placeholders})
        ''', accessible_emails)
        users_dict = {row[0]: row[1] or row[0] for row in cursor.fetchall()}
        conn.close()
        
        # Build WHERE clauses
        eval_where_parts = [f"user_email IN ({placeholders})"]
        handbook_where_parts = [f"user_email IN ({placeholders})"]
        eval_params = list(accessible_emails)
        handbook_params = list(accessible_emails)
        
        if filter_user and filter_user in accessible_emails:
            eval_where_parts.append("user_email = ?")
            handbook_where_parts.append("user_email = ?")
            eval_params.append(filter_user)
            handbook_params.append(filter_user)
        
        if date_from:
            eval_where_parts.append("DATE(timestamp) >= ?")
            handbook_where_parts.append("DATE(timestamp) >= ?")
            eval_params.append(date_from)
            handbook_params.append(date_from)
        
        if date_to:
            eval_where_parts.append("DATE(timestamp) <= ?")
            handbook_where_parts.append("DATE(timestamp) <= ?")
            eval_params.append(date_to)
            handbook_params.append(date_to)
        
        eval_where_clause = " WHERE " + " AND ".join(eval_where_parts)
        handbook_where_clause = " WHERE " + " AND ".join(handbook_where_parts)
        
        # Get recent evaluations
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        cursor.execute(f'''
            SELECT 
                'evaluation' as type,
                user_email,
                filename,
                job_title,
                oorwin_job_id,
                match_percentage,
                timestamp
            FROM evaluations
            {eval_where_clause}
            ORDER BY timestamp DESC
            LIMIT {limit}
        ''', eval_params)
        
        evaluations = cursor.fetchall()
        
        # Get recent handbooks
        cursor.execute(f'''
            SELECT 
                'handbook' as type,
                user_email,
                NULL as filename,
                job_title,
                oorwin_job_id,
                NULL as match_percentage,
                timestamp
            FROM recruiter_handbooks
            {handbook_where_clause}
            ORDER BY timestamp DESC
            LIMIT {limit}
        ''', handbook_params)
        
        handbooks = cursor.fetchall()
        conn.close()
        
        # Combine and sort by timestamp
        activities = []
        for row in evaluations:
            activities.append({
                'type': 'evaluation',
                'user_email': row[1],
                'user_name': users_dict.get(row[1], row[1]),
                'filename': row[2],
                'job_title': row[3],
                'job_id': row[4],
                'match_percentage': row[5],
                'timestamp': row[6]
            })
        
        for row in handbooks:
            activities.append({
                'type': 'handbook',
                'user_email': row[1],
                'user_name': users_dict.get(row[1], row[1]),
                'filename': None,
                'job_title': row[3],
                'job_id': row[4],
                'match_percentage': None,
                'timestamp': row[6]
            })
        
        # Sort by timestamp descending
        activities.sort(key=lambda x: x['timestamp'], reverse=True)
        
        # Return top N
        activities = activities[:limit]
        
        return jsonify({
            'success': True,
            'activities': activities
        })
        
    except Exception as e:
        logging.error(f"Error fetching recent activity: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/analytics/top-jobs', methods=['GET'])
@login_required
def get_top_jobs():
    """API endpoint for top jobs by activity (filtered by team)"""
    try:
        user_email = session['user'].get('email')
        limit = int(request.args.get('limit', 10))
        filter_user = request.args.get('user_email', '')
        filter_team = request.args.get('team', '')
        date_from = request.args.get('date_from', '')
        date_to = request.args.get('date_to', '')
        
        accessible_emails = get_accessible_user_emails(user_email)
        
        # Filter by team if specified
        if filter_team:
            conn = sqlite3.connect(DATABASE_NAME)
            cursor = conn.cursor()
            cursor.execute('SELECT email FROM users WHERE team = ?', (filter_team,))
            team_emails = [row[0] for row in cursor.fetchall()]
            conn.close()
            accessible_emails = [e for e in accessible_emails if e in team_emails]
        
        if not accessible_emails:
            return jsonify({
                'success': True,
                'jobs': []
            })
        
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # Build WHERE clause
        where_parts = ["e.oorwin_job_id IS NOT NULL AND e.oorwin_job_id != ''"]
        params = []
        
        # Filter by accessible users
        placeholders = ','.join(['?'] * len(accessible_emails))
        where_parts.append(f"e.user_email IN ({placeholders})")
        params.extend(accessible_emails)
        
        # Filter by specific user if selected
        if filter_user and filter_user in accessible_emails:
            where_parts.append("e.user_email = ?")
            params.append(filter_user)
        
        # Filter by date range
        if date_from:
            where_parts.append("DATE(e.timestamp) >= ?")
            params.append(date_from)
        
        if date_to:
            where_parts.append("DATE(e.timestamp) <= ?")
            params.append(date_to)
        
        where_clause = " WHERE " + " AND ".join(where_parts)
        
        cursor.execute(f'''
            SELECT 
                e.oorwin_job_id,
                e.job_title,
                COUNT(*) as eval_count,
                AVG(e.match_percentage) as avg_score,
                MAX(e.timestamp) as last_active
            FROM evaluations e
            {where_clause}
            GROUP BY e.oorwin_job_id, e.job_title
            ORDER BY eval_count DESC
            LIMIT {limit}
        ''', params)
        
        rows = cursor.fetchall()
        conn.close()
        
        jobs = []
        for row in rows:
            jobs.append({
                'job_id': row[0],
                'job_title': row[1],
                'eval_count': row[2],
                'avg_score': round(row[3], 1) if row[3] else 0,
                'last_active': row[4]
            })
        
        return jsonify({
            'success': True,
            'jobs': jobs
        })
        
    except Exception as e:
        logging.error(f"Error fetching top jobs: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/handbook/<int:handbook_id>', methods=['GET'])
def get_single_handbook(handbook_id):
    """API endpoint to get a single handbook by ID"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT id, job_title, oorwin_job_id, job_description, timestamp, markdown_content
            FROM recruiter_handbooks
            WHERE id = ?
        ''', (handbook_id,))
        
        row = cursor.fetchone()
        conn.close()
        
        if not row:
            return jsonify({
                'success': False,
                'message': 'Handbook not found'
            }), 404
        
        handbook = {
            'id': row[0],
            'job_title': row[1],
            'oorwin_job_id': row[2],
            'job_description': row[3],
            'timestamp': row[4],
            'markdown_content': row[5]
        }
        
        return jsonify({
            'success': True,
            'handbook': handbook
        })
        
    except Exception as e:
        logging.error(f"Error fetching handbook {handbook_id}: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/handbooks-only', methods=['GET'])
def get_handbooks_only():
    """API endpoint for handbooks-only history"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT id, job_title, oorwin_job_id, timestamp, markdown_content
            FROM recruiter_handbooks
            ORDER BY timestamp DESC
        ''')
        
        handbooks = []
        for row in cursor.fetchall():
            handbooks.append({
                'id': row[0],
                'job_title': row[1],
                'oorwin_job_id': row[2],
                'timestamp': row[3],
                'markdown_content': row[4][:500] + '...' if row[4] and len(row[4]) > 500 else row[4]  # Truncate for preview
            })
        
        conn.close()
        
        return jsonify({
            'success': True,
            'handbooks': handbooks
        })
        
    except Exception as e:
        logging.error(f"Error fetching handbooks-only history: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/evaluations-only', methods=['GET'])
def get_evaluations_only():
    """API endpoint for evaluations-only history"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT id, filename, job_title, oorwin_job_id, match_percentage, timestamp
            FROM evaluations
            ORDER BY timestamp DESC
        ''')
        
        evaluations = []
        for row in cursor.fetchall():
            evaluations.append({
                'id': row[0],
                'filename': row[1],
                'job_title': row[2],
                'oorwin_job_id': row[3] if row[3] else 'N/A',
                'match_percentage': row[4],
                'timestamp': row[5]
            })
        
        conn.close()
        
        return jsonify({
            'success': True,
            'evaluations': evaluations
        })
        
    except Exception as e:
        logging.error(f"Error fetching evaluations-only history: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/job-centric-history', methods=['GET'])
@login_required
def get_job_centric_history():
    """API endpoint for job-centric grouped history (filtered by role)"""
    try:
        # Get accessible user emails based on role
        user_email = session['user'].get('email')
        accessible_emails = get_accessible_user_emails(user_email)
        
        # Debug logging
        logging.info(f"Job-centric history API - Current user: {user_email}, Accessible emails: {accessible_emails}")
        
        if not accessible_emails:
            return jsonify({
                'success': True,
                'jobs': []
            })
        
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Build placeholders for SQL queries
        placeholders = ','.join(['?'] * len(accessible_emails))
        
        # Get all unique job IDs from both tables, ordered by most recent activity (filtered by accessible users)
        # Also include evaluations/handbooks with NULL user_email for backward compatibility
        cursor.execute(f'''
            SELECT oorwin_job_id, MAX(timestamp) as last_activity
            FROM (
                SELECT oorwin_job_id, timestamp FROM evaluations 
                WHERE oorwin_job_id IS NOT NULL AND oorwin_job_id != ""
                AND (user_email IN ({placeholders}) OR user_email IS NULL)
                UNION ALL
                SELECT oorwin_job_id, timestamp FROM recruiter_handbooks 
                WHERE oorwin_job_id IS NOT NULL AND oorwin_job_id != ""
                AND (user_email IN ({placeholders}) OR user_email IS NULL)
            )
            GROUP BY oorwin_job_id
            ORDER BY last_activity DESC
        ''', accessible_emails + accessible_emails)
        
        job_ids = [row[0] for row in cursor.fetchall()]
        
        result = []
        for job_id in job_ids:
            # Get job title (prefer from handbooks, fallback to evaluations) - filtered by accessible users
            # Also include handbooks/evaluations with NULL user_email for backward compatibility
            cursor.execute(f'''
                SELECT job_title FROM recruiter_handbooks 
                WHERE oorwin_job_id = ? AND job_title IS NOT NULL AND job_title != ""
                AND (user_email IN ({placeholders}) OR user_email IS NULL)
                ORDER BY timestamp DESC LIMIT 1
            ''', (job_id,) + tuple(accessible_emails))
            
            job_title_row = cursor.fetchone()
            job_title = job_title_row[0] if job_title_row else None
            
            if not job_title:
                cursor.execute(f'''
                    SELECT job_title FROM evaluations 
                    WHERE oorwin_job_id = ? AND job_title IS NOT NULL AND job_title != ""
                    AND (user_email IN ({placeholders}) OR user_email IS NULL)
                    ORDER BY timestamp DESC LIMIT 1
                ''', (job_id,) + tuple(accessible_emails))
                job_title_row = cursor.fetchone()
                job_title = job_title_row[0] if job_title_row else "N/A"
            
            # Count handbooks (filtered by team members only - team-based filtering)
            # Also include handbooks with NULL user_email for backward compatibility
            cursor.execute(f'''
                SELECT COUNT(*) FROM recruiter_handbooks 
                WHERE oorwin_job_id = ? AND (user_email IN ({placeholders}) OR user_email IS NULL)
            ''', (job_id,) + tuple(accessible_emails))
            handbooks_count = cursor.fetchone()[0]
            
            # Get handbook creators (unique users who generated handbooks) - filtered by accessible users
            cursor.execute(f'''
                SELECT DISTINCT user_email FROM recruiter_handbooks 
                WHERE oorwin_job_id = ? AND user_email IS NOT NULL AND user_email != ""
                AND user_email IN ({placeholders})
            ''', (job_id,) + tuple(accessible_emails))
            hb_creator_emails = [row[0] for row in cursor.fetchall()]
            
            # Get user names for handbook creators
            hb_generated_by = []
            if hb_creator_emails:
                hb_placeholders = ','.join(['?'] * len(hb_creator_emails))
                cursor.execute(f'''
                    SELECT email, name FROM users 
                    WHERE email IN ({hb_placeholders})
                ''', hb_creator_emails)
                for row in cursor.fetchall():
                    hb_generated_by.append({
                        'email': row[0],
                        'name': row[1] or row[0]
                    })
            
            # Get evaluated resumes (filenames) - filtered by team members only (team-based filtering)
            # Also include evaluations with NULL user_email for backward compatibility
            cursor.execute(f'''
                SELECT filename, timestamp, match_percentage, user_email FROM evaluations 
                WHERE oorwin_job_id = ? AND (user_email IN ({placeholders}) OR user_email IS NULL)
                ORDER BY timestamp DESC
            ''', (job_id,) + tuple(accessible_emails))
            evaluations = cursor.fetchall()
            evaluations_count = len(evaluations)
            
            # Debug logging
            logging.info(f"Job {job_id} - Found {evaluations_count} evaluations (accessible emails: {accessible_emails})")
            eval_user_emails = set()  # Track unique evaluators
            for eval_row in evaluations:
                if eval_row[3]:  # user_email (not NULL/empty)
                    eval_user_emails.add(eval_row[3])
            
            # Get user names for resume evaluators
            eval_user_names = {}
            if eval_user_emails:
                eval_email_list = list(eval_user_emails)
                eval_placeholders = ','.join(['?'] * len(eval_email_list))
                cursor.execute(f'''
                    SELECT email, name FROM users 
                    WHERE email IN ({eval_placeholders})
                ''', eval_email_list)
                for row in cursor.fetchall():
                    email, name = row[0], row[1] or row[0]
                    eval_user_names[email] = name
            
            # Build resume list including evaluator information per resume
            resume_list = []
            res_evaluated_by = []
            for eval_row in evaluations:
                evaluator_email = eval_row[3]
                evaluator_name = eval_user_names.get(evaluator_email, evaluator_email) if evaluator_email else None
                resume_list.append({
                    'filename': eval_row[0],
                    'timestamp': eval_row[1],
                    'match_percentage': eval_row[2],
                    'evaluator_email': evaluator_email,
                    'evaluator_name': evaluator_name
                })
                
                # Build one "Res Evaluated By" entry per evaluation (even if the same user repeats)
                if evaluator_email:
                    res_evaluated_by.append({
                        'email': evaluator_email,
                        'name': evaluator_name or evaluator_email
                    })
                else:
                    res_evaluated_by.append({
                        'email': None,
                        'name': 'Unknown'
                    })
            
            # Get first created date (earliest timestamp from both tables) - filtered by accessible users
            # Also include evaluations/handbooks with NULL user_email for backward compatibility
            cursor.execute(f'''
                SELECT MIN(timestamp) FROM (
                    SELECT timestamp FROM evaluations 
                    WHERE oorwin_job_id = ? AND (user_email IN ({placeholders}) OR user_email IS NULL)
                    UNION ALL
                    SELECT timestamp FROM recruiter_handbooks 
                    WHERE oorwin_job_id = ? AND (user_email IN ({placeholders}) OR user_email IS NULL)
                )
            ''', (job_id,) + tuple(accessible_emails) + (job_id,) + tuple(accessible_emails))
            first_created = cursor.fetchone()[0]
            
            # Get last activity date (latest timestamp from both tables) - filtered by accessible users
            # Also include evaluations/handbooks with NULL user_email for backward compatibility
            cursor.execute(f'''
                SELECT MAX(timestamp) FROM (
                    SELECT timestamp FROM evaluations 
                    WHERE oorwin_job_id = ? AND (user_email IN ({placeholders}) OR user_email IS NULL)
                    UNION ALL
                    SELECT timestamp FROM recruiter_handbooks 
                    WHERE oorwin_job_id = ? AND (user_email IN ({placeholders}) OR user_email IS NULL)
                )
            ''', (job_id,) + tuple(accessible_emails) + (job_id,) + tuple(accessible_emails))
            last_activity = cursor.fetchone()[0]
            
            result.append({
                'job_id': job_id,
                'job_title': job_title,
                'handbooks_count': handbooks_count,
                'hb_generated_by': hb_generated_by,
                'evaluations_count': evaluations_count,
                'resume_list': resume_list,
                'res_evaluated_by': res_evaluated_by,
                'first_created': first_created,
                'last_activity': last_activity
            })
        
        conn.close()
        
        return jsonify({
            'success': True,
            'jobs': result
        })
        
    except Exception as e:
        logging.error(f"Error fetching job-centric history: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/handbooks-by-job/<job_id>', methods=['GET'])
@login_required
def get_handbooks_by_job(job_id):
    """API endpoint to get all handbooks for a specific job ID (filtered by role)"""
    try:
        # Get accessible user emails based on role
        user_email = session['user'].get('email')
        accessible_emails = get_accessible_user_emails(user_email)
        
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        if not accessible_emails:
            return jsonify({
                'success': True,
                'handbooks': []
            })
        
        # Build placeholders for SQL query
        placeholders = ','.join(['?'] * len(accessible_emails))
        
        cursor.execute(f'''
            SELECT id, job_title, job_description, additional_context, markdown_content, timestamp, user_email
            FROM recruiter_handbooks
            WHERE oorwin_job_id = ? AND user_email IN ({placeholders})
            ORDER BY timestamp DESC
        ''', (job_id,) + tuple(accessible_emails))
        
        rows = cursor.fetchall()
        
        # Get unique user emails from handbooks
        user_emails = set()
        for row in rows:
            if row[6]:  # user_email
                user_emails.add(row[6])
        
        # Fetch user names
        user_names = {}
        if user_emails:
            email_list = list(user_emails)
            user_placeholders = ','.join(['?'] * len(email_list))
            cursor.execute(f'''
                SELECT email, name FROM users
                WHERE email IN ({user_placeholders})
            ''', email_list)
            for email, name in cursor.fetchall():
                user_names[email] = name or email
        
        conn.close()
        
        handbooks = []
        for row in rows:
            user_email = row[6]
            user_name = user_names.get(user_email, user_email) if user_email else 'Unknown'
            
            handbooks.append({
                'id': row[0],
                'job_title': row[1],
                'job_description': row[2],
                'additional_context': row[3],
                'markdown_content': row[4],
                'timestamp': row[5],
                'user_email': user_email,
                'user_name': user_name
            })
        
        return jsonify({
            'success': True,
            'handbooks': handbooks
        })
        
    except Exception as e:
        logging.error(f"Error fetching handbooks for job {job_id}: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/download-evaluation-pdf', methods=['POST'])
@login_required
def download_evaluation_pdf():
    """API endpoint to download resume evaluation as PDF"""
    try:
        data = request.get_json()
        evaluation_data = data.get('evaluation_data', {})
        
        if not evaluation_data:
            return jsonify({
                'success': False,
                'message': 'No evaluation data to download'
            }), 400
        
        logging.info("Generating PDF from evaluation data (PeopleLogic layout)...")
        pdf_data = build_evaluation_pdf_bytes(evaluation_data)
        filename = evaluation_data.get('filename', 'Unknown')
        
        logging.info("PDF generated successfully")
        
        # Return PDF as response
        response = make_response(pdf_data)
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = f'attachment; filename=Resume_Evaluation_{filename}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
        
        return response
        
    except Exception as e:
        logging.error(f"Error generating evaluation PDF: {str(e)}", exc_info=True)
        return jsonify({
            'success': False,
            'message': f'Failed to generate PDF: {str(e)}'
        }), 500


@app.route('/api/download-resume/<int:eval_id>', methods=['GET'])
@login_required
def download_resume(eval_id):
    """Download the original resume file associated with a given evaluation."""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        cursor.execute(
            '''
            SELECT filename, resume_path
            FROM evaluations
            WHERE id = ?
            ''',
            (eval_id,),
        )
        row = cursor.fetchone()
        conn.close()

        if not row:
            return jsonify({'success': False, 'message': 'Evaluation not found'}), 404

        filename, resume_path_from_db = row[0], row[1]

        # We now store only filename in resume_path. Normalize to filename and rebuild full path.
        resume_filename = None
        if resume_path_from_db:
            resume_filename = os.path.basename(str(resume_path_from_db).strip().strip('"').strip("'"))

        if not resume_filename:
            # Backward compatibility: fall back to filename column
            if not filename:
                return jsonify(
                    {
                        'success': False,
                        'message': 'Resume file path not available for this evaluation',
                    }
                ), 404
            resume_filename = os.path.basename(filename)

        resume_path = os.path.join(app.config['UPLOAD_FOLDER'], resume_filename)
        resume_path = os.path.normpath(resume_path)

        if not os.path.exists(resume_path):
            logging.error(f"Resume file not found for evaluation {eval_id}: {resume_path}")
            return jsonify(
                {
                    'success': False,
                    'message': f'Resume file not found on server: {os.path.basename(resume_path)}',
                }
            ), 404

        download_name = os.path.basename(resume_path) if resume_path else (filename or 'resume')
        return send_file(resume_path, as_attachment=True, download_name=download_name)

    except Exception as e:
        logging.error(f"Error downloading resume for evaluation {eval_id}: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'message': f'Failed to download resume: {str(e)}'}), 500

@app.route('/api/download-evaluation-with-resume', methods=['POST'])
@login_required
def download_evaluation_with_resume():
    """API endpoint to download resume evaluation merged with original resume PDF"""
    try:
        if not PDF_MERGE_AVAILABLE:
            return jsonify({
                'success': False,
                'message': 'PDF merging not available. Please install PyPDF2.'
            }), 500
        
        data = request.get_json()
        logging.info(f"[MERGE DEBUG] Incoming JSON: {data}")
        evaluation_data = data.get('evaluation_data', {}) or {}
        evaluation_id = data.get('evaluation_id')
        resume_path_from_request = data.get('resume_path')

        if not evaluation_id:
            return jsonify({
                'success': False,
                'message': 'No evaluation_id provided for merged download.'
            }), 400

        # Validate that evaluation_id is NOT a UUID (UUIDs contain hyphens)
        # The database uses integer IDs, not UUIDs
        if isinstance(evaluation_id, str) and '-' in str(evaluation_id):
            logging.error(f"[MERGE_DOWNLOAD_ERROR] Received UUID instead of DB ID: {evaluation_id}")
            return jsonify({
                'success': False,
                'message': f'Invalid evaluation ID format. Received UUID "{evaluation_id}" but database uses integer IDs. '
                          f'Please wait for evaluation to complete and try again, or refresh the page.'
            }), 400

        # Convert to integer if it's a numeric string
        try:
            evaluation_id = int(evaluation_id)
        except (ValueError, TypeError):
            logging.error(f"[MERGE_DOWNLOAD_ERROR] Invalid evaluation_id type: {type(evaluation_id)}, value: {evaluation_id}")
            return jsonify({
                'success': False,
                'message': f'Invalid evaluation ID format. Expected integer, got: {type(evaluation_id).__name__}'
            }), 400

        # Look up resume location from the database based on evaluation_id
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        cursor.execute(
            "SELECT filename, resume_path FROM evaluations WHERE id = ?",
            (evaluation_id,)
        )
        row = cursor.fetchone()
        conn.close()

        if not row:
            logging.error(f"[MERGE_DOWNLOAD_ERROR] Evaluation {evaluation_id} not found in database.")
            return jsonify({
                'success': False,
                'message': f'Evaluation {evaluation_id} not found in database.'
            }), 404

        filename_from_db, resume_path_from_db = row

        # Normalize stored value to a filename (basename)
        # Prefer resume_path from request if provided (frontend sends it), otherwise use DB value
        resume_filename = None
        if resume_path_from_request:
            # Frontend sent resume_path, use it (but validate it exists)
            resume_filename = os.path.basename(str(resume_path_from_request).strip().strip('"').strip("'"))
            logging.info(f"[MERGE DEBUG] Using resume_path from request: {resume_filename}")
        elif resume_path_from_db:
            resume_filename = os.path.basename(str(resume_path_from_db).strip().strip('"').strip("'"))
            logging.info(f"[MERGE DEBUG] Using resume_path from DB: {resume_filename}")
        elif filename_from_db:
            resume_filename = os.path.basename(str(filename_from_db).strip().strip('"').strip("'"))
            logging.info(f"[MERGE DEBUG] Using filename from DB: {resume_filename}")

        if not resume_filename:
            logging.error(
                f"[MERGE_DOWNLOAD_ERROR] No resume filename for evaluation {evaluation_id}. "
                f"DB resume_path={resume_path_from_db}, filename={filename_from_db}"
            )
            return jsonify({
                'success': False,
                'message': 'No resume file reference is stored for this evaluation in the database.'
            }), 404

        # Construct full path to resume file in uploads folder
        # app.config['UPLOAD_FOLDER'] is set to: os.path.join(app.root_path, "uploads")
        uploads_folder = app.config.get('UPLOAD_FOLDER', os.path.join(app.root_path, "uploads"))
        full_resume_path = os.path.join(uploads_folder, resume_filename)
        full_resume_path = os.path.normpath(full_resume_path)
        
        # Log the path resolution for debugging
        logging.info(f"[MERGE DEBUG] Uploads folder: {uploads_folder}")
        logging.info(f"[MERGE DEBUG] Resume filename: {resume_filename}")
        logging.info(f"[MERGE DEBUG] Full resume path: {full_resume_path}")
        logging.info(f"[MERGE DEBUG] Path exists: {os.path.exists(full_resume_path)}")

        # Check if resume file exists
        if not os.path.exists(full_resume_path):
            logging.error(
                "[MERGE_DOWNLOAD_ERROR] Resume file not found on disk. "
                f"Expected at: {full_resume_path}"
            )
            logging.error(f"[MERGE_DOWNLOAD_ERROR] Uploads folder contents: {os.listdir(uploads_folder) if os.path.exists(uploads_folder) else 'Folder does not exist'}")
            return jsonify({
                'success': False,
                'message': f'Resume file not found on server at expected location. '
                           f'Expected file name: \"{resume_filename}\" inside the uploads folder at: {uploads_folder}'
            }), 404
        
        # Check if resume is a PDF (only PDFs can be merged)
        if not resume_filename.lower().endswith('.pdf'):
            logging.error(
                "[MERGE_DOWNLOAD_ERROR] Resume is not a PDF. "
                f"Got extension for file \"{resume_filename}\""
            )
            return jsonify({
                'success': False,
                'message': 'Only PDF resumes can be merged. The stored resume is not a PDF. '
                           'Please upload a PDF resume and run a new evaluation.'
            }), 400
        
        logging.info(f"Generating merged PDF: evaluation + resume from {full_resume_path}")
        
        # Same full PeopleLogic evaluation layout as standalone PDF download
        # Concise 1–2 page evaluation (no interview questions), then resume PDF
        evaluation_pdf_bytes = build_evaluation_pdf_bytes(evaluation_data, concise=True)
        
        # Merge evaluation PDF with resume PDF
        merger = PdfMerger()
        
        # Add evaluation PDF
        evaluation_pdf_io = BytesIO(evaluation_pdf_bytes)
        merger.append(evaluation_pdf_io)
        
        # Add original resume PDF
        with open(full_resume_path, 'rb') as resume_file:
            merger.append(resume_file)
        
        # Create merged PDF in memory
        merged_pdf_buffer = BytesIO()
        merger.write(merged_pdf_buffer)
        merger.close()
        evaluation_pdf_io.close()
        
        # Get merged PDF data
        merged_pdf_data = merged_pdf_buffer.getvalue()
        merged_pdf_buffer.close()
        
        logging.info("Merged PDF generated successfully")
        
        # Return merged PDF as response
        filename = evaluation_data.get('filename', 'Resume_With_Evaluation')
        filename_base = os.path.splitext(filename)[0] if filename else 'Resume_With_Evaluation'
        
        response = make_response(merged_pdf_data)
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = f'attachment; filename=Resume_With_Evaluation_{filename_base}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
        
        return response
        
    except Exception as e:
        logging.error(f"Error generating merged PDF: {str(e)}", exc_info=True)
        return jsonify({
            'success': False,
            'message': f'Failed to generate merged PDF: {str(e)}'
        }), 500

@app.route('/api/evaluation-full/<int:eval_id>', methods=['GET'])
def get_evaluation_full(eval_id):
    """API endpoint to get full evaluation data for viewing"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Fetch evaluation data
        cursor.execute('''
            SELECT 
                e.id, e.filename, e.job_title, e.job_description,
                e.match_percentage, e.match_factors, e.profile_summary,
                e.missing_keywords, e.job_stability, e.career_progression,
                e.oorwin_job_id, e.timestamp, e.candidate_fit_analysis, e.over_under_qualification, e.resume_path
            FROM evaluations e
            WHERE e.id = ?
        ''', (eval_id,))
        
        row = cursor.fetchone()
        
        if not row:
            conn.close()
            return jsonify({
                'success': False,
                'message': 'Evaluation not found'
            }), 404
        
        # Fetch interview questions
        cursor.execute('''
            SELECT technical_questions, nontechnical_questions, behavioral_questions
            FROM interview_questions
            WHERE evaluation_id = ?
        ''', (eval_id,))
        
        questions_row = cursor.fetchone()
        conn.close()
        
        # Parse JSON fields
        import json
        resume_path_from_db = row[14] if len(row) > 14 else None

        # Temporary debug logging to inspect stored paths and existence
        logging.info(f"[DEBUG] Raw DB resume_path value: {resume_path_from_db}")

        resolved_resume_path = None

        # Normalize stored value to a filename (basename) when present
        if resume_path_from_db:
            filename_only = os.path.basename(str(resume_path_from_db).strip().strip('"').strip("'"))
        else:
            filename_only = None

        # If no resume_path stored, fall back to filename column
        if not filename_only:
            filename_col = row[1] if row[1] else ''
            filename_only = os.path.basename(filename_col) if filename_col else None

        if filename_only:
            # Reconstruct full absolute path from UPLOAD_FOLDER and filename
            candidate_path = os.path.join(app.config['UPLOAD_FOLDER'], filename_only)
            candidate_path = os.path.normpath(candidate_path)

            # Handle relative paths: if not absolute, make it relative to app root
            if not os.path.isabs(candidate_path):
                candidate_path = os.path.join(app.root_path, candidate_path)
                candidate_path = os.path.normpath(candidate_path)

            logging.info(f"[DEBUG] Reconstructed resume path candidate: {candidate_path}")
            logging.info(f"[DEBUG] Absolute resolved path: {os.path.abspath(candidate_path)}")
            logging.info(f"[DEBUG] Exists check: {os.path.exists(os.path.abspath(candidate_path))}")

            if os.path.exists(candidate_path):
                resolved_resume_path = candidate_path
                logging.info(f"[DEBUG] Resolved path: {resolved_resume_path}")
                logging.info(f"[DEBUG] Exists check: {os.path.exists(resolved_resume_path)}")
            else:
                logging.warning(f"Resume file not found at reconstructed path: {candidate_path}")

        # If we still don't have a valid path, keep it as None so frontend can handle gracefully
        
        evaluation = {
            'id': row[0],
            'filename': row[1],
            'job_title': row[2],
            'job_description': row[3],
            'match_percentage': row[4],
            'match_percentage_str': str(int(row[4])) + '%' if row[4] else '0%',
            'match_factors': json.loads(row[5]) if row[5] else {},
            'profile_summary': row[6],
            'missing_keywords': json.loads(row[7]) if row[7] else [],
            'job_stability': json.loads(row[8]) if row[8] else {},
            'career_progression': json.loads(row[9]) if row[9] else {},
            'oorwin_job_id': row[10],
            'timestamp': row[11],
            # Parse new fields from database (or use empty if not present)
            'candidate_fit_analysis': json.loads(row[12]) if (len(row) > 12 and row[12]) else {},
            'over_under_qualification': row[13] if (len(row) > 13 and row[13]) else '',
            # Expose resolved absolute path only if file exists; otherwise None
            'resume_path': resolved_resume_path
        }
        
        # Helper function to normalize questions (convert objects to strings)
        def normalize_questions(questions_list):
            if not questions_list:
                return []
            normalized = []
            for q in questions_list:
                if isinstance(q, str):
                    normalized.append(q)
                elif isinstance(q, dict):
                    # Extract question text from common property names
                    normalized.append(q.get('question') or q.get('text') or q.get('content') or q.get('value') or str(q))
                else:
                    normalized.append(str(q))
            return normalized
        
        if questions_row:
            tech_raw = json.loads(questions_row[0]) if questions_row[0] else []
            nontech_raw = json.loads(questions_row[1]) if questions_row[1] else []
            behavioral_raw = json.loads(questions_row[2]) if questions_row[2] else []
            
            evaluation['technical_questions'] = normalize_questions(tech_raw)
            evaluation['nontechnical_questions'] = normalize_questions(nontech_raw)
            evaluation['behavioral_questions'] = normalize_questions(behavioral_raw)
        else:
            evaluation['technical_questions'] = []
            evaluation['nontechnical_questions'] = []
            evaluation['behavioral_questions'] = []
        
        return jsonify({
            'success': True,
            'evaluation': evaluation
        })
        
    except Exception as e:
        logging.error(f"Error fetching full evaluation {eval_id}: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/evaluations-by-job/<job_id>', methods=['GET'])
@login_required
def get_evaluations_by_job(job_id):
    """API endpoint to get all evaluations for a specific job ID (filtered by team membership)"""
    try:
        # Get accessible user emails based on team membership
        user_email = session['user'].get('email')
        accessible_emails = get_accessible_user_emails(user_email)
        
        if not accessible_emails:
            return jsonify({
                'success': True,
                'evaluations': []
            })
        
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Build placeholders for SQL query
        placeholders = ','.join(['?'] * len(accessible_emails))
        
        # Get evaluations for this job (with evaluator email)
        cursor.execute(f'''
            SELECT id, filename, match_percentage, timestamp, user_email
            FROM evaluations
            WHERE oorwin_job_id = ? AND user_email IN ({placeholders})
            ORDER BY timestamp DESC
        ''', (job_id,) + tuple(accessible_emails))
        rows = cursor.fetchall()

        # Collect unique evaluator emails
        evaluator_emails = {row[4] for row in rows if row[4]}

        # Map emails to user names
        evaluator_names = {}
        if evaluator_emails:
            email_list = list(evaluator_emails)
            email_placeholders = ','.join(['?'] * len(email_list))
            cursor.execute(f'''
                SELECT email, name FROM users
                WHERE email IN ({email_placeholders})
            ''', email_list)
            for email, name in cursor.fetchall():
                evaluator_names[email] = name or email

        conn.close()
        
        evaluations = []
        for row in rows:
            email = row[4]
            # Normalize match_percentage for API consumers (avoid None in JS comparisons)
            try:
                mp_value = int(row[2]) if row[2] is not None else 0
            except (ValueError, TypeError):
                mp_value = 0
            evaluations.append({
                'id': row[0],
                'filename': row[1],
                'match_percentage': mp_value,
                'timestamp': row[3],
                'user_email': email,
                'evaluator_name': evaluator_names.get(email, email) if email else 'Unknown'
            })
        
        return jsonify({
            'success': True,
            'evaluations': evaluations
        })
        
    except Exception as e:
        logging.error(f"Error fetching evaluations for job {job_id}: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/download-handbook-pdf', methods=['POST'])
async def download_handbook_pdf():
    """API endpoint to download recruiter handbook as PDF"""
    try:
        data = request.get_json()
        markdown_content = data.get('markdown_content', '').strip()
        job_title = (data.get('job_title') or '').strip()
        oorwin_job_id = (data.get('oorwin_job_id') or '').strip()
        
        if not markdown_content:
            return jsonify({
                'success': False,
                'message': 'No content to download'
            }), 400
        
        logging.info("Generating PDF from handbook content (branded handbook_pdf)...")
        pdf_data = build_handbook_pdf_bytes(markdown_content, job_title, oorwin_job_id)
        
        logging.info("PDF generated successfully")
        
        # Return PDF as response
        response = make_response(pdf_data)
        response.headers['Content-Type'] = 'application/pdf'
        
        # Format filename as RH_JobTitle_JobID.pdf
        safe_title = re.sub(r'[^A-Za-z0-9 _\-]+', '', job_title) or 'Handbook'
        safe_job = re.sub(r'[^A-Za-z0-9 _\-]+', '', oorwin_job_id) if oorwin_job_id else ''
        
        # Build filename: RH_JobTitle_JobID.pdf
        filename_parts = ["RH", safe_title]
        if safe_job:
            filename_parts.append(safe_job)
        
        filename = "_".join(filename_parts) + '.pdf'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        logging.error(f"Error generating PDF: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500


def get_default_interview_questions(job_title):
    """Generate default interview questions based on job title"""
    # Default technical questions based on common job titles
    technical_questions = {
        "software": [
            "Describe your experience with different programming languages and frameworks.",
            "How do you approach debugging a complex issue in your code?",
            "Explain your understanding of object-oriented programming principles.",
            "How do you ensure code quality and maintainability?",
            "Describe a challenging technical problem you solved recently."
        ],
        "data": [
            "Explain the difference between supervised and unsupervised learning.",
            "How do you handle missing or inconsistent data in your analysis?",
            "Describe your experience with SQL and database optimization.",
            "What tools and libraries do you use for data visualization?",
            "How do you validate the results of your data analysis?"
        ],
        "manager": [
            "How do you approach resource allocation in a project?",
            "Describe your experience with agile methodologies.",
            "How do you handle conflicts within your team?",
            "What metrics do you use to measure project success?",
            "How do you ensure your team meets deadlines and quality standards?"
        ],
        "analyst": [
            "Describe your approach to gathering requirements from stakeholders.",
            "How do you prioritize features or improvements?",
            "What tools do you use for data analysis and reporting?",
            "How do you communicate complex findings to non-technical stakeholders?",
            "Describe a situation where your analysis led to a significant business decision."
        ],
        "designer": [
            "How do you approach the design process for a new project?",
            "Describe your experience with different design tools and software.",
            "How do you incorporate user feedback into your designs?",
            "How do you balance aesthetics with functionality?",
            "Describe a design challenge you faced and how you overcame it."
        ]
    }
    
    # Default non-technical questions
    nontechnical_questions = [
        "How do you prioritize your work when dealing with multiple deadlines?",
        "Describe a situation where you had to collaborate with a difficult team member.",
        "How do you stay updated with the latest trends and developments in your field?",
        "Describe your ideal work environment and company culture.",
        "How do you handle feedback and criticism?"
    ]
    
    # Determine which set of technical questions to use based on job title
    job_title_lower = job_title.lower()
    selected_technical_questions = []
    
    if any(keyword in job_title_lower for keyword in ["developer", "engineer", "programmer", "software", "code", "web"]):
        selected_technical_questions = technical_questions["software"]
    elif any(keyword in job_title_lower for keyword in ["data", "analytics", "scientist", "ml", "ai"]):
        selected_technical_questions = technical_questions["data"]
    elif any(keyword in job_title_lower for keyword in ["manager", "director", "lead", "head"]):
        selected_technical_questions = technical_questions["manager"]
    elif any(keyword in job_title_lower for keyword in ["analyst", "business", "product"]):
        selected_technical_questions = technical_questions["analyst"]
    elif any(keyword in job_title_lower for keyword in ["designer", "ux", "ui", "graphic"]):
        selected_technical_questions = technical_questions["designer"]
    else:
        # If no match, use a mix of questions
        selected_technical_questions = [
            technical_questions["software"][0],
            technical_questions["analyst"][0],
            technical_questions["manager"][0],
            "Describe your technical skills that are most relevant to this position.",
            "What technical challenges are you looking forward to tackling in this role?"
        ]
    
    return selected_technical_questions, nontechnical_questions

# Batch evaluate multiple resumes against the same JD
@app.route('/evaluate-batch', methods=['POST'])
def evaluate_batch():
    """Evaluate multiple resumes against the same JD and return a comparison ranking."""
    try:
        if 'resumes' not in request.files:
            return jsonify({'success': False, 'error': 'No resumes provided'}), 400

        files = request.files.getlist('resumes')
        if not files:
            return jsonify({'success': False, 'error': 'No files received'}), 400

        job_title = request.form.get('job_title')
        job_description = request.form.get('job_description')
        if not job_title or not job_description:
            return jsonify({'success': False, 'error': 'Missing job title or description'}), 400

        additional_context = request.form.get('additional_context', '').strip()
        additional_context_block = f"**Additional Context (client constraints/preference):** {additional_context}" if additional_context else ""

        results = []
        for f in files:
            if f.filename == '' or not allowed_file(f.filename):
                continue
            filename = secure_filename(f.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            f.save(file_path)

            resume_text = extract_text_from_file(file_path)
            if not resume_text:
                continue

            formatted_prompt = input_prompt_template.format(
                resume_text=resume_text,
                job_description=job_description,
                additional_context_block=additional_context_block
            )
            main_response = asyncio.run(async_gemini_generate(formatted_prompt))
            if not main_response:
                continue

            match_percentage_str = main_response.get('JD Match', '0%')
            match_percentage = int(str(match_percentage_str).strip('%') or 0)
            # Derive strengths from Match Factors if available (dict of dimension->score/comment)
            top_strengths = []
            mf = main_response.get('Match Factors', {})
            if isinstance(mf, dict) and mf:
                # If numeric scores, sort desc; else take top keys
                try:
                    sorted_items = sorted(mf.items(), key=lambda kv: float(str(kv[1]).split('%')[0]) if isinstance(kv[1], str) and '%' in kv[1] else float(kv[1]), reverse=True)
                except Exception:
                    sorted_items = list(mf.items())
                for k, v in sorted_items[:5]:
                    top_strengths.append(f"{k}")
            # Fallback: use extracted keywords from profile summary heuristics
            if not top_strengths:
                ps = main_response.get('Profile Summary', '') or ''
                words = [w.strip('.,;:()').title() for w in ps.split() if len(w) > 3]
                uniq = []
                for w in words:
                    if w not in uniq:
                        uniq.append(w)
                top_strengths = uniq[:5]

            key_gaps = list(main_response.get('MissingKeywords', [])) if isinstance(main_response.get('MissingKeywords', []), list) else []

            results.append({
                'filename': filename,
                'match_percentage': match_percentage,
                'top_strengths': top_strengths,
                'key_gaps': key_gaps
            })

        results.sort(key=lambda x: x['match_percentage'], reverse=True)
        if not results:
            return jsonify({'success': False, 'error': 'Failed to evaluate uploaded resumes'}), 500

        # Build a recruiter-style markdown comparison report (compact)
        def eval_mark(mark_score):
            if mark_score >= 75:
                return '✅'
            if mark_score >= 55:
                return '⚠️'
            return '❌'

        # JD Summary placeholder (kept short)
        md_lines = []
        md_lines.append('# 🧭 JD Summary')
        md_lines.append('(Concise summary of the role. Auto-generated placeholders — edit as needed.)')
        md_lines.append('')
        md_lines.append('| JD Pillar | Key Expectations |')
        md_lines.append('|------------|------------------|')
        md_lines.append('| Role Objective | Define and deliver measurable impact for the business |')
        md_lines.append('| Core Focus Areas | Execution, stakeholder alignment, metrics |')
        md_lines.append('| Key Competencies | Problem solving, delivery, collaboration |')
        md_lines.append('| Consulting & Client Engagement | Discovery, advisory, influence |')
        md_lines.append('| AI / Analytics / Domain | Practical awareness and usage |')
        md_lines.append('| Cultural Fit | Ownership, clarity, bias for action |')
        md_lines.append('\n---\n')

        # Per-candidate sections
        for r in results:
            name = r['filename']
            md_lines.append('# 🧩 Candidate Summary')
            md_lines.append(f'**Name:** {name}')
            md_lines.append('**Current Role:** —')
            md_lines.append('**Experience:** —')
            md_lines.append('**Industry / Domain Expertise:** —')
            md_lines.append('**Education:** —')
            md_lines.append('**Location:** —')
            md_lines.append(f"**Key Themes / Keywords:** {', '.join(r.get('top_strengths', [])[:6]) or '—'}")
            md_lines.append('\n---\n')

            # Ensure non-empty strengths/gaps
            if not r.get('top_strengths'):
                base = os.path.splitext(name)[0]
                heur = [w.title() for w in base.replace('_',' ').replace('-',' ').split() if len(w) > 2][:3]
                r['top_strengths'] = heur or ['General delivery', 'Stakeholder collaboration']
            if not r.get('key_gaps'):
                r['key_gaps'] = ['No critical gap surfaced']

            md_lines.append('# 📊 Comparative Fit Analysis (JD vs Resume)')
            md_lines.append('| **Dimension** | **Evaluation** | **Commentary** |')
            md_lines.append('|----------------|----------------|----------------|')
            mark = eval_mark(r['match_percentage'])
            sig = ", ".join(r.get("top_strengths", [])[:3]) or '—'
            gap_one = (r.get('key_gaps') or ['—'])[0]
            md_lines.append(f'| Domain Expertise | {mark} | Signals: {sig} |')
            md_lines.append(f'| Consulting & Advisory Orientation | {mark} | Based on profile narrative |')
            md_lines.append(f'| AI / Analytics Awareness | {mark} | Tooling/awareness inferred |')
            md_lines.append(f'| Account Growth / Leadership | {mark} | Team/initiative ownership |')
            md_lines.append(f'| Client Gravitas (C-suite Influence) | {mark} | Stakeholder influence indicators |')
            md_lines.append(f'| Communication & Storytelling | {mark} | Clarity of outcomes |')
            md_lines.append(f'| Technical or Delivery Depth | {mark} | Depth vs breadth balance |')
            md_lines.append(f'| Cultural Fit (Consulting + Innovation) | {mark} | Bias for action, collaboration |')
            md_lines.append('')

            md_lines.append('# 💪 Key Strengths')
            strengths = r.get('top_strengths', [])[:5] or ['General delivery', 'Collaboration']
            for s in strengths:
                md_lines.append(f'- {s}')
            md_lines.append('')

            md_lines.append('# ⚠️ Gaps / Risks')
            md_lines.append('| Gap | Explanation | Impact |')
            md_lines.append('|------|-------------|---------|')
            gaps = r.get('key_gaps', [])[:3] or ['No critical gap surfaced']
            for g in gaps:
                md_lines.append(f'| {g} | — | Medium |')
            md_lines.append('')

            md_lines.append('# 🧾 Scorecard Summary')
            def to_star(score):
                # Map 0-100 to 1-5
                return max(1, min(5, round(score/20)))
            star = to_star(r['match_percentage'])
            md_lines.append('| Category | Rating (1–5) | Comment |')
            md_lines.append('|-----------|--------------|----------|')
            for cat in ['Domain Fit','Consulting Gravitas','AI / Analytics Awareness','Account Growth Leadership','Client Relationship / Communication','Cultural Fit']:
                md_lines.append(f'| {cat} | {star} | Derived from resume signals |')
            overall10 = round(r['match_percentage']/10, 1)
            verdict = '✅ Strong Fit' if r['match_percentage']>=75 else ('⚠️ Partial Fit' if r['match_percentage']>=55 else '❌ Not a Fit')
            md_lines.append('')
            md_lines.append(f'**Overall Fit Score:** {overall10} / 10  ')
            md_lines.append(f'**Verdict:** {verdict}')
            md_lines.append('\n---\n')

            md_lines.append('# ✅ Final Recruiter Verdict')
            md_lines.append('> Candidate shows relevant capability signals with room to validate consulting gravitas and delivery depth. Recommend next-step screening focused on stakeholder influence, structured problem solving, and measurable impact.')
            md_lines.append('')

        # Summary comparison table
        md_lines.append('# ⚖️ Multi-Candidate Comparison Table')
        header = '| **Criteria** | ' + ' | '.join([r['filename'] for r in results]) + ' |'
        sep = '|---------------|' + '|'.join(['----------------' for _ in results]) + '|'
        md_lines.append(header)
        md_lines.append(sep)
        def row_line(label):
            vals = []
            for r in results:
                vals.append(f"{round(r['match_percentage']/10,1)} / 10")
            return f"| {label} | " + " | ".join(vals) + " |"
        for crit in ['Domain Relevance','Consulting Orientation','AI / Analytics Exposure','Client Growth Aptitude','Cultural Fit','**Overall Fit Score**']:
            md_lines.append(row_line(crit))

        report_markdown = '\n'.join(md_lines)

        return jsonify({'success': True, 'results': results, 'report_markdown': report_markdown})
    except Exception as e:
        logging.error(f"Error in evaluate_batch: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

if __name__ == "__main__":
    # Initialize logging
    logging.basicConfig(level=logging.INFO,
                       format='%(asctime)s - %(levelname)s - %(message)s')

    # Initialize database
    init_db()
    
    # Update database schema
    update_db_schema()
    
    try:
        # Initialize Pinecone safely
        vectorstore = initialize_pinecone()
        
        # Build BM25 index
        logging.info("🔍 Building BM25 index...")
        build_bm25_index(POLICIES_FOLDER)
        
        # Set up LLM and QA chain
        logging.info("🤖 Setting up LLM and QA chain...")
        llm, qa_chain, retriever = setup_llm_chain()
        
        # Start Flask server with ASGI support using hypercorn
        logging.info("🌐 Starting server...")
        from hypercorn.config import Config
        from hypercorn.asyncio import serve

        config = Config()
        config.bind = ["localhost:5000"]
        config.use_reloader = True
        
        asyncio.run(serve(asgi_app, config))
        
    except Exception as e:
        logging.error(f"❌ Startup error: {str(e)}")
        raise
